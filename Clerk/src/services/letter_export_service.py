"""
Service for exporting Good Faith letters to various formats.

Handles PDF, DOCX, and HTML export with proper formatting.
"""

import io
from datetime import datetime

from reportlab.lib.pagesizes import letter as PAGE_SIZE
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown

from src.models.deficiency_models import GeneratedLetter
from src.utils.logger import get_logger

logger = get_logger("letter_export_service")


class LetterExportService:
    """Handles export of letters to various formats."""

    def __init__(self):
        """Initialize export service."""
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()

    def _setup_custom_styles(self):
        """Setup custom styles for PDF generation."""
        # Legal letter style
        self.styles.add(
            ParagraphStyle(
                name="LegalNormal",
                parent=self.styles["Normal"],
                fontSize=11,
                leading=14,
                alignment=TA_JUSTIFY,
                spaceAfter=12,
            )
        )

        # Letterhead style
        self.styles.add(
            ParagraphStyle(
                name="Letterhead",
                parent=self.styles["Heading1"],
                fontSize=14,
                alignment=TA_CENTER,
                spaceAfter=24,
            )
        )

        # Date style
        self.styles.add(
            ParagraphStyle(
                name="DateLine",
                parent=self.styles["Normal"],
                fontSize=11,
                spaceAfter=18,
            )
        )

        # Signature style
        self.styles.add(
            ParagraphStyle(
                name="Signature",
                parent=self.styles["Normal"],
                fontSize=11,
                spaceAfter=6,
            )
        )

    async def export_to_pdf(self, letter: GeneratedLetter) -> bytes:
        """
        Export letter to PDF format.

        Args:
            letter: GeneratedLetter to export

        Returns:
            bytes: PDF content
        """
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=PAGE_SIZE,
            rightMargin=1 * inch,
            leftMargin=1 * inch,
            topMargin=1 * inch,
            bottomMargin=1 * inch,
        )

        # Create story for PDF
        story = []

        # Add letterhead if in metadata
        if letter.metadata.get("letterhead"):
            letterhead = letter.metadata["letterhead"]
            story.append(Paragraph(letterhead["firm_name"], self.styles["Letterhead"]))
            if letterhead.get("address"):
                story.append(
                    Paragraph(letterhead["address"], self.styles["Letterhead"])
                )
            story.append(Spacer(1, 0.5 * inch))

        # Add date
        date_str = datetime.utcnow().strftime("%B %d, %Y")
        story.append(Paragraph(date_str, self.styles["DateLine"]))

        # Process letter content
        # Split by double newlines to identify paragraphs
        paragraphs = letter.content.split("\n\n")

        for para in paragraphs:
            if para.strip():
                # Handle special formatting
                if para.startswith("RE:"):
                    # Subject line - bold
                    story.append(
                        Paragraph(f"<b>{para}</b>", self.styles["LegalNormal"])
                    )
                elif para.strip().endswith(":") and len(para.strip()) < 50:
                    # Likely a salutation or section header
                    story.append(Paragraph(para, self.styles["LegalNormal"]))
                else:
                    # Regular paragraph
                    # Convert markdown-style formatting
                    formatted_para = self._convert_markdown_to_reportlab(para)
                    story.append(Paragraph(formatted_para, self.styles["LegalNormal"]))

                story.append(Spacer(1, 12))

        # Build PDF
        try:
            doc.build(story)
            buffer.seek(0)
            pdf_content = buffer.read()

            logger.info(
                f"Generated PDF for letter {letter.id} ({len(pdf_content)} bytes)"
            )
            return pdf_content

        except Exception as e:
            logger.error(f"PDF generation failed: {str(e)}")
            raise
        finally:
            buffer.close()

    async def export_to_docx(self, letter: GeneratedLetter) -> bytes:
        """
        Export letter to DOCX format.

        Args:
            letter: GeneratedLetter to export

        Returns:
            bytes: DOCX content
        """
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add letterhead if configured
        if letter.metadata.get("letterhead"):
            letterhead = letter.metadata["letterhead"]

            # Firm name
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(letterhead["firm_name"])
            run.font.size = Pt(14)
            run.font.bold = True

            # Address
            if letterhead.get("address"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(letterhead["address"]).font.size = Pt(11)

            doc.add_paragraph()  # Blank line

        # Add date
        date_str = datetime.utcnow().strftime("%B %d, %Y")
        p = doc.add_paragraph(date_str)
        p.add_run().add_break()

        # Process letter content
        paragraphs = letter.content.split("\n\n")

        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph()

                # Handle special formatting
                if para.startswith("RE:"):
                    # Subject line - bold
                    run = p.add_run(para)
                    run.font.bold = True
                    run.font.size = Pt(11)
                else:
                    # Regular paragraph
                    # Process inline formatting
                    self._add_formatted_text_to_docx(p, para)

                # Set paragraph formatting
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.space_after = Pt(12)

        # Save to buffer
        buffer = io.BytesIO()
        try:
            doc.save(buffer)
            buffer.seek(0)
            docx_content = buffer.read()

            logger.info(
                f"Generated DOCX for letter {letter.id} ({len(docx_content)} bytes)"
            )
            return docx_content

        except Exception as e:
            logger.error(f"DOCX generation failed: {str(e)}")
            raise
        finally:
            buffer.close()

    async def export_to_html(self, letter: GeneratedLetter) -> str:
        """
        Export letter to HTML format.

        Args:
            letter: GeneratedLetter to export

        Returns:
            str: HTML content
        """
        # Convert markdown-style formatting to HTML
        html_content = markdown.markdown(
            letter.content, extensions=["nl2br", "tables", "fenced_code"]
        )

        # Build full HTML document
        letterhead_html = ""
        if letter.metadata.get("letterhead"):
            letterhead = letter.metadata["letterhead"]
            letterhead_html = f"""
            <div class="letterhead">
                <h1>{letterhead["firm_name"]}</h1>
                {"<p>" + letterhead.get("address", "") + "</p>" if letterhead.get("address") else ""}
            </div>
            """

        template = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Good Faith Letter - {letter.case_name}</title>
    <style>
        body {{
            font-family: 'Times New Roman', Times, serif;
            font-size: 11pt;
            line-height: 1.5;
            max-width: 8.5in;
            margin: 1in auto;
            padding: 0 1in;
            color: #000;
            background: #fff;
        }}
        .letterhead {{
            text-align: center;
            margin-bottom: 2em;
        }}
        .letterhead h1 {{
            font-size: 14pt;
            margin: 0;
        }}
        .date {{
            margin-bottom: 1.5em;
        }}
        p {{
            text-align: justify;
            margin-bottom: 1em;
        }}
        .signature-block {{
            margin-top: 3em;
        }}
        @media print {{
            body {{
                margin: 0;
                padding: 0;
            }}
        }}
    </style>
</head>
<body>
    {letterhead_html}
    <div class="date">{datetime.utcnow().strftime("%B %d, %Y")}</div>
    <div class="letter-content">
        {html_content}
    </div>
</body>
</html>
        """

        logger.info(f"Generated HTML for letter {letter.id}")
        return template.strip()

    def _convert_markdown_to_reportlab(self, text: str) -> str:
        """
        Convert markdown formatting to ReportLab XML tags.

        Args:
            text: Text with markdown formatting

        Returns:
            str: Text with ReportLab XML formatting
        """
        # Bold: **text** or __text__ -> <b>text</b>
        import re

        text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
        text = re.sub(r"__(.+?)__", r"<b>\1</b>", text)

        # Italic: *text* or _text_ -> <i>text</i>
        text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
        text = re.sub(r"_(.+?)_", r"<i>\1</i>", text)

        # Escape XML special characters
        text = text.replace("&", "&amp;")
        text = text.replace("<", "&lt;").replace(">", "&gt;")

        # Restore our formatting tags
        text = text.replace("&lt;b&gt;", "<b>").replace("&lt;/b&gt;", "</b>")
        text = text.replace("&lt;i&gt;", "<i>").replace("&lt;/i&gt;", "</i>")

        return text

    def _add_formatted_text_to_docx(self, paragraph, text: str):
        """
        Add text with inline formatting to DOCX paragraph.

        Args:
            paragraph: python-docx paragraph object
            text: Text to add with potential formatting
        """
        import re

        # Pattern to find formatted sections
        pattern = r"(\*\*(.+?)\*\*)|(__(.+?)__)|(\*(.+?)\*)|(_(.+?)_)|([^*_]+)"

        for match in re.finditer(pattern, text):
            if match.group(1):  # **bold**
                run = paragraph.add_run(match.group(2))
                run.font.bold = True
            elif match.group(3):  # __bold__
                run = paragraph.add_run(match.group(4))
                run.font.bold = True
            elif match.group(5):  # *italic*
                run = paragraph.add_run(match.group(6))
                run.font.italic = True
            elif match.group(7):  # _italic_
                run = paragraph.add_run(match.group(8))
                run.font.italic = True
            else:  # Plain text
                paragraph.add_run(match.group(9))

            # Set font for all runs
            if hasattr(paragraph.runs[-1], "font"):
                paragraph.runs[-1].font.size = Pt(11)
                paragraph.runs[-1].font.name = "Times New Roman"
