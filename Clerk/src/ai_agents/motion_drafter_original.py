"""
AI Motion Drafting Agent - Enhanced Version with Direct Database Access
Implements section-by-section legal motion generation following best practices
"""

import asyncio
import logging
from typing import List, Dict, Any, Optional
from datetime import datetime
from dataclasses import dataclass, field
from enum import Enum
import re
import json

from pydantic_ai import Agent
from pydantic_ai.models.openai import OpenAIModel
from openai import AsyncOpenAI
import tiktoken

from src.vector_storage.qdrant_store import QdrantVectorStore
from src.vector_storage.embeddings import EmbeddingGenerator
from src.ai_agents.motion_cache_manager import motion_cache, context_cache
from src.ai_agents.outline_cache_manager import outline_cache
from src.utils.timeout_monitor import TimeoutMonitor, ProgressTracker
from config.settings import settings

logger = logging.getLogger(__name__)


class SectionType(Enum):
    """Types of sections in a legal motion"""

    INTRODUCTION = "introduction"
    STATEMENT_OF_FACTS = "statement_of_facts"
    LEGAL_STANDARD = "legal_standard"
    ARGUMENT = "argument"
    SUB_ARGUMENT = "sub_argument"
    CONCLUSION = "conclusion"
    PRAYER_FOR_RELIEF = "prayer_for_relief"


class DocumentLength(Enum):
    """Target document lengths with expanded ranges"""

    SHORT = (15, 20)  # 15-20 pages
    MEDIUM = (20, 30)  # 20-30 pages
    LONG = (30, 40)  # 30-40 pages
    COMPREHENSIVE = (35, 50)  # 35-50 pages for complex motions


@dataclass
class OutlineSection:
    """Enhanced outline section with better structure"""

    id: str
    title: str
    section_type: SectionType
    content_points: List[str]
    legal_authorities: List[str]
    target_length: int  # Target word count
    parent_id: Optional[str] = None
    children: List["OutlineSection"] = field(default_factory=list)
    context_summary: Optional[str] = None
    hook_options: List[str] = field(default_factory=list)
    themes: List[str] = field(default_factory=list)
    key_facts: List[Dict[str, Any]] = field(default_factory=list)
    counter_arguments: List[Dict[str, str]] = field(default_factory=list)


@dataclass
class DraftedSection:
    """Enhanced drafted section with quality metrics"""

    outline_section: OutlineSection
    content: str
    word_count: int
    citations_used: List[str]
    citations_verified: Dict[str, bool]  # Citation -> verification status
    expansion_cycles: int
    confidence_score: float
    needs_revision: bool = False
    revision_notes: List[str] = field(default_factory=list)
    consistency_score: float = 0.0
    factual_accuracy_score: float = 0.0
    argument_strength_score: float = 0.0
    transitions: Dict[str, str] = field(default_factory=dict)  # To next/from previous


@dataclass
class MotionDraft:
    """Enhanced motion draft with comprehensive metadata"""

    title: str
    case_name: str  # Can be derived from database_name or set as a display value
    sections: List[DraftedSection]
    total_word_count: int
    total_page_estimate: int
    creation_timestamp: datetime
    outline_source: Dict[str, Any]
    coherence_score: float = 0.0
    review_notes: List[str] = field(default_factory=list)
    quality_metrics: Dict[str, float] = field(default_factory=dict)
    citation_index: Dict[str, List[str]] = field(
        default_factory=dict
    )  # Citation -> sections using it
    consistency_issues: List[Dict[str, Any]] = field(default_factory=list)


class EnhancedMotionDraftingAgent:
    """Enhanced AI agent for drafting legal motions with direct database access"""

    def __init__(self):
        """Initialize the enhanced motion drafting agent"""
        self.vector_store = QdrantVectorStore()
        self.embedding_generator = EmbeddingGenerator()
        self.openai_client = AsyncOpenAI(api_key=settings.openai.api_key)

        # Initialize OpenAI models with proper configuration
        # Note: If using Claude models through OpenAI-compatible API, set base_url
        self.gpt4_model = OpenAIModel("gpt-4.1-mini-2025-04-14")

        # For now, use GPT-4 for all agents to ensure compatibility
        self.primary_model = self.gpt4_model

        # Initialize AI agents with enhanced prompts
        self.section_writer = self._create_enhanced_section_writer()
        self.section_expander = self._create_enhanced_section_expander()
        self.transition_writer = self._create_enhanced_transition_writer()
        self.document_reviewer = self._create_enhanced_document_reviewer()
        self.consistency_checker = self._create_consistency_checker()
        self.citation_verifier = self._create_citation_verifier()

        # Initialize tokenizer
        self.tokenizer = tiktoken.encoding_for_model("gpt-4")

        # Enhanced configuration
        self.words_per_page = 250
        self.max_expansion_cycles = 5
        self.min_confidence_threshold = 0.75
        self.citation_patterns = self._compile_citation_patterns()

        # Track document-wide context
        self.document_context = {
            "themes": [],
            "key_arguments": [],
            "terminology": {},
            "citations_used": set(),
            "fact_chronology": [],
        }

        logger.info("EnhancedMotionDraftingAgent initialized successfully")

    def _create_enhanced_section_writer(self) -> Agent:
        """Create enhanced section writer with ABCDE framework and CoT prompting"""
        try:
            # Use primary model for section writing
            agent = Agent(
                self.primary_model,
                system_prompt="""You are an expert legal writer specializing in comprehensive motion drafting.

## AUDIENCE
You are writing for a trial court judge who needs detailed legal analysis, not summaries.
The document will be reviewed by opposing counsel looking for weaknesses.
Your writing must be authoritative, precise, and persuasive.

## BACKGROUND
You are drafting sections of formal legal motions (15-40 pages) that will be filed in court.
Each section must contribute substantive content to reach the target length.
This is NOT a brief or summary - it requires comprehensive analysis.

## CONSTRAINTS
1. Use formal legal writing style appropriate for court filings
2. Follow IRAC/CRAC structure (Issue, Rule, Application, Conclusion)
3. All citations must be in proper Bluebook format
4. Never create fictional citations - only use provided authorities
5. Maintain consistent terminology throughout
6. Each section must meet its target word count

## DETAILED PARAMETERS
- Write comprehensively, not concisely
- Develop each point with multiple paragraphs of analysis
- Include detailed factual application to legal standards
- Address counterarguments preemptively
- Use topic sentences and smooth transitions
- Incorporate policy considerations where relevant
- Analyze case law holdings in detail
- Apply multi-factor tests thoroughly

## EVALUATION CRITERIA
Your output will be evaluated on:
1. Completeness - Did you address all outline points?
2. Depth - Is the analysis thorough and detailed?
3. Length - Does it meet the target word count?
4. Legal accuracy - Are citations and law correctly stated?
5. Persuasiveness - Is the argument compelling?
6. Coherence - Does it flow logically?

## CHAIN-OF-THOUGHT PROCESS
Before writing, think through:
1. What is the core legal issue this section addresses?
2. What are the relevant legal standards and elements?
3. How do the facts of our case apply to each element?
4. What would opposing counsel argue?
5. How can we preemptively counter those arguments?
6. What policy reasons support our position?""",
                result_type=str,
            )
            logger.info("Section writer agent created successfully")
            return agent
        except Exception as e:
            logger.error(f"Failed to create section writer agent: {str(e)}")
            raise

    def _create_enhanced_section_expander(self) -> Agent:
        """Create enhanced expander with explicit anti-brevity instructions"""
        try:
            agent = Agent(
                self.primary_model,
                system_prompt="""You are an expert legal writer focused on expanding and enriching legal arguments.

## YOUR MISSION
Transform concise legal sections into comprehensive, detailed analyses suitable for 20-40 page motions.
You are the opposite of a summarizer - you ADD substance, never remove it.

## EXPANSION TECHNIQUES
1. **Deepen Legal Analysis**
   - Explain the evolution of legal standards
   - Discuss split authorities and circuit conflicts
   - Analyze legislative history and intent
   - Examine policy rationales behind rules

2. **Enhance Factual Application**
   - Apply facts to EACH element separately
   - Use multiple factual examples per point
   - Create detailed factual analogies to case law
   - Distinguish opposing cases factually

3. **Strengthen Arguments**
   - Add "even if" alternative arguments
   - Include "moreover" and "furthermore" points
   - Develop subsidiary arguments
   - Address potential weaknesses proactively

4. **Add Supporting Material**
   - Incorporate relevant statutory text
   - Quote key language from controlling cases
   - Add parenthetical explanations for citations
   - Include see also citations for additional support

## CRITICAL REQUIREMENTS
- NEVER summarize or condense
- NEVER use phrases like "in summary" or "briefly stated"
- ALWAYS add new substantive content
- Target at least 40% increase in word count per expansion
- Maintain formal legal writing style
- Preserve all existing content while adding new material""",
                result_type=str,
            )
            logger.info("Section expander agent created successfully")
            return agent
        except Exception as e:
            logger.error(f"Failed to create section expander agent: {str(e)}")
            raise

    def _create_enhanced_transition_writer(self) -> Agent:
        """Create transition writer for document coherence"""
        try:
            agent = Agent(
                self.primary_model,
                system_prompt="""You are an expert legal writer specializing in document coherence and flow.

Create sophisticated transitions that:
1. **Summarize** - Crystallize the key holding/conclusion from the previous section
2. **Connect** - Show logical relationship to the next argument
3. **Preview** - Introduce what's coming next
4. **Maintain Momentum** - Keep the reader engaged and moving forward

Transition types to use:
- Logical progression: "Having established X, we now turn to Y"
- Building arguments: "This principle extends further when we consider"
- Contrast: "While X is true, it does not preclude Y"
- Causation: "As a direct result of X, Y necessarily follows"

Requirements:
- 2-3 paragraphs for major section transitions
- 1 paragraph for subsection transitions
- Reference specific holdings or facts
- Use sophisticated connecting language
- Maintain formal legal tone""",
                result_type=str,
            )
            logger.info("Transition writer agent created successfully")
            return agent
        except Exception as e:
            logger.error(f"Failed to create transition writer agent: {str(e)}")
            raise

    def _create_enhanced_document_reviewer(self) -> Agent:
        """Create comprehensive document reviewer"""
        try:
            agent = Agent(
                self.primary_model,
                system_prompt="""You are a senior legal editor and former appellate judge reviewing motions.

Conduct a comprehensive review examining:

## LEGAL ACCURACY
- Verify all citations are properly formatted
- Confirm legal standards are correctly stated
- Check that case holdings are accurately represented
- Ensure statutory language is quoted correctly

## ARGUMENT QUALITY
- Assess logical flow and persuasiveness
- Identify gaps in reasoning
- Evaluate strength of factual application
- Check for unaddressed counterarguments

## CONSISTENCY
- Terminology usage throughout document
- Consistent characterization of facts
- Uniform citation format
- Coherent theory of the case

## COMPLETENESS
- All outline points addressed
- Sufficient detail and analysis
- Adequate factual support
- Comprehensive legal authority

## PROFESSIONALISM
- Appropriate tone for court
- No inflammatory language
- Respectful treatment of opposing party
- Clear and precise language

Provide output as JSON with:
- readiness_score (1-10)
- section_scores (dict of section_id: score)
- consistency_issues (list of specific problems)
- missing_elements (list of what's missing)
- revision_priorities (ordered list of what needs fixing)
- specific_edits (dict of section_id: list of edits)""",
                result_type=Dict[str, Any],
            )
            logger.info("Document reviewer agent created successfully")
            return agent
        except Exception as e:
            logger.error(f"Failed to create document reviewer agent: {str(e)}")
            raise

    def _create_consistency_checker(self) -> Agent:
        """Create agent for checking document consistency"""
        try:
            agent = Agent(
                self.primary_model,
                system_prompt="""You are a legal document consistency analyzer.

Check for inconsistencies in:
1. **Terminology** - Same terms used throughout (Agreement vs Contract)
2. **Facts** - Dates, names, amounts consistent
3. **Legal positions** - No contradictory arguments
4. **Citations** - Same cases cited consistently
5. **Party names** - Consistent references to parties

Output JSON with:
- inconsistencies: list of {type, location, issue, suggestion}
- terminology_map: dict of preferred terms
- fact_conflicts: list of conflicting facts
- citation_issues: list of citation problems""",
                result_type=Dict[str, Any],
            )
            logger.info("Consistency checker agent created successfully")
            return agent
        except Exception as e:
            logger.error(f"Failed to create consistency checker agent: {str(e)}")
            raise

    def _create_citation_verifier(self) -> Agent:
        """Create agent for citation verification"""
        try:
            agent = Agent(
                self.primary_model,
                system_prompt="""You are a legal citation verification specialist.

Verify that:
1. All citations follow Bluebook format
2. Case names are properly formatted
3. Pincites are included where needed
4. Parentheticals explain relevance
5. Signals (see, see also, cf.) are used correctly

Check against the provided case documents to ensure citations are real.

Output JSON with:
- verified_citations: dict of citation: boolean
- format_issues: list of formatting problems
- missing_pincites: list of citations needing page numbers
- fictional_citations: list of any made-up citations""",
                result_type=Dict[str, Any],
            )
            logger.info("Citation verifier agent created successfully")
            return agent
        except Exception as e:
            logger.error(f"Failed to create citation verifier agent: {str(e)}")
            raise

    async def draft_motion(
        self,
        outline: Dict[str, Any],
        database_name: str,
        target_length: DocumentLength = DocumentLength.MEDIUM,
        motion_title: Optional[str] = None,
        opposing_motion_text: Optional[str] = None,
    ) -> MotionDraft:
        """
        Enhanced motion drafting with direct database access

        Args:
            outline: Structured outline from the outline generation phase
            database_name: Name of the database/collection containing case documents
            target_length: Target document length (SHORT, MEDIUM, LONG, COMPREHENSIVE)
            motion_title: Optional title for the motion

        Returns:
            Complete MotionDraft object
        """
        try:
            start_time = datetime.utcnow()
            logger.info(
                f"[MOTION_DRAFTER] Starting enhanced motion draft for database: {database_name} at {start_time}"
            )
            logger.info(
                f"[MOTION_DRAFTER] Target length: {target_length}, Motion title: {motion_title}"
            )
            logger.info(f"[MOTION_DRAFTER] Outline keys: {list(outline.keys())}")

            # Initialize timeout monitor and progress tracker
            timeout_monitor = TimeoutMonitor(
                operation_name=f"Motion Drafting ({database_name})",
                warning_threshold=60,  # 1 minute warning
                critical_threshold=120,  # 2 minute critical warning
            )
            timeout_monitor.log_progress("Motion drafting started")

            # Initialize document context
            logger.info("[MOTION_DRAFTER] Initializing document context")
            self.document_context = {
                "themes": outline.get("themes", []),
                "key_arguments": [],
                "terminology": {},
                "citations_used": set(),
                "fact_chronology": [],
                "database_name": database_name,
                "opposing_motion_text": opposing_motion_text,
            }
            logger.info(
                f"[MOTION_DRAFTER] Document themes: {self.document_context['themes']}"
            )

            # Parse outline with enhanced structure
            logger.info("[MOTION_DRAFTER] Parsing outline structure")
            timeout_monitor.log_progress("Parsing outline structure")
            outline_sections = self._parse_enhanced_outline(outline, target_length)
            logger.info(
                f"[MOTION_DRAFTER] Parsed {len(outline_sections)} sections from outline"
            )
            for i, section in enumerate(outline_sections):
                logger.info(
                    f"[MOTION_DRAFTER] Section {i + 1}: {section.title} (target: {section.target_length} words)"
                )

            # Initialize progress tracker for sections
            progress_tracker = ProgressTracker(
                total_steps=len(outline_sections)
                + 4,  # sections + context + review + consistency + finalization
                operation_name=f"Motion Drafting ({database_name})",
            )

            # Retrieve comprehensive case context directly from database
            logger.info(
                f"[MOTION_DRAFTER] Retrieving case context from database: {database_name}"
            )
            context_start_time = datetime.utcnow()
            timeout_monitor.log_progress(
                "Retrieving case context", {"database": database_name}
            )
            progress_tracker.next_step(
                "Retrieving case context", f"Database: {database_name}"
            )

            case_context = await self._retrieve_enhanced_case_context(
                database_name, outline_sections
            )

            context_end_time = datetime.utcnow()
            logger.info(
                f"[MOTION_DRAFTER] Case context retrieval completed in {context_end_time - context_start_time}"
            )
            logger.info(
                f"[MOTION_DRAFTER] Retrieved context: {len(case_context.get('case_facts', []))} facts, {len(case_context.get('legal_authorities', []))} authorities"
            )
            timeout_monitor.log_progress(
                "Case context retrieved",
                {
                    "facts": len(case_context.get("case_facts", [])),
                    "authorities": len(case_context.get("legal_authorities", [])),
                },
            )

            # Store initial context in cache
            motion_id = f"{database_name}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}"
            logger.info(
                f"[MOTION_DRAFTER] Storing initial context in cache with motion_id: {motion_id}"
            )
            try:
                await context_cache.store_section_context(
                    motion_id=motion_id,
                    section_id="initial_context",
                    context={
                        "case_context": case_context,
                        "themes": self.document_context["themes"],
                        "outline": outline,
                        "database_name": database_name,
                    },
                )
                logger.info("[MOTION_DRAFTER] Context cached successfully")
            except Exception as e:
                logger.warning(f"[MOTION_DRAFTER] Failed to cache context: {str(e)}")

            # Draft sections with progressive context building
            logger.info("[MOTION_DRAFTER] Starting section drafting process")
            drafted_sections = []
            total_words = 0
            sections_start_time = datetime.utcnow()

            for i, section in enumerate(outline_sections):
                section_start_time = datetime.utcnow()
                logger.info(
                    f"[MOTION_DRAFTER] === Drafting section {i + 1}/{len(outline_sections)}: {section.title} ==="
                )
                logger.info(
                    f"[MOTION_DRAFTER] Section target length: {section.target_length} words"
                )
                logger.info(
                    f"[MOTION_DRAFTER] Section type: {section.section_type.value}"
                )

                # Update progress tracking
                timeout_monitor.log_progress(
                    f"Drafting section {i + 1}: {section.title}",
                    {
                        "section_type": section.section_type.value,
                        "target_length": section.target_length,
                    },
                )
                progress_tracker.next_step(
                    f"Section {i + 1}: {section.title}",
                    f"Type: {section.section_type.value}, Target: {section.target_length} words",
                )

                # Build cumulative context
                logger.info(
                    f"[MOTION_DRAFTER] Building cumulative context for section {i + 1}"
                )
                cumulative_context = self._build_cumulative_context(
                    drafted_sections, case_context
                )
                logger.info(
                    f"[MOTION_DRAFTER] Cumulative context: {len(cumulative_context.get('citations_used', set()))} citations used so far"
                )

                # Draft with Chain-of-Thought approach
                logger.info(
                    f"[MOTION_DRAFTER] Starting Chain-of-Thought drafting for section {i + 1}"
                )
                cot_start_time = datetime.utcnow()

                try:
                    drafted_section = await self._draft_section_with_cot(
                        section, case_context, cumulative_context
                    )
                    cot_end_time = datetime.utcnow()
                    logger.info(
                        f"[MOTION_DRAFTER] CoT drafting completed in {cot_end_time - cot_start_time}"
                    )
                    logger.info(
                        f"[MOTION_DRAFTER] Initial draft: {drafted_section.word_count} words, confidence: {drafted_section.confidence_score:.2f}"
                    )
                except Exception as e:
                    logger.error(
                        f"[MOTION_DRAFTER] Error drafting section {i + 1}: {str(e)}",
                        exc_info=True,
                    )
                    # Create a placeholder section on error
                    drafted_section = DraftedSection(
                        outline_section=section,
                        content=f"[ERROR DRAFTING SECTION: {str(e)}]",
                        word_count=0,
                        citations_used=[],
                        citations_verified={},
                        expansion_cycles=0,
                        confidence_score=0.0,
                    )

                # Expand to meet length requirements
                logger.info(
                    f"[MOTION_DRAFTER] Checking if expansion needed (current: {drafted_section.word_count}, target: {section.target_length})"
                )
                expansion_cycles = 0
                previous_word_count = drafted_section.word_count

                while drafted_section.word_count < section.target_length * 0.9:
                    expansion_cycles += 1
                    logger.info(
                        f"[MOTION_DRAFTER] Starting expansion cycle {expansion_cycles} for section {i + 1}"
                    )
                    expansion_start_time = datetime.utcnow()

                    try:
                        drafted_section = await self._expand_section_intelligently(
                            drafted_section,
                            section.target_length,
                            case_context,
                            cumulative_context,
                        )

                        expansion_end_time = datetime.utcnow()
                        logger.info(
                            f"[MOTION_DRAFTER] Expansion cycle {expansion_cycles} completed in {expansion_end_time - expansion_start_time}"
                        )
                        logger.info(
                            f"[MOTION_DRAFTER] After expansion: {drafted_section.word_count} words"
                        )
                    except Exception as e:
                        logger.error(
                            f"[MOTION_DRAFTER] Error in expansion cycle {expansion_cycles}: {str(e)}"
                        )
                        break

                    # Check if expansion actually added words
                    if drafted_section.word_count <= previous_word_count:
                        logger.warning(
                            f"[MOTION_DRAFTER] No words added in expansion cycle {expansion_cycles}, breaking to prevent infinite loop"
                        )
                        break

                    # Update previous word count for next iteration
                    previous_word_count = drafted_section.word_count

                    # Use local expansion_cycles counter instead of drafted_section.expansion_cycles
                    if expansion_cycles >= self.max_expansion_cycles:
                        logger.warning(
                            f"[MOTION_DRAFTER] Max expansion cycles ({self.max_expansion_cycles}) reached for {section.title}"
                        )
                        break

                # Generate sophisticated transitions
                logger.info(
                    f"[MOTION_DRAFTER] Generating transitions for section {i + 1}"
                )
                if i > 0:
                    logger.info(
                        "[MOTION_DRAFTER] Creating transition from previous section"
                    )
                    transition_start_time = datetime.utcnow()
                    try:
                        transition_from_previous = (
                            await self._create_sophisticated_transition(
                                drafted_sections[-1], drafted_section
                            )
                        )
                        drafted_section.transitions["from_previous"] = (
                            transition_from_previous
                        )
                        transition_end_time = datetime.utcnow()
                        logger.info(
                            f"[MOTION_DRAFTER] Transition from previous created in {transition_end_time - transition_start_time}"
                        )
                    except Exception as e:
                        logger.error(
                            f"[MOTION_DRAFTER] Error creating transition: {str(e)}"
                        )

                if i < len(outline_sections) - 1:
                    logger.info(
                        "[MOTION_DRAFTER] Creating preview transition to next section"
                    )
                    # Preview transition to next section
                    preview_start_time = datetime.utcnow()
                    try:
                        transition_to_next = await self._preview_next_section(
                            drafted_section, outline_sections[i + 1]
                        )
                        drafted_section.transitions["to_next"] = transition_to_next
                        preview_end_time = datetime.utcnow()
                        logger.info(
                            f"[MOTION_DRAFTER] Preview transition created in {preview_end_time - preview_start_time}"
                        )
                    except Exception as e:
                        logger.error(
                            f"[MOTION_DRAFTER] Error creating preview transition: {str(e)}"
                        )

                # Check section consistency
                logger.info(
                    f"[MOTION_DRAFTER] Checking section consistency for section {i + 1}"
                )
                consistency_start_time = datetime.utcnow()
                try:
                    consistency_issues = await self._check_section_consistency(
                        drafted_section, drafted_sections, self.document_context
                    )
                    consistency_end_time = datetime.utcnow()
                    logger.info(
                        f"[MOTION_DRAFTER] Consistency check completed in {consistency_end_time - consistency_start_time}"
                    )

                    if consistency_issues:
                        logger.info(
                            f"[MOTION_DRAFTER] Found {len(consistency_issues)} consistency issues, resolving..."
                        )
                        resolve_start_time = datetime.utcnow()
                        drafted_section = await self._resolve_consistency_issues(
                            drafted_section, consistency_issues
                        )
                        resolve_end_time = datetime.utcnow()
                        logger.info(
                            f"[MOTION_DRAFTER] Consistency issues resolved in {resolve_end_time - resolve_start_time}"
                        )
                    else:
                        logger.info("[MOTION_DRAFTER] No consistency issues found")
                except Exception as e:
                    logger.error(
                        f"[MOTION_DRAFTER] Error in consistency check: {str(e)}"
                    )

                drafted_sections.append(drafted_section)
                total_words += drafted_section.word_count

                # Update document context
                logger.info("[MOTION_DRAFTER] Updating document context")
                self._update_document_context(drafted_section)

                section_end_time = datetime.utcnow()
                section_duration = section_end_time - section_start_time
                logger.info(
                    f"[MOTION_DRAFTER] === Section {i + 1} completed in {section_duration} ==="
                )
                logger.info(
                    f"[MOTION_DRAFTER] Section stats: {drafted_section.word_count} words, {len(drafted_section.citations_used)} citations"
                )
                logger.info(
                    f"[MOTION_DRAFTER] Running totals: {total_words} total words across {len(drafted_sections)} sections"
                )

            sections_end_time = datetime.utcnow()
            sections_duration = sections_end_time - sections_start_time
            logger.info(f"[MOTION_DRAFTER] All sections drafted in {sections_duration}")

            # Create motion draft
            logger.info("[MOTION_DRAFTER] Creating motion draft object")
            # Use database name as case name for display, or extract from outline
            display_case_name = outline.get(
                "case_name", database_name.replace("_", " ").title()
            )
            generated_title = motion_title or self._generate_motion_title(outline)

            logger.info(f"[MOTION_DRAFTER] Motion title: {generated_title}")
            logger.info(f"[MOTION_DRAFTER] Case name: {display_case_name}")

            motion_draft = MotionDraft(
                title=generated_title,
                case_name=display_case_name,
                sections=drafted_sections,
                total_word_count=total_words,
                total_page_estimate=total_words // self.words_per_page,
                creation_timestamp=datetime.utcnow(),
                outline_source=outline,
            )

            logger.info(
                f"[MOTION_DRAFTER] Motion draft created: {motion_draft.total_page_estimate} pages, {motion_draft.total_word_count} words"
            )

            # Comprehensive review process
            logger.info("[MOTION_DRAFTER] Starting comprehensive review process")
            review_start_time = datetime.utcnow()
            timeout_monitor.log_progress("Starting comprehensive review")
            progress_tracker.next_step(
                "Comprehensive review", "Checking quality, citations, consistency"
            )

            try:
                motion_draft = await self._comprehensive_review_process(
                    motion_draft, case_context
                )
            except Exception as e:
                logger.error(
                    f"[MOTION_DRAFTER] Error in comprehensive review: {str(e)}"
                )
                motion_draft.review_notes.append(
                    f"Review process encountered errors: {str(e)}"
                )

            review_end_time = datetime.utcnow()
            logger.info(
                f"[MOTION_DRAFTER] Comprehensive review completed in {review_end_time - review_start_time}"
            )
            timeout_monitor.log_progress("Comprehensive review completed")

            # Final consistency check
            logger.info("[MOTION_DRAFTER] Starting final consistency pass")
            consistency_start_time = datetime.utcnow()
            timeout_monitor.log_progress("Final consistency check")
            progress_tracker.next_step(
                "Final consistency check", "Document-wide consistency verification"
            )

            try:
                motion_draft = await self._final_consistency_pass(motion_draft)
            except Exception as e:
                logger.error(
                    f"[MOTION_DRAFTER] Error in final consistency pass: {str(e)}"
                )

            consistency_end_time = datetime.utcnow()
            logger.info(
                f"[MOTION_DRAFTER] Final consistency pass completed in {consistency_end_time - consistency_start_time}"
            )
            timeout_monitor.log_progress("Final consistency check completed")

            # Build citation index
            logger.info("[MOTION_DRAFTER] Building citation index")
            timeout_monitor.log_progress("Building citation index")
            progress_tracker.next_step(
                "Finalization", "Building citation index and final cleanup"
            )

            motion_draft.citation_index = self._build_citation_index(motion_draft)
            logger.info(
                f"[MOTION_DRAFTER] Citation index built: {len(motion_draft.citation_index)} unique citations"
            )

            # Finalize monitoring
            end_time = datetime.utcnow()
            total_duration = end_time - start_time
            timeout_monitor.log_progress(
                "Motion drafting completed",
                {
                    "total_pages": motion_draft.total_page_estimate,
                    "total_words": motion_draft.total_word_count,
                    "quality_score": motion_draft.coherence_score,
                    "sections": len(motion_draft.sections),
                    "citations": len(motion_draft.citation_index),
                },
            )
            timeout_monitor.finish(success=True)
            progress_tracker.finish()

            logger.info("[MOTION_DRAFTER] === MOTION DRAFTING COMPLETED ====")
            logger.info(f"[MOTION_DRAFTER] Total duration: {total_duration}")
            logger.info(
                f"[MOTION_DRAFTER] Final stats: {motion_draft.total_page_estimate} pages, {motion_draft.total_word_count} words"
            )
            logger.info(
                f"[MOTION_DRAFTER] Quality score: {motion_draft.coherence_score:.2f}"
            )
            logger.info(f"[MOTION_DRAFTER] Sections: {len(motion_draft.sections)}")
            logger.info(
                f"[MOTION_DRAFTER] Citations: {len(motion_draft.citation_index)}"
            )

            return motion_draft

        except Exception as e:
            error_time = datetime.utcnow()
            error_duration = (
                error_time - start_time if "start_time" in locals() else "unknown"
            )
            logger.error(
                f"[MOTION_DRAFTER] ERROR in enhanced motion drafting after {error_duration}: {str(e)}",
                exc_info=True,
            )
            logger.error(
                f"[MOTION_DRAFTER] Database: {database_name}, Target length: {target_length}"
            )
            logger.error(
                f"[MOTION_DRAFTER] Outline keys: {list(outline.keys()) if outline else 'None'}"
            )
            if "drafted_sections" in locals():
                logger.error(
                    f"[MOTION_DRAFTER] Sections completed before error: {len(drafted_sections)}"
                )

            # Log timeout monitoring summary on error
            if "timeout_monitor" in locals():
                summary = timeout_monitor.get_summary()
                logger.error(f"[MOTION_DRAFTER] Timeout monitor summary: {summary}")
                timeout_monitor.finish(success=False)

            if "progress_tracker" in locals():
                logger.error(
                    f"[MOTION_DRAFTER] Progress when error occurred: {progress_tracker.current_step}/{progress_tracker.total_steps}"
                )

            raise

    def _parse_enhanced_outline(
        self, outline: Dict[str, Any], target_length: DocumentLength
    ) -> List[OutlineSection]:
        """Parse outline with enhanced structure from doc-converter format"""
        sections = []

        # Calculate target distribution
        min_pages, max_pages = target_length.value
        target_words = ((min_pages + max_pages) // 2) * self.words_per_page

        # Handle the doc-converter format
        # Check if this is the new format with array wrapper
        outline_data = outline
        if isinstance(outline, list) and len(outline) > 0:
            outline_data = outline[0]

        # Extract style notes for the document
        if "style_notes" in outline_data:
            self.document_context["style_notes"] = outline_data["style_notes"]

        # Get outline sections
        outline_sections = outline_data.get("sections", [])
        if not outline_sections and "arguments" in outline_data:
            outline_sections = outline_data["arguments"]

        # Calculate intelligent word distribution
        word_distribution = self._calculate_word_distribution(
            outline_sections, target_words
        )

        for idx, section_data in enumerate(outline_sections):
            # Parse the new format from doc-converter
            parsed_section = self._parse_doc_converter_section(
                section_data, idx, word_distribution
            )
            sections.append(parsed_section)

        return sections

    def _parse_doc_converter_section(
        self, section_data: Dict[str, Any], idx: int, word_distribution: Dict[str, int]
    ) -> OutlineSection:
        """Parse a section in the doc-converter format"""
        # Clean the section data first
        section_data = self._clean_outline_section(section_data)

        section_type = self._determine_section_type_from_data(section_data)

        # Extract content from the structured format
        content_points = []
        hook_options = []
        themes = []
        key_facts = []
        counter_arguments = []
        legal_authorities = []

        # Process the content array
        content_items = section_data.get("content", [])
        for item in content_items[:10]:  # Limit items processed
            item_type = item.get("type", "")

            if item_type == "field":
                label = item.get("label", "")
                value = item.get("value", "")

                if label == "Hook":
                    # Parse hook options more aggressively
                    if "||" in value:
                        first_hook = value.split("||")[0].strip()
                        if first_hook.startswith("Option"):
                            first_hook = (
                                first_hook.split(":", 1)[1].strip()
                                if ":" in first_hook
                                else first_hook
                            )
                        # Clean quotes and reduce length
                        first_hook = first_hook.strip('"').strip()
                        hook_options = [first_hook[:150]]  # Reduce to 150 chars
                    else:
                        hook_options = [value.strip('"').strip()[:150]]
                elif label == "Theme":
                    # Parse themes more carefully
                    if "||" in value:
                        theme_parts = value.split("||")
                        themes = [
                            t.strip()[:100] for t in theme_parts[:3]
                        ]  # Max 3 themes
                    else:
                        themes = [value[:100]]
                else:
                    # For other fields, just add a truncated version
                    content_points.append(f"{label}: {value[:200]}")

        # Create the outline section
        section = OutlineSection(
            id=f"section_{idx}",
            title=section_data.get(
                "heading", section_data.get("title", f"Section {idx + 1}")
            ),
            section_type=section_type,
            content_points=content_points
            if content_points
            else [section_data.get("summary", "")],
            legal_authorities=legal_authorities,
            target_length=word_distribution.get(f"section_{idx}", 500),
            hook_options=hook_options,
            themes=themes if themes else self.document_context.get("themes", []),
            key_facts=key_facts,
            counter_arguments=counter_arguments,
        )

        # Parse any subsections in the same format
        if "subsections" in section_data:
            subsection_words = section.target_length // (
                len(section_data["subsections"]) + 1
            )
            for sub_idx, subsection_data in enumerate(section_data["subsections"]):
                subsection = self._parse_doc_converter_section(
                    subsection_data,
                    f"{idx}_{sub_idx}",
                    {f"section_{idx}_{sub_idx}": subsection_words},
                )
                subsection.parent_id = section.id
                section.children.append(subsection)

        return section

    def _determine_section_type_from_data(
        self, section_data: Dict[str, Any]
    ) -> SectionType:
        """Determine section type from doc-converter data"""
        # First check the explicit type field
        explicit_type = section_data.get("type", "").lower()

        type_mapping = {
            "introduction": SectionType.INTRODUCTION,
            "facts": SectionType.STATEMENT_OF_FACTS,
            "legal_standard": SectionType.LEGAL_STANDARD,
            "argument": SectionType.ARGUMENT,
            "conclusion": SectionType.CONCLUSION,
            "prayer": SectionType.PRAYER_FOR_RELIEF,
        }

        if explicit_type in type_mapping:
            return type_mapping[explicit_type]

        # Fall back to title analysis
        title = section_data.get("heading", section_data.get("title", "")).lower()
        return self._determine_section_type(title)

    def _calculate_word_distribution(
        self, outline_sections: List[Dict], target_words: int
    ) -> Dict[str, int]:
        """Intelligently distribute words based on section importance"""
        distribution = {}

        # Assign weights based on section type and content
        weights = {}
        total_weight = 0

        for idx, section in enumerate(outline_sections):
            title_lower = section.get("title", "").lower()

            # Assign weights based on section type
            if "introduction" in title_lower:
                weight = 0.5
            elif "facts" in title_lower or "background" in title_lower:
                weight = 1.5
            elif "standard" in title_lower:
                weight = 0.8
            elif "conclusion" in title_lower:
                weight = 0.5
            elif "prayer" in title_lower:
                weight = 0.3
            else:  # Main arguments
                weight = 1.2
                # Add weight for subsections
                weight += 0.3 * len(section.get("subsections", []))

            weights[f"section_{idx}"] = weight
            total_weight += weight

        # Distribute words based on weights
        for section_id, weight in weights.items():
            distribution[section_id] = int((weight / total_weight) * target_words)

        return distribution

    def _extract_authorities(self, section_data: Dict) -> List[str]:
        """Extract legal authorities from various fields in the outline"""
        authorities = []

        # Direct authorities field
        authorities.extend(section_data.get("authorities", []))
        authorities.extend(section_data.get("legal_authorities", []))

        # Extract from key_authorities
        if "key_authorities" in section_data:
            for auth in section_data["key_authorities"]:
                if isinstance(auth, dict):
                    authorities.append(auth.get("citation", ""))
                else:
                    authorities.append(str(auth))

        # Extract from content if present
        content = section_data.get("content", "")
        if content:
            # Use regex to find citations
            citations = self._extract_citations_from_text(content)
            authorities.extend(citations)

        return list(
            set(filter(None, authorities))
        )  # Remove duplicates and empty strings

    async def _retrieve_enhanced_case_context(
        self, database_name: str, outline_sections: List[OutlineSection]
    ) -> Dict[str, Any]:
        """Retrieve comprehensive case context directly from database"""

        # Check cache first
        cache_key = f"case_context_{database_name}"
        cached_context = await motion_cache.get(
            "case_context", database_name=database_name
        )
        if cached_context:
            logger.info("Using cached case context")
            return cached_context

        context = {
            "case_facts": [],
            "legal_authorities": [],
            "expert_reports": [],
            "key_documents": [],
            "medical_records": [],
            "motion_practice": [],
            "fact_chronology": [],
            "themes": [],
            "opposing_motion": self.document_context.get("opposing_motion_text", ""),
        }

        # Build comprehensive search queries
        search_queries = self._build_comprehensive_search_queries(outline_sections)

        # LIMIT the number of queries to prevent overload
        max_queries_per_type = 5

        for query_type, queries in search_queries.items():
            # Limit queries
            limited_queries = queries[:max_queries_per_type]
            logger.info(f"Processing {len(limited_queries)} queries for {query_type}")

            for query in limited_queries:
                try:
                    # Clean and validate query
                    query = query.strip()
                    if not query or len(query) < 3:
                        continue

                    # Truncate overly long queries
                    if len(query) > 200:
                        query = query[:200]

                    logger.debug(f"Searching for: {query[:100]}...")

                    # Generate embedding for the query
                    query_embedding, _ = self.embedding_generator.generate_embedding(
                        query
                    )

                    # Perform hybrid search with timeout
                    results = await asyncio.wait_for(
                        self.vector_store.hybrid_search(
                            collection_name=database_name,
                            query=query,
                            query_embedding=query_embedding,
                            limit=5,  # Reduced from 10
                            enable_reranking=False,
                        ),
                        timeout=10,  # 10 second timeout per search
                    )

                    # Categorize results
                    for result in results:
                        if result.score > 0.6:
                            doc_type = result.metadata.get("document_type", "").lower()

                            if "medical" in doc_type:
                                context["medical_records"].append(result)
                            elif "expert" in doc_type:
                                context["expert_reports"].append(result)
                            elif "motion" in doc_type:
                                context["motion_practice"].append(result)
                            elif query_type == "legal_authorities":
                                context["legal_authorities"].append(
                                    {
                                        "citation": query,
                                        "content": result.content,
                                        "document": result.metadata.get(
                                            "document_name", "Unknown"
                                        ),
                                        "score": result.score,
                                        "context": result.metadata.get(
                                            "context_summary", ""
                                        ),
                                    }
                                )
                            else:
                                context["case_facts"].append(result)

                except Exception as e:
                    logger.error(f"Error searching for '{query}': {str(e)}")
                    continue

        # Build fact chronology
        context["fact_chronology"] = self._build_fact_chronology(context["case_facts"])

        # Extract themes
        context["themes"] = self._extract_case_themes(context)

        # Cache the context
        await motion_cache.set(
            "case_context", context, ttl=7200, **{"database_name": database_name}
        )

        logger.info(
            f"Retrieved comprehensive context from {database_name}: {len(context['legal_authorities'])} authorities, "
            f"{len(context['case_facts'])} facts, {len(context['expert_reports'])} expert reports, "
            f"{len(context['medical_records'])} medical records"
        )

        return context

    def _build_comprehensive_search_queries(
        self, outline_sections: List[OutlineSection]
    ) -> Dict[str, List[str]]:
        """Build comprehensive search queries based on outline"""
        queries = {
            "legal_authorities": [],
            "factual_searches": [],
            "expert_searches": [],
            "procedural_history": [],
        }

        # Extract all legal authorities
        for section in outline_sections:
            queries["legal_authorities"].extend(section.legal_authorities)
            for child in section.children:
                queries["legal_authorities"].extend(child.legal_authorities)

        # Build factual search queries
        fact_patterns = set()
        for section in outline_sections:
            # Extract fact patterns from key facts
            for fact in section.key_facts:
                if isinstance(fact, dict):
                    fact_patterns.add(fact.get("description", ""))
                else:
                    fact_patterns.add(str(fact))

            # Extract from content points - ensure they're strings
            for point in section.content_points[:3]:  # Top 3 points
                if isinstance(point, str):
                    fact_patterns.add(point)
                else:
                    fact_patterns.add(str(point))

        queries["factual_searches"] = list(fact_patterns)

        # Expert-specific searches
        queries["expert_searches"] = [
            "expert report",
            "expert opinion",
            "expert analysis",
            "expert testimony",
            "expert conclusions",
        ]

        # Procedural history searches
        queries["procedural_history"] = [
            "complaint filed",
            "motion to dismiss",
            "discovery order",
            "scheduling order",
            "previous motions",
        ]

        return queries

    def _clean_outline_content(self, content: str, max_length: int = 500) -> str:
        """Clean and truncate outline content to prevent token overflow"""
        if not content:
            return ""

        # Handle multi-option format (e.g., "Option 1: ... || Option 2: ...")
        if "||" in content:
            options = content.split("||")
            # Take only the first option and clean it
            content = options[0].strip()
            # Remove "Option 1:" prefix if present
            if content.startswith("Option"):
                content = (
                    content.split(":", 1)[1].strip() if ":" in content else content
                )

        # Truncate if too long
        if len(content) > max_length:
            content = content[:max_length] + "..."

        return content

    def _clean_outline_section(self, section_data: Dict[str, Any]) -> Dict[str, Any]:
        """Clean a section to prevent token overflow"""
        cleaned = section_data.copy()

        # Clean content array
        if "content" in cleaned and isinstance(cleaned["content"], list):
            for item in cleaned["content"]:
                if item.get("type") == "field" and "value" in item:
                    item["value"] = self._clean_outline_content(item["value"])
                elif item.get("type") == "paragraph" and "text" in item:
                    item["text"] = self._clean_outline_content(
                        item["text"], max_length=1000
                    )

        return cleaned

    def _determine_section_type(self, title: str) -> SectionType:
        """Determine section type from title"""
        title_lower = title.lower()

        if any(
            term in title_lower for term in ["introduction", "preliminary", "overview"]
        ):
            return SectionType.INTRODUCTION
        elif any(term in title_lower for term in ["facts", "background", "factual"]):
            return SectionType.STATEMENT_OF_FACTS
        elif any(
            term in title_lower for term in ["standard", "law", "legal framework"]
        ):
            return SectionType.LEGAL_STANDARD
        elif any(term in title_lower for term in ["conclusion", "summary"]):
            return SectionType.CONCLUSION
        elif any(term in title_lower for term in ["prayer", "relief", "request"]):
            return SectionType.PRAYER_FOR_RELIEF
        else:
            return SectionType.ARGUMENT

    async def _draft_section_with_cot(
        self,
        section: OutlineSection,
        case_context: Dict[str, Any],
        cumulative_context: Dict[str, Any],
    ) -> DraftedSection:
        """Draft section using Chain-of-Thought approach"""

        try:
            # LIMIT and CLEAN the content points to prevent token overflow
            content_points_limited = []
            for point in section.content_points[:5]:  # Max 5 points
                cleaned_point = self._clean_outline_content(
                    point, 200
                )  # Max 200 chars per point
                if cleaned_point:
                    content_points_limited.append(cleaned_point)

            legal_authorities_limited = section.legal_authorities[
                :7
            ]  # Max 7 authorities

            # Also limit hook options more aggressively
            if section.hook_options:
                section.hook_options = [
                    self._clean_outline_content(hook, 150)
                    for hook in section.hook_options[:2]
                ]  # Max 2 hooks, 150 chars each

            # First, use CoT to plan the section
            planning_prompt = f"""Plan the drafting of this legal section using step-by-step reasoning:

Section: {section.title}
Type: {section.section_type.value}
Target Length: {section.target_length} words

Content Points to Address (top {len(content_points_limited)}):
{chr(10).join(f"- {self._clean_outline_content(point, 200)}" for point in content_points_limited)}

Required Authorities (top {len(legal_authorities_limited)}):
{chr(10).join(f"- {auth}" for auth in legal_authorities_limited)}

Provide a detailed plan for drafting this section."""

            # Get section plan
            logger.info(f"[COT] Planning section: {section.title}")
            section_plan_result = await asyncio.wait_for(
                self.section_writer.run(planning_prompt),
                timeout=30,  # 30 second timeout
            )
            section_plan = (
                str(section_plan_result.data)
                if hasattr(section_plan_result, "data")
                else str(section_plan_result)
            )
            logger.info(f"[COT] Section plan created, length: {len(section_plan)}")

            # Now draft based on the plan
            relevant_facts = self._filter_relevant_facts(
                case_context["case_facts"], section
            )
            relevant_authorities = self._filter_relevant_authorities(
                case_context["legal_authorities"], section.legal_authorities
            )

            # Use hook options if this is the introduction
            hook_content = ""
            if (
                section.section_type == SectionType.INTRODUCTION
                and section.hook_options
            ):
                hook_content = f"\nHook Options:\n{chr(10).join(f'{i + 1}. {hook}' for i, hook in enumerate(section.hook_options))}\n\nChoose the most compelling hook and develop it fully."

            opposing_motion_context = ""
            if self.document_context.get("opposing_motion_text"):
                opposing_motion_context = f"""

Opposing Motion to Respond To:
{self.document_context["opposing_motion_text"]}

Key Arguments from Opposing Counsel to Address:
- Analyze and respond to their primary arguments
- Identify weaknesses in their legal reasoning
- Counter their factual assertions with our evidence
"""

            drafting_prompt = f"""Draft this section of the legal motion based on the following plan:

{section_plan}

Section Details:
- Title: {section.title}
- Target Length: {section.target_length} words (THIS IS A MINIMUM - aim for 10-20% more)
- Document Themes: {", ".join(self.document_context["themes"])}
{hook_content}
{opposing_motion_context}

Relevant Case Facts:
{self._format_enhanced_facts(relevant_facts, section.key_facts)}

Legal Authority Analysis:
{self._format_enhanced_authorities(relevant_authorities)}

Previous Sections Summary:
{cumulative_context.get("summary", "This is the first section.")}

Key Citations Already Used:
{", ".join(list(cumulative_context.get("citations_used", set()))[:10])}

CRITICAL INSTRUCTIONS:
1. This is a COMPREHENSIVE legal document, not a summary
2. Develop EACH point with multiple paragraphs of analysis
3. Include detailed factual application for every legal element
4. Add parenthetical explanations for all case citations
5. Use "Furthermore," "Moreover," "Additionally" to expand points
6. Include "even if" arguments to show thoroughness
7. DO NOT use conclusory statements without support
8. DO NOT write in bullet points - use flowing paragraphs
9. Target {section.target_length} words MINIMUM"""

            # Generate comprehensive draft
            logger.info("[COT] Drafting section content")
            initial_content_result = await asyncio.wait_for(
                self.section_writer.run(drafting_prompt),
                timeout=60,  # 60 second timeout for content generation
            )
            initial_content = (
                str(initial_content_result.data)
                if hasattr(initial_content_result, "data")
                else str(initial_content_result)
            )
            logger.info(
                f"[COT] Initial draft generated, length: {len(initial_content.split())} words"
            )

            # Verify and enhance citations
            enhanced_content = await self._enhance_citations(
                initial_content, relevant_authorities
            )

            # Calculate metrics
            word_count = len(enhanced_content.split())
            citations_used = self._extract_citations_from_text(enhanced_content)
            confidence_score = self._calculate_section_confidence(
                enhanced_content, section, word_count
            )

            # Create drafted section
            drafted_section = DraftedSection(
                outline_section=section,
                content=enhanced_content,
                word_count=word_count,
                citations_used=citations_used,
                citations_verified={},  # Will be populated in verification step
                expansion_cycles=1,
                confidence_score=confidence_score,
                consistency_score=1.0,  # Will be updated in consistency check
                factual_accuracy_score=self._calculate_factual_accuracy(
                    enhanced_content, relevant_facts
                ),
                argument_strength_score=self._calculate_argument_strength(
                    enhanced_content, section
                ),
            )

            # Draft subsections if present
            if section.children:
                for child in section.children:
                    child_context = {
                        **cumulative_context,
                        "parent_section": enhanced_content[-1000:],  # Last 1000 chars
                    }
                    child_draft = await self._draft_section_with_cot(
                        child, case_context, child_context
                    )

                    # Integrate subsection
                    drafted_section.content += f"\n\n{child_draft.content}"
                    drafted_section.word_count += child_draft.word_count
                    drafted_section.citations_used.extend(child_draft.citations_used)

            return drafted_section

        except Exception as e:
            logger.error(
                f"[COT] Error in _draft_section_with_cot: {str(e)}", exc_info=True
            )
            # Return a minimal section on error
            return DraftedSection(
                outline_section=section,
                content=f"[ERROR: Unable to draft section - {str(e)}]",
                word_count=0,
                citations_used=[],
                citations_verified={},
                expansion_cycles=0,
                confidence_score=0.0,
            )

    async def _expand_section_intelligently(
        self,
        drafted_section: DraftedSection,
        target_length: int,
        case_context: Dict[str, Any],
        cumulative_context: Dict[str, Any],
    ) -> DraftedSection:
        """Intelligent expansion that adds substantive content"""

        try:
            current_length = drafted_section.word_count
            expansion_needed = target_length - current_length

            # Prevent division by zero when current_length is 0
            if current_length > 0:
                expansion_percentage = (expansion_needed / current_length) * 100
            else:
                expansion_percentage = 100.0  # Fallback when starting from 0 words

            logger.info(
                f"Expanding {drafted_section.outline_section.title} by {expansion_percentage:.1f}% "
                f"({expansion_needed} words)"
            )

            # Identify expansion opportunities
            expansion_analysis = await self._analyze_expansion_opportunities(
                drafted_section.content, drafted_section.outline_section
            )

            expansion_prompt = f"""Expand this legal section by adding {expansion_needed} words of substantive content.

Current Section ({current_length} words):
{drafted_section.content}

Expansion Analysis:
{expansion_analysis}

Available Additional Context:
- Unused Legal Authorities: {self._get_unused_authorities(drafted_section, case_context)}
- Additional Facts: {self._get_additional_facts(drafted_section, case_context)}
- Policy Considerations: {self._get_policy_considerations(drafted_section.outline_section)}

EXPANSION REQUIREMENTS:
1. Add {expansion_needed} words MINIMUM
2. DO NOT summarize or remove existing content
3. Add new paragraphs of analysis, not just phrases
4. Develop these specific areas:
   - Deeper analysis of case law holdings
   - More detailed factual application
   - Additional counterargument rebuttals
   - Policy and practical implications
   - Alternative legal theories
   - Procedural considerations

Insert new content at natural break points.
Mark insertions with [EXPANSION START] and [EXPANSION END] tags."""

            # Generate expansion
            expanded_content_result = await self.section_expander.run(expansion_prompt)
            expanded_content = (
                str(expanded_content_result.data)
                if hasattr(expanded_content_result, "data")
                else str(expanded_content_result)
            )

            # Integrate expansion seamlessly
            integrated_content = self._integrate_expansion_seamlessly(
                drafted_section.content, expanded_content
            )

            # Update section
            new_word_count = len(integrated_content.split())
            drafted_section.content = integrated_content
            drafted_section.word_count = new_word_count
            drafted_section.expansion_cycles += 1

            # Extract new citations
            new_citations = self._extract_citations_from_text(expanded_content)
            drafted_section.citations_used = list(
                set(drafted_section.citations_used + new_citations)
            )

            # Recalculate quality scores
            drafted_section.confidence_score = self._calculate_section_confidence(
                integrated_content, drafted_section.outline_section, new_word_count
            )

            return drafted_section

        except Exception as e:
            logger.error(f"Error in _expand_section_intelligently: {str(e)}")
            # Return the section unchanged on error
            return drafted_section

    async def _create_sophisticated_transition(
        self, from_section: DraftedSection, to_section: DraftedSection
    ) -> str:
        """Create sophisticated multi-paragraph transitions"""

        try:
            # Extract key holdings and conclusions
            from_conclusions = self._extract_conclusions(from_section.content)
            to_premises = self._extract_opening_premises(to_section.content)

            transition_prompt = f"""Create a sophisticated 2-3 paragraph transition between these sections:

FROM SECTION: {from_section.outline_section.title}
Key Conclusions:
{chr(10).join(f"- {c}" for c in from_conclusions)}

TO SECTION: {to_section.outline_section.title}  
Opening Premises:
{chr(10).join(f"- {p}" for p in to_premises)}

Transition Type Needed: {self._determine_transition_type(from_section, to_section)}

Document Themes: {", ".join(self.document_context["themes"])}

Create a transition that:
1. First paragraph: Crystallize the key holding/victory from the previous section
2. Second paragraph: Bridge logically to the next argument
3. Optional third paragraph: Preview why the next argument strengthens our position

Use sophisticated legal transition language:
- "Having established that..., we now demonstrate..."
- "This principle finds further support when examining..."
- "The foregoing analysis necessarily implicates..."
- "Building upon this foundation..."

Make it substantive, not perfunctory."""

            transition_result = await self.transition_writer.run(transition_prompt)
            transition = (
                str(transition_result.data)
                if hasattr(transition_result, "data")
                else str(transition_result)
            )
            return transition

        except Exception as e:
            logger.error(f"Error creating transition: {str(e)}")
            return ""

    async def _check_section_consistency(
        self,
        current_section: DraftedSection,
        previous_sections: List[DraftedSection],
        document_context: Dict[str, Any],
    ) -> List[Dict[str, Any]]:
        """Check section consistency with document"""

        try:
            consistency_prompt = f"""Check this section for consistency with the rest of the document:

Current Section:
{current_section.content[:2000]}...

Document Context:
- Established Terminology: {json.dumps(document_context.get("terminology", {}))}
- Key Facts Used: {json.dumps([f["description"] for f in document_context.get("fact_chronology", [])[:10]])}
- Citations Already Used: {list(document_context.get("citations_used", set()))[:20]}

Previous Sections Summary:
{self._summarize_previous_sections(previous_sections)}

Check for:
1. Terminology consistency (same terms for same concepts)
2. Factual consistency (dates, names, amounts)
3. Legal position consistency (no contradictions)
4. Citation consistency (same cases cited the same way)
5. Tone and style consistency

Output JSON with any inconsistencies found."""

            try:
                result_raw = await self.consistency_checker.run(consistency_prompt)
                result = result_raw.data if hasattr(result_raw, "data") else result_raw
            except Exception as e:
                logger.error(f"Consistency checker failed: {str(e)}")
                # Return empty issues on failure instead of crashing
                return []

            consistency_data = (
                result if isinstance(result, dict) else {"inconsistencies": []}
            )
            return consistency_data.get("inconsistencies", [])

        except Exception as e:
            logger.error(f"Error in consistency check: {str(e)}")
            return []

    async def _comprehensive_review_process(
        self, motion_draft: MotionDraft, case_context: Dict[str, Any]
    ) -> MotionDraft:
        """Enhanced comprehensive review process"""

        try:
            # Compile full document
            full_text = self._compile_full_document(motion_draft)

            # Multi-stage review
            # Stage 1: Legal accuracy and citations
            citation_review = await self._review_citations(motion_draft)

            # Stage 2: Consistency check
            consistency_review = await self._review_consistency(motion_draft)

            # Stage 3: Argument strength
            argument_review = await self._review_argument_strength(
                motion_draft, case_context
            )

            # Stage 4: Overall quality
            review_prompt = f"""Conduct a comprehensive review of this legal motion:

{full_text[:10000]}... [truncated for length]

Document Metrics:
- Total Pages: {motion_draft.total_page_estimate}
- Total Words: {motion_draft.total_word_count}
- Sections: {len(motion_draft.sections)}
- Citations: {len(set(sum([s.citations_used for s in motion_draft.sections], [])))}

Previous Review Results:
- Citation Issues: {len(citation_review.get("format_issues", []))}
- Consistency Issues: {len(consistency_review.get("inconsistencies", []))}
- Weak Arguments: {len(argument_review.get("weak_sections", []))}

Evaluate:
1. Overall readiness for filing (1-10)
2. Professional tone and presentation
3. Persuasiveness of arguments
4. Compliance with court rules
5. Any critical issues requiring immediate fix

Provide comprehensive feedback as JSON."""

            overall_review_raw = await self.document_reviewer.run(review_prompt)
            overall_review = (
                overall_review_raw.data
                if hasattr(overall_review_raw, "data")
                else overall_review_raw
            )

            # Parse and apply review feedback
            review_data = (
                overall_review
                if isinstance(overall_review, dict)
                else {"readiness_score": 7, "section_scores": {}, "specific_edits": []}
            )

            # Update motion draft with review results
            motion_draft.coherence_score = review_data.get("readiness_score", 7) / 10.0
            motion_draft.review_notes = review_data.get("specific_edits", [])

            # Quality metrics
            motion_draft.quality_metrics = {
                "readiness_score": review_data.get("readiness_score", 7),
                "citation_accuracy": 1.0
                - (
                    len(citation_review.get("format_issues", []))
                    / max(len(motion_draft.citation_index), 1)
                ),
                "consistency_score": 1.0
                - (len(consistency_review.get("inconsistencies", [])) / 10.0),
                "argument_strength": argument_review.get("average_strength", 0.7),
                "professionalism": review_data.get("professionalism_score", 8) / 10.0,
            }

            # Mark sections needing revision
            for section_id, issues in review_data.get("section_issues", {}).items():
                for section in motion_draft.sections:
                    if section.outline_section.id == section_id:
                        section.needs_revision = True
                        section.revision_notes.extend(issues)

            # Apply critical fixes if needed
            if motion_draft.coherence_score < 0.7:
                motion_draft = await self._apply_critical_fixes(
                    motion_draft, review_data, citation_review, consistency_review
                )

            logger.info("[REVIEW] Comprehensive review process completed")
            return motion_draft

        except Exception as e:
            logger.error(f"[REVIEW] Error in comprehensive review: {str(e)}")
            motion_draft.coherence_score = 0.7
            motion_draft.review_notes = ["Review completed with errors"]
            return motion_draft

    async def _final_consistency_pass(self, motion_draft: MotionDraft) -> MotionDraft:
        """Final consistency pass across entire document"""

        try:
            # Build terminology map
            terminology_map = {}
            for section in motion_draft.sections:
                # Extract key terms
                terms = self._extract_key_terms(section.content)
                for term in terms:
                    if term not in terminology_map:
                        terminology_map[term] = []
                    terminology_map[term].append(section.outline_section.id)

            # Check for inconsistent usage
            inconsistencies = []
            for term, occurrences in terminology_map.items():
                if len(set(occurrences)) > 1:
                    # Term appears in multiple sections - check consistency
                    variants = self._find_term_variants(term, motion_draft)
                    if variants:
                        inconsistencies.append(
                            {
                                "term": term,
                                "variants": variants,
                                "sections": occurrences,
                            }
                        )

            # Apply fixes
            if inconsistencies:
                motion_draft = await self._fix_terminology_inconsistencies(
                    motion_draft, inconsistencies
                )

            # Update consistency issues
            motion_draft.consistency_issues = inconsistencies

            return motion_draft

        except Exception as e:
            logger.error(f"Error in final consistency pass: {str(e)}")
            return motion_draft

    def _build_citation_index(self, motion_draft: MotionDraft) -> Dict[str, List[str]]:
        """Build index of citations to sections using them"""
        citation_index = {}

        for section in motion_draft.sections:
            for citation in section.citations_used:
                if citation not in citation_index:
                    citation_index[citation] = []
                citation_index[citation].append(section.outline_section.id)

        return citation_index

    # Helper methods

    def _compile_citation_patterns(self) -> List[re.Pattern]:
        """Compile regex patterns for citation extraction"""
        patterns = [
            # Federal cases
            re.compile(r"\d+\s+U\.S\.\s+\d+(?:\s+\(\d{4}\))?"),
            re.compile(r"\d+\s+F\.\d+d\s+\d+(?:\s+\([^)]+\d{4}\))?"),
            re.compile(r"\d+\s+F\.\s*Supp\.\s*\d+d?\s+\d+(?:\s+\([^)]+\d{4}\))?"),
            re.compile(r"\d+\s+S\.\s*Ct\.\s+\d+(?:\s+\(\d{4}\))?"),
            # State cases
            re.compile(r"\d+\s+[A-Z][a-z]+\.(?:\s+\d+d)?\s+\d+(?:\s+\([^)]+\d{4}\))?"),
            # Case names
            re.compile(r"[A-Z][a-zA-Z]+\s+v\.\s+[A-Z][a-zA-Z]+(?:,\s+\d+)?"),
            # Statutes
            re.compile(r"\d+\s+U\.S\.C\.\s+§+\s*\d+[\w\-\.]*"),
            re.compile(r"[A-Z][a-z]+\.\s+Stat\.\s+(?:Ann\.\s+)?§+\s*\d+[\w\-\.]*"),
            # Regulations
            re.compile(r"\d+\s+C\.F\.R\.\s+§*\s*\d+\.\d+"),
            # Rules
            re.compile(r"Fed\.\s*R\.\s*(?:Civ|Crim|Evid|App)\.\s*P\.\s*\d+[a-z]*"),
        ]
        return patterns

    def _extract_citations_from_text(self, text: str) -> List[str]:
        """Extract all citations from text using compiled patterns"""
        citations = []

        for pattern in self.citation_patterns:
            matches = pattern.findall(text)
            citations.extend(matches)

        # Clean and deduplicate
        citations = list(set(citation.strip() for citation in citations))

        return citations

    def _calculate_section_confidence(
        self, content: str, section: OutlineSection, word_count: int
    ) -> float:
        """Calculate confidence score for section"""
        score = 0.0
        max_score = 100.0

        # Length score (30 points)
        if word_count >= section.target_length:
            score += 30
        else:
            score += (word_count / section.target_length) * 30

        # Content points addressed (25 points)
        points_addressed = sum(
            1
            for point in section.content_points
            if any(keyword in content.lower() for keyword in point.lower().split()[:5])
        )
        if section.content_points:
            score += (points_addressed / len(section.content_points)) * 25
        else:
            score += 25

        # Citations included (25 points)
        citations_found = sum(
            1
            for auth in section.legal_authorities
            if auth in content or auth.replace(" ", "") in content
        )
        if section.legal_authorities:
            score += (citations_found / len(section.legal_authorities)) * 25
        else:
            score += 25

        # Structure quality (20 points)
        structure_indicators = [
            "first",
            "second",
            "third",
            "moreover",
            "furthermore",
            "additionally",
            "however",
            "nevertheless",
            "therefore",
            "accordingly",
            "issue",
            "rule",
            "application",
            "conclusion",
        ]
        structure_score = sum(
            2 for indicator in structure_indicators if indicator in content.lower()
        )
        score += min(structure_score, 20)

        return score / max_score

    def _calculate_factual_accuracy(
        self, content: str, relevant_facts: List[Any]
    ) -> float:
        """Calculate factual accuracy score"""
        if not relevant_facts:
            return 1.0

        facts_referenced = 0
        content_lower = content.lower()

        for fact in relevant_facts:
            fact_text = fact.content if hasattr(fact, "content") else str(fact)
            # Check if key parts of fact are referenced
            key_words = fact_text.lower().split()[:10]  # First 10 words
            if sum(1 for word in key_words if word in content_lower) >= 3:
                facts_referenced += 1

        return min(facts_referenced / len(relevant_facts), 1.0)

    def _calculate_argument_strength(
        self, content: str, section: OutlineSection
    ) -> float:
        """Calculate argument strength score"""
        score = 0.0

        # Check for IRAC/CRAC structure
        if all(
            indicator in content.lower()
            for indicator in ["issue", "rule", "application"]
        ):
            score += 0.3

        # Check for counter-argument handling
        counter_indicators = [
            "even if",
            "assuming",
            "although",
            "despite",
            "nevertheless",
        ]
        if any(indicator in content.lower() for indicator in counter_indicators):
            score += 0.2

        # Check for policy arguments
        policy_indicators = [
            "policy",
            "purpose",
            "intent",
            "rationale",
            "public interest",
        ]
        if any(indicator in content.lower() for indicator in policy_indicators):
            score += 0.2

        # Check for strong conclusion language
        conclusion_indicators = [
            "therefore",
            "accordingly",
            "thus",
            "clearly",
            "established",
        ]
        if any(indicator in content.lower() for indicator in conclusion_indicators):
            score += 0.2

        # Check citation density (citations per 100 words)
        citations = self._extract_citations_from_text(content)
        citation_density = len(citations) / (len(content.split()) / 100)
        if citation_density >= 1.0:  # At least 1 citation per 100 words
            score += 0.1

        return min(score, 1.0)

    def _format_enhanced_facts(
        self, relevant_facts: List[Any], key_facts: List[Dict]
    ) -> str:
        """Format facts with enhanced detail"""
        formatted = []

        # Add key facts from outline
        if key_facts:
            formatted.append("Key Facts from Outline:")
            for i, fact in enumerate(key_facts[:5], 1):
                if isinstance(fact, dict):
                    formatted.append(f"{i}. {fact.get('description', str(fact))}")
                else:
                    formatted.append(f"{i}. {str(fact)}")

        # Add retrieved facts
        if relevant_facts:
            formatted.append("\nRelevant Case Documents:")
            for i, fact in enumerate(relevant_facts[:5], 1):
                if hasattr(fact, "content"):
                    content = fact.content[:300]
                    doc_name = (
                        fact.metadata.get("document_name", "Unknown")
                        if hasattr(fact, "metadata")
                        else "Unknown"
                    )
                    formatted.append(f"{i}. [{doc_name}] {content}...")
                else:
                    formatted.append(f"{i}. {str(fact)[:300]}...")

        return "\n".join(formatted) if formatted else "No specific facts available."

    def _format_enhanced_authorities(self, authorities: List[Dict[str, Any]]) -> str:
        """Format authorities with enhanced detail"""
        if not authorities:
            return "No specific authorities retrieved."

        formatted = []
        for auth in authorities[:7]:  # Top 7 authorities
            citation = auth.get("citation", "")
            content = auth.get("content", "")[:250]
            context = auth.get("context", "")
            score = auth.get("score", 0)

            formatted.append(f"""- {citation} (Relevance: {score:.2f})
  Holding: {content}...
  Context: {context}""")

        return "\n\n".join(formatted)

    def _build_cumulative_context(
        self, drafted_sections: List[DraftedSection], case_context: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Build cumulative context from drafted sections"""
        context = {
            "summary": "",
            "key_holdings": [],
            "citations_used": set(),
            "themes_developed": [],
            "facts_established": [],
        }

        if not drafted_sections:
            context["summary"] = "This is the first section of the motion."
            return context

        # Build summary of previous sections
        summaries = []
        for section in drafted_sections[-3:]:  # Last 3 sections
            # Extract key holding
            conclusion = self._extract_section_conclusion(section.content)
            summaries.append(f"{section.outline_section.title}: {conclusion}")

            # Collect citations
            context["citations_used"].update(section.citations_used)

            # Extract established facts
            facts = self._extract_established_facts(section.content)
            context["facts_established"].extend(facts)

        context["summary"] = " ".join(summaries)
        context["themes_developed"] = list(self.document_context.get("themes", []))

        return context

    def _extract_section_conclusion(self, content: str) -> str:
        """Extract the conclusion from a section"""
        # Look for conclusion indicators
        conclusion_markers = [
            "therefore",
            "accordingly",
            "thus",
            "in conclusion",
            "for these reasons",
            "as demonstrated",
        ]

        paragraphs = content.split("\n\n")

        # Check last few paragraphs for conclusion
        for para in reversed(paragraphs[-3:]):
            para_lower = para.lower()
            if any(marker in para_lower for marker in conclusion_markers):
                return para[:200] + "..."

        # Default to last paragraph
        return paragraphs[-1][:200] + "..." if paragraphs else ""

    def _extract_established_facts(self, content: str) -> List[str]:
        """Extract facts that have been established in the section"""
        facts = []

        # Look for fact establishment patterns
        fact_patterns = [
            r"The record (?:shows|demonstrates|establishes) that ([^.]+)",
            r"It is undisputed that ([^.]+)",
            r"The evidence (?:shows|demonstrates) that ([^.]+)",
            r"(?:Plaintiff|Defendant) (?:admits|concedes) that ([^.]+)",
        ]

        for pattern in fact_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            facts.extend(matches)

        return facts[:5]  # Top 5 facts

    def _determine_transition_type(
        self, from_section: DraftedSection, to_section: DraftedSection
    ) -> str:
        """Determine the type of transition needed"""
        from_type = from_section.outline_section.section_type
        to_type = to_section.outline_section.section_type

        if (
            from_type == SectionType.STATEMENT_OF_FACTS
            and to_type == SectionType.LEGAL_STANDARD
        ):
            return "facts_to_law"
        elif (
            from_type == SectionType.LEGAL_STANDARD and to_type == SectionType.ARGUMENT
        ):
            return "law_to_application"
        elif from_type == SectionType.ARGUMENT and to_type == SectionType.ARGUMENT:
            return "argument_progression"
        elif to_type == SectionType.CONCLUSION:
            return "to_conclusion"
        else:
            return "general_progression"

    def _extract_conclusions(self, content: str) -> List[str]:
        """Extract key conclusions from section content"""
        conclusions = []

        # Split into paragraphs
        paragraphs = content.split("\n\n")

        # Look in last 3 paragraphs
        for para in paragraphs[-3:]:
            if any(
                marker in para.lower()
                for marker in [
                    "therefore",
                    "accordingly",
                    "thus",
                    "established",
                    "demonstrated",
                ]
            ):
                # Extract the key sentence
                sentences = para.split(". ")
                for sent in sentences:
                    if any(
                        marker in sent.lower()
                        for marker in ["therefore", "accordingly", "thus"]
                    ):
                        conclusions.append(sent.strip())
                        break

        return conclusions[:2]  # Top 2 conclusions

    def _extract_opening_premises(self, content: str) -> List[str]:
        """Extract opening premises from section content"""
        premises = []

        # Get first few paragraphs
        paragraphs = content.split("\n\n")[:3]

        for para in paragraphs:
            # Look for premise indicators
            if any(
                indicator in para.lower()
                for indicator in [
                    "principle",
                    "rule",
                    "standard",
                    "established",
                    "recognized",
                ]
            ):
                # Extract first sentence as premise
                first_sentence = para.split(". ")[0]
                premises.append(first_sentence)

        return premises[:2]  # Top 2 premises

    async def _analyze_expansion_opportunities(
        self, content: str, section: OutlineSection
    ) -> str:
        """Analyze where and how to expand section"""

        analysis = []

        # Check coverage of outline points
        uncovered_points = []
        content_lower = content.lower()
        for point in section.content_points:
            if not any(
                keyword in content_lower for keyword in point.lower().split()[:3]
            ):
                uncovered_points.append(point)

        if uncovered_points:
            analysis.append(f"Uncovered outline points: {', '.join(uncovered_points)}")

        # Check citation depth
        citations = self._extract_citations_from_text(content)
        citations_without_parentheticals = [
            c for c in citations if not re.search(r"\([^)]+\)$", c)
        ]
        if citations_without_parentheticals:
            analysis.append(
                f"Citations lacking parenthetical explanations: {len(citations_without_parentheticals)}"
            )

        # Check for thin paragraphs
        paragraphs = content.split("\n\n")
        thin_paragraphs = [i for i, p in enumerate(paragraphs) if len(p.split()) < 50]
        if thin_paragraphs:
            analysis.append(
                f"Thin paragraphs that need development: {len(thin_paragraphs)}"
            )

        # Check for missing counter-arguments
        if "even if" not in content_lower and "assuming" not in content_lower:
            analysis.append("No counter-argument analysis found")

        # Check for missing policy discussion
        if not any(
            term in content_lower for term in ["policy", "purpose", "rationale"]
        ):
            analysis.append("No policy discussion found")

        return "\n".join(analysis) if analysis else "Section appears well-developed"

    def _filter_relevant_facts(
        self, facts: List[Any], section: OutlineSection
    ) -> List[Any]:
        """Filter facts relevant to a specific section"""
        relevant = []

        # Create search terms from section content
        search_terms = []
        search_terms.extend(section.content_points)
        search_terms.extend(section.legal_authorities)
        search_terms.append(section.title)

        # Score each fact by relevance
        for fact in facts:
            relevance_score = 0
            fact_content = (
                fact.content.lower() if hasattr(fact, "content") else str(fact).lower()
            )

            for term in search_terms:
                if term.lower() in fact_content:
                    relevance_score += 1

            if relevance_score > 0:
                relevant.append((relevance_score, fact))

        # Sort by relevance and return top facts
        relevant.sort(key=lambda x: x[0], reverse=True)
        return [fact for _, fact in relevant[:5]]

    def _filter_relevant_authorities(
        self, authorities: List[Dict[str, Any]], required_citations: List[str]
    ) -> List[Dict[str, Any]]:
        """Filter authorities relevant to required citations"""
        relevant = []

        for auth in authorities:
            citation = auth.get("citation", "")
            # Check if this authority matches any required citation
            for required in required_citations:
                if (
                    required.lower() in citation.lower()
                    or citation.lower() in required.lower()
                ):
                    relevant.append(auth)
                    break

        return relevant

    def _get_unused_authorities(
        self, section: DraftedSection, case_context: Dict[str, Any]
    ) -> str:
        """Get authorities not yet used in section"""
        used_citations = set(section.citations_used)

        unused = []
        for auth in case_context.get("legal_authorities", []):
            citation = auth.get("citation", "")
            if citation and citation not in used_citations:
                unused.append(f"- {citation}: {auth.get('content', '')[:100]}...")

        return "\n".join(unused[:5]) if unused else "All key authorities have been used"

    def _get_additional_facts(
        self, section: DraftedSection, case_context: Dict[str, Any]
    ) -> str:
        """Get additional facts not yet used"""
        content_lower = section.content.lower()

        unused_facts = []
        for fact in case_context.get("case_facts", [])[:10]:
            fact_text = fact.content if hasattr(fact, "content") else str(fact)
            # Simple check if fact is mentioned
            key_words = fact_text.lower().split()[:5]  # First 5 words
            if not any(word in content_lower for word in key_words):
                unused_facts.append(f"- {fact_text[:150]}...")

        return (
            "\n".join(unused_facts[:5])
            if unused_facts
            else "Key facts have been incorporated"
        )

    def _get_policy_considerations(self, section: OutlineSection) -> str:
        """Get relevant policy considerations for section type"""
        policy_map = {
            SectionType.ARGUMENT: [
                "- Promoting judicial efficiency",
                "- Ensuring fair notice to defendants",
                "- Protecting fundamental rights",
                "- Maintaining consistency in legal standards",
            ],
            SectionType.LEGAL_STANDARD: [
                "- Legislative intent behind the statute",
                "- Public policy goals of the legal framework",
                "- Balance between competing interests",
            ],
        }

        policies = policy_map.get(
            section.section_type, ["General policy considerations"]
        )
        return "\n".join(policies)

    def _integrate_expansion_seamlessly(self, original: str, expansion: str) -> str:
        """Integrate expansion content seamlessly"""

        # Remove expansion markers if present
        expansion = expansion.replace("[EXPANSION START]", "").replace(
            "[EXPANSION END]", ""
        )

        # Parse both contents
        original_paragraphs = original.split("\n\n")
        expansion_parts = expansion.split("\n\n")

        # Smart integration based on content
        integrated = []
        expansion_idx = 0

        for i, para in enumerate(original_paragraphs):
            integrated.append(para)

            # Insert expansion after substantive paragraphs
            if (
                len(para.split()) > 100
                and expansion_idx < len(expansion_parts)
                and i < len(original_paragraphs) - 2
            ):  # Not near conclusion
                # Check if expansion fits contextually
                if self._check_contextual_fit(para, expansion_parts[expansion_idx]):
                    integrated.append(expansion_parts[expansion_idx])
                    expansion_idx += 1

        # Add remaining expansion before conclusion
        while expansion_idx < len(expansion_parts):
            # Insert before last paragraph (usually conclusion)
            integrated.insert(-1, expansion_parts[expansion_idx])
            expansion_idx += 1

        return "\n\n".join(integrated)

    def _check_contextual_fit(self, paragraph: str, expansion: str) -> bool:
        """Check if expansion fits contextually after paragraph"""

        # Get last sentence of paragraph
        last_sentence = paragraph.split(". ")[-1].lower()

        # Get first sentence of expansion
        first_sentence = expansion.split(". ")[0].lower()

        # Check for logical connectors
        connectors = [
            "furthermore",
            "moreover",
            "additionally",
            "this principle",
            "this rule",
            "this analysis",
            "similarly",
        ]

        return any(conn in first_sentence for conn in connectors)

    async def _enhance_citations(
        self, content: str, authorities: List[Dict[str, Any]]
    ) -> str:
        """Enhance citations with parentheticals and signals"""

        enhanced = content

        for auth in authorities:
            citation = auth.get("citation", "")
            if citation in content:
                # Check if citation already has parenthetical
                pattern = re.compile(rf"{re.escape(citation)}(?:\s*\([^)]+\))?")
                matches = pattern.findall(enhanced)

                for match in matches:
                    if "(" not in match:  # No parenthetical yet
                        # Create parenthetical from authority content
                        holding = auth.get("content", "")[:100]
                        parenthetical = f" (holding that {holding.lower()})"
                        enhanced = enhanced.replace(match, match + parenthetical, 1)

        return enhanced

    def _update_document_context(self, section: DraftedSection):
        """Update document-wide context with section information"""

        # Update citations used
        self.document_context["citations_used"].update(section.citations_used)

        # Extract and update key arguments
        if section.outline_section.section_type == SectionType.ARGUMENT:
            key_argument = self._extract_section_conclusion(section.content)[:100]
            self.document_context["key_arguments"].append(
                {"section": section.outline_section.title, "argument": key_argument}
            )

        # Update terminology
        terms = self._extract_key_terms(section.content)
        for term in terms:
            if term not in self.document_context["terminology"]:
                self.document_context["terminology"][term] = section.outline_section.id

    def _extract_key_terms(self, content: str) -> List[str]:
        """Extract key legal terms from content"""

        # Common legal terms to track
        term_patterns = [
            r"\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b",  # Capitalized terms
            r"\b(?:plaintiff|defendant|appellant|appellee)\b",
            r"\b(?:agreement|contract|lease|license)\b",
            r"\b(?:negligence|liability|damages|breach)\b",
        ]

        terms = []
        for pattern in term_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            terms.extend(matches)

        return list(set(terms))[:20]  # Top 20 unique terms

    def _build_fact_chronology(self, facts: List[Any]) -> List[Dict[str, Any]]:
        """Build chronological ordering of facts"""

        chronology = []

        for fact in facts:
            content = fact.content if hasattr(fact, "content") else str(fact)

            # Extract dates
            date_patterns = [
                r"(\d{1,2}/\d{1,2}/\d{2,4})",
                r"(\w+\s+\d{1,2},\s+\d{4})",
                r"(\d{4}-\d{2}-\d{2})",
            ]

            for pattern in date_patterns:
                dates = re.findall(pattern, content)
                for date in dates:
                    chronology.append(
                        {
                            "date": date,
                            "description": content[:100],
                            "source": fact.metadata.get("document_name", "Unknown")
                            if hasattr(fact, "metadata")
                            else "Unknown",
                        }
                    )

        # Sort by date (simple approach - would need proper date parsing in production)
        return chronology[:20]  # Top 20 events

    def _extract_case_themes(self, context: Dict[str, Any]) -> List[str]:
        """Extract overarching themes from case context"""

        themes = []

        # Analyze content for common themes
        all_content = []
        for category in ["case_facts", "expert_reports", "legal_authorities"]:
            for item in context.get(category, [])[:10]:
                if hasattr(item, "content"):
                    all_content.append(item.content)
                elif isinstance(item, dict):
                    all_content.append(item.get("content", ""))

        combined_text = " ".join(all_content).lower()

        # Theme detection patterns
        theme_patterns = {
            "safety": ["safety", "dangerous", "hazard", "risk", "harm"],
            "corporate_responsibility": [
                "corporate",
                "company",
                "business",
                "employer",
            ],
            "negligence": ["negligent", "breach", "duty", "reasonable care"],
            "causation": ["caused", "resulted", "led to", "because of"],
            "damages": ["damages", "injury", "harm", "loss", "suffering"],
        }

        for theme, keywords in theme_patterns.items():
            if sum(1 for kw in keywords if kw in combined_text) >= 2:
                themes.append(theme)

        return themes[:5]  # Top 5 themes

    def _summarize_previous_sections(self, sections: List[DraftedSection]) -> str:
        """Summarize previous sections for consistency check"""

        summaries = []
        for section in sections[-3:]:  # Last 3 sections
            summary = f"{section.outline_section.title}: "
            summary += self._extract_section_conclusion(section.content)[:100]
            summaries.append(summary)

        return "\n".join(summaries)

    async def _resolve_consistency_issues(
        self, section: DraftedSection, issues: List[Dict[str, Any]]
    ) -> DraftedSection:
        """Resolve identified consistency issues"""

        try:
            if not issues:
                return section

            fixes_needed = []
            for issue in issues:
                fixes_needed.append(
                    f"- {issue['type']}: {issue['issue']} -> {issue.get('suggestion', 'needs fix')}"
                )

            fix_prompt = f"""Fix the following consistency issues in this section:

Issues to Fix:
{chr(10).join(fixes_needed)}

Current Section:
{section.content}

Make minimal changes to resolve the issues while preserving the substance."""

            fixed_content_result = await self.section_writer.run(fix_prompt)
            fixed_content = (
                str(fixed_content_result.data)
                if hasattr(fixed_content_result, "data")
                else str(fixed_content_result)
            )

            section.content = fixed_content
            section.word_count = len(fixed_content.split())
            section.consistency_score = 1.0 - (len(issues) / 10.0)  # Simple scoring

            return section

        except Exception as e:
            logger.error(f"Error resolving consistency issues: {str(e)}")
            return section

    async def _review_citations(self, motion_draft: MotionDraft) -> Dict[str, Any]:
        """Review all citations in the motion"""

        try:
            all_citations = []
            for section in motion_draft.sections:
                for citation in section.citations_used:
                    all_citations.append(
                        {"citation": citation, "section": section.outline_section.id}
                    )

            citation_text = "\n".join([c["citation"] for c in all_citations])

            review_prompt = f"""Review these legal citations for format and accuracy:

{citation_text}

Check each citation for:
1. Proper Bluebook format
2. Missing pincites where needed
3. Consistent formatting across document
4. Any obviously fictional citations

Return JSON with findings."""

            result_raw = await self.citation_verifier.run(review_prompt)
            result = result_raw.data if hasattr(result_raw, "data") else result_raw

            return (
                result
                if isinstance(result, dict)
                else {"format_issues": [], "fictional_citations": []}
            )

        except Exception as e:
            logger.error(f"Error reviewing citations: {str(e)}")
            return {"format_issues": [], "fictional_citations": []}

    async def _review_consistency(self, motion_draft: MotionDraft) -> Dict[str, Any]:
        """Review document-wide consistency"""

        try:
            # Compile key information
            consistency_data = {
                "terminology": list(
                    self.document_context.get("terminology", {}).keys()
                )[:20],
                "party_names": self._extract_party_names(motion_draft),
                "key_dates": self._extract_key_dates(motion_draft),
                "citations": list(motion_draft.citation_index.keys())[:20],
            }

            full_text = self._compile_full_document(motion_draft)[
                :5000
            ]  # First 5000 chars

            review_prompt = f"""Check this document for consistency issues:

Document Sample:
{full_text}

Key Terms to Check:
{json.dumps(consistency_data, indent=2)}

Look for:
1. Inconsistent terminology (same concept, different terms)
2. Inconsistent party references
3. Conflicting dates or facts
4. Contradictory legal positions

Return JSON with any inconsistencies found."""

            result_raw = await self.consistency_checker.run(review_prompt)
            result = result_raw.data if hasattr(result_raw, "data") else result_raw

            return result if isinstance(result, dict) else {"inconsistencies": []}

        except Exception as e:
            logger.error(f"Error reviewing consistency: {str(e)}")
            return {"inconsistencies": []}

    async def _review_argument_strength(
        self, motion_draft: MotionDraft, case_context: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Review strength of legal arguments"""

        try:
            section_strengths = {}
            weak_sections = []

            for section in motion_draft.sections:
                if section.outline_section.section_type == SectionType.ARGUMENT:
                    strength = section.argument_strength_score
                    section_strengths[section.outline_section.id] = strength

                    if strength < 0.6:
                        weak_sections.append(
                            {
                                "section_id": section.outline_section.id,
                                "title": section.outline_section.title,
                                "strength": strength,
                                "issues": self._identify_argument_weaknesses(section),
                            }
                        )

            avg_strength = sum(section_strengths.values()) / max(
                len(section_strengths), 1
            )

            return {
                "average_strength": avg_strength,
                "section_strengths": section_strengths,
                "weak_sections": weak_sections,
            }

        except Exception as e:
            logger.error(f"Error reviewing argument strength: {str(e)}")
            return {
                "average_strength": 0.7,
                "section_strengths": {},
                "weak_sections": [],
            }

    def _identify_argument_weaknesses(self, section: DraftedSection) -> List[str]:
        """Identify specific weaknesses in legal argument"""

        weaknesses = []
        content_lower = section.content.lower()

        # Check for missing IRAC elements
        if "issue" not in content_lower:
            weaknesses.append("No clear issue statement")
        if "rule" not in content_lower and "law" not in content_lower:
            weaknesses.append("Weak legal rule statement")
        if "application" not in content_lower and "apply" not in content_lower:
            weaknesses.append("Limited application of law to facts")

        # Check citation support
        paragraphs = section.content.split("\n\n")
        unsupported_paragraphs = sum(
            1
            for para in paragraphs
            if len(para.split()) > 50
            and not any(c in para for c in section.citations_used)
        )
        if unsupported_paragraphs > 2:
            weaknesses.append(
                f"{unsupported_paragraphs} paragraphs lack citation support"
            )

        # Check for counter-argument handling
        if not any(
            term in content_lower for term in ["even if", "assuming", "although"]
        ):
            weaknesses.append("No counter-argument analysis")

        return weaknesses

    def _compile_full_document(self, motion_draft: MotionDraft) -> str:
        """Compile full document text with transitions"""

        full_text = f"{motion_draft.title}\n\n{motion_draft.case_name}\n\n"

        for i, section in enumerate(motion_draft.sections):
            # Section heading
            full_text += f"{section.outline_section.title}\n\n"
            full_text += section.content

            # Add transition if available
            if "to_next" in section.transitions and i < len(motion_draft.sections) - 1:
                full_text += f"\n\n{section.transitions['to_next']}"

            full_text += "\n\n"

        return full_text

    def _extract_party_names(self, motion_draft: MotionDraft) -> List[str]:
        """Extract party names from document"""

        party_names = set()

        # Common patterns
        patterns = [
            r"Plaintiff[s]?\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)",
            r"Defendant[s]?\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)",
            r"([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+v\.\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)",
        ]

        for section in motion_draft.sections[:3]:  # Check first 3 sections
            for pattern in patterns:
                matches = re.findall(pattern, section.content)
                if isinstance(matches[0], tuple) if matches else False:
                    for match in matches:
                        party_names.update(match)
                else:
                    party_names.update(matches)

        return list(party_names)[:10]

    def _extract_key_dates(self, motion_draft: MotionDraft) -> List[str]:
        """Extract key dates from document"""

        dates = set()
        date_patterns = [
            r"\b\d{1,2}/\d{1,2}/\d{2,4}\b",
            r"\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}\b",
            r"\b\d{4}-\d{2}-\d{2}\b",
        ]

        for section in motion_draft.sections:
            for pattern in date_patterns:
                matches = re.findall(pattern, section.content)
                dates.update(matches)

        return list(dates)[:20]

    async def _apply_critical_fixes(
        self,
        motion_draft: MotionDraft,
        overall_review: Dict[str, Any],
        citation_review: Dict[str, Any],
        consistency_review: Dict[str, Any],
    ) -> MotionDraft:
        """Apply critical fixes based on review feedback"""

        try:
            logger.info("Applying critical fixes to motion draft")

            # Fix citation issues
            if citation_review.get("format_issues"):
                motion_draft = await self._fix_citation_formatting(
                    motion_draft, citation_review["format_issues"]
                )

            # Fix consistency issues
            if consistency_review.get("inconsistencies"):
                motion_draft = await self._fix_terminology_inconsistencies(
                    motion_draft, consistency_review["inconsistencies"]
                )

            # Fix weak sections
            for section_id in overall_review.get("revision_priorities", [])[:3]:
                for i, section in enumerate(motion_draft.sections):
                    if section.outline_section.id == section_id:
                        logger.info(
                            f"Revising weak section: {section.outline_section.title}"
                        )
                        motion_draft.sections[i] = await self._revise_weak_section(
                            section,
                            overall_review.get("specific_edits", {}).get(
                                section_id, []
                            ),
                        )

            return motion_draft

        except Exception as e:
            logger.error(f"Error applying critical fixes: {str(e)}")
            return motion_draft

    async def _fix_citation_formatting(
        self, motion_draft: MotionDraft, format_issues: List[Dict[str, Any]]
    ) -> MotionDraft:
        """Fix citation formatting issues"""

        for section in motion_draft.sections:
            content = section.content

            for issue in format_issues:
                if issue.get("section") == section.outline_section.id:
                    old_citation = issue.get("citation", "")
                    corrected = issue.get("corrected", "")

                    if old_citation and corrected:
                        content = content.replace(old_citation, corrected)

            section.content = content

        return motion_draft

    async def _fix_terminology_inconsistencies(
        self, motion_draft: MotionDraft, inconsistencies: List[Dict[str, Any]]
    ) -> MotionDraft:
        """Fix terminology inconsistencies across document"""

        # Build replacement map
        replacements = {}
        for issue in inconsistencies:
            if issue["type"] == "terminology":
                variants = issue.get("variants", [])
                if variants:
                    preferred = variants[0]  # Use first as preferred
                    for variant in variants[1:]:
                        replacements[variant] = preferred

        # Apply replacements
        for section in motion_draft.sections:
            content = section.content
            for old_term, new_term in replacements.items():
                # Case-sensitive replacement
                content = content.replace(old_term, new_term)
                # Also handle lowercase
                content = content.replace(old_term.lower(), new_term.lower())

            section.content = content
            section.word_count = len(content.split())

        return motion_draft

    async def _revise_weak_section(
        self, section: DraftedSection, revision_notes: List[str]
    ) -> DraftedSection:
        """Revise a weak section based on feedback"""

        try:
            revision_prompt = f"""Revise this legal section to address the following issues:

Issues to Address:
{chr(10).join(f"- {note}" for note in revision_notes)}

Current Section:
{section.content}

Requirements:
1. Strengthen legal analysis
2. Add missing elements (citations, counter-arguments, etc.)
3. Improve argument flow and persuasiveness
4. Maintain current length or expand if needed
5. Preserve all existing strong points

Provide the complete revised section."""

            revised_content_result = await self.section_writer.run(revision_prompt)
            revised_content = (
                str(revised_content_result.data)
                if hasattr(revised_content_result, "data")
                else str(revised_content_result)
            )

            # Update section
            section.content = revised_content
            section.word_count = len(revised_content.split())
            section.needs_revision = False
            section.revision_notes = []

            # Recalculate quality scores
            section.confidence_score = self._calculate_section_confidence(
                revised_content, section.outline_section, section.word_count
            )
            section.argument_strength_score = self._calculate_argument_strength(
                revised_content, section.outline_section
            )

            return section

        except Exception as e:
            logger.error(f"Error revising weak section: {str(e)}")
            return section

    def _find_term_variants(self, term: str, motion_draft: MotionDraft) -> List[str]:
        """Find variants of a term used in the document"""

        variants = set([term])
        term_lower = term.lower()

        # Common variations to check
        if "agreement" in term_lower:
            check_terms = ["Agreement", "Contract", "contract", "agreement"]
        elif "plaintiff" in term_lower:
            check_terms = ["Plaintiff", "plaintiff", "Petitioner", "petitioner"]
        elif "defendant" in term_lower:
            check_terms = ["Defendant", "defendant", "Respondent", "respondent"]
        else:
            check_terms = [term, term.lower(), term.upper()]

        # Check each section
        for section in motion_draft.sections:
            for check_term in check_terms:
                if check_term in section.content and check_term != term:
                    variants.add(check_term)

        return list(variants)

    async def _preview_next_section(
        self, current_section: DraftedSection, next_section: OutlineSection
    ) -> str:
        """Create preview transition to next section"""

        try:
            prompt = f"""Create a brief preview transition (1 paragraph) that:

Current section: {current_section.outline_section.title}
Next section: {next_section.title}

Preview why the next section strengthens our position.
Use language like:
- "As we will demonstrate in the next section..."
- "This principle becomes even clearer when we examine..."
- "The following section will show..."

Keep it to 2-3 sentences."""

            preview_result = await self.transition_writer.run(prompt)
            preview = (
                str(preview_result.data)
                if hasattr(preview_result, "data")
                else str(preview_result)
            )
            return preview

        except Exception as e:
            logger.error(f"Error creating preview transition: {str(e)}")
            return ""

    def _generate_motion_title(self, outline: Dict[str, Any]) -> str:
        """Generate motion title from outline"""

        # Check various possible fields
        if "title" in outline:
            return outline["title"]
        if "motion_title" in outline:
            return outline["motion_title"]
        if "motion_type" in outline:
            return outline["motion_type"]

        # Try to infer from sections
        if "sections" in outline and outline["sections"]:
            # Look for motion type in first section
            first_section = outline["sections"][0]
            if "title" in first_section:
                return f"Motion: {first_section['title']}"

        return "Legal Motion"

    def export_to_docx(self, motion_draft: MotionDraft, output_path: str):
        """Export motion draft to DOCX format"""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Create document
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add title
        title = doc.add_heading(motion_draft.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add case caption
        doc.add_paragraph(motion_draft.case_name)
        doc.add_paragraph()

        # Add each section with transitions
        for i, section in enumerate(motion_draft.sections):
            # Section heading
            doc.add_heading(section.outline_section.title, level=1)

            # Add transition from previous if available
            if "from_previous" in section.transitions and i > 0:
                transition_para = doc.add_paragraph(
                    section.transitions["from_previous"]
                )
                transition_para.style.font.italic = True
                doc.add_paragraph()

            # Section content
            for paragraph in section.content.split("\n\n"):
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph)
                    p.style.font.size = Pt(12)
                    p.style.font.name = "Times New Roman"

            # Add transition to next if available
            if "to_next" in section.transitions and i < len(motion_draft.sections) - 1:
                doc.add_paragraph()
                transition_para = doc.add_paragraph(section.transitions["to_next"])
                transition_para.style.font.italic = True

        # Add metadata page
        doc.add_page_break()
        doc.add_heading("Document Metadata", level=1)

        # Basic metadata
        doc.add_paragraph(f"Total Words: {motion_draft.total_word_count:,}")
        doc.add_paragraph(f"Estimated Pages: {motion_draft.total_page_estimate}")
        doc.add_paragraph(
            f"Created: {motion_draft.creation_timestamp.strftime('%B %d, %Y at %I:%M %p')}"
        )
        doc.add_paragraph(f"Document Quality Score: {motion_draft.coherence_score:.2%}")

        # Quality metrics
        if motion_draft.quality_metrics:
            doc.add_heading("Quality Metrics", level=2)
            for metric, score in motion_draft.quality_metrics.items():
                doc.add_paragraph(f"• {metric.replace('_', ' ').title()}: {score:.2%}")

        # Citation index
        if motion_draft.citation_index:
            doc.add_heading("Citation Index", level=2)
            doc.add_paragraph(
                f"Total Unique Citations: {len(motion_draft.citation_index)}"
            )

            # List top citations
            sorted_citations = sorted(
                motion_draft.citation_index.items(),
                key=lambda x: len(x[1]),
                reverse=True,
            )[:10]

            for citation, sections in sorted_citations:
                doc.add_paragraph(f"• {citation} (used in {len(sections)} sections)")

        # Review notes
        if motion_draft.review_notes:
            doc.add_heading("Review Notes", level=2)
            for note in motion_draft.review_notes:
                doc.add_paragraph(f"• {note}")

        # Save document
        doc.save(output_path)
        logger.info(f"Motion exported to {output_path}")

    async def draft_motion_with_cache(
        self,
        outline: Dict[str, Any],
        database_name: str,
        target_length: DocumentLength = DocumentLength.MEDIUM,
        motion_title: Optional[str] = None,
        opposing_motion_text: Optional[str] = None,
        outline_id: Optional[str] = None,  # If already cached
    ) -> MotionDraft:
        """
        Enhanced motion drafting that caches outline and processes sections individually
        """
        start_time = datetime.utcnow()
        logger.info(
            f"[MOTION_DRAFTER] Starting cached motion draft for database: {database_name}"
        )

        # Cache the outline if not already cached
        if not outline_id:
            outline_id = await outline_cache.cache_outline(outline, database_name)
            logger.info(f"[MOTION_DRAFTER] Cached outline with ID: {outline_id}")
        else:
            # Verify outline exists in cache
            cached_outline = await outline_cache.get_outline(outline_id)
            if not cached_outline:
                # Re-cache if missing
                outline_id = await outline_cache.cache_outline(outline, database_name)
                logger.info(
                    f"[MOTION_DRAFTER] Re-cached missing outline with ID: {outline_id}"
                )

        # Get outline structure for processing
        outline_structure = await outline_cache.get_outline_structure(outline_id)
        if not outline_structure:
            raise ValueError(f"Failed to get outline structure for ID: {outline_id}")

        logger.info(f"[MOTION_DRAFTER] Processing {len(outline_structure)} sections")

        # Initialize timeout monitor
        timeout_monitor = TimeoutMonitor(
            operation_name=f"Motion Drafting ({database_name})",
            warning_threshold=60,
            critical_threshold=180,  # Increased to 3 minutes
        )
        timeout_monitor.log_progress("Cached motion drafting started")

        # Initialize document context with minimal data
        self.document_context = {
            "themes": outline.get("themes", outline.get("style_notes", [])),
            "key_arguments": [],
            "terminology": {},
            "citations_used": set(),
            "fact_chronology": [],
            "database_name": database_name,
            "opposing_motion_text": opposing_motion_text,
            "outline_id": outline_id,  # Store for reference
        }

        # Calculate word distribution based on structure
        min_pages, max_pages = target_length.value
        target_words = ((min_pages + max_pages) // 2) * self.words_per_page
        word_distribution = self._calculate_word_distribution_from_structure(
            outline_structure, target_words
        )

        # Retrieve case context once - convert structure back to sections for compatibility
        logger.info("[MOTION_DRAFTER] Retrieving case context")
        full_outline_sections = await self._convert_structure_to_sections(
            outline_id, outline_structure
        )
        case_context = await self._retrieve_enhanced_case_context(
            database_name, full_outline_sections
        )

        # Process sections individually
        drafted_sections = []
        total_words = 0

        for i, section_info in enumerate(outline_structure):
            section_start_time = datetime.utcnow()
            logger.info(
                f"[MOTION_DRAFTER] Processing section {i + 1}/{len(outline_structure)}: {section_info['heading']}"
            )

            # Update timeout monitor
            timeout_monitor.log_progress(
                f"Drafting section {i + 1}: {section_info['heading']}"
            )

            # Retrieve only this section from cache
            section_data = await outline_cache.get_section(
                outline_id, section_info["index"]
            )
            if not section_data:
                logger.error(f"Failed to retrieve section {i} from cache")
                continue

            # Create minimal OutlineSection for this section only
            outline_section = self._create_minimal_outline_section(
                section_data, section_info, word_distribution.get(f"section_{i}", 500)
            )

            # Build cumulative context (lightweight)
            cumulative_context = self._build_lightweight_cumulative_context(
                drafted_sections, case_context
            )

            # Draft the section
            try:
                drafted_section = await self._draft_section_efficiently(
                    outline_section,
                    section_data,  # Pass full section data separately
                    case_context,
                    cumulative_context,
                )

                # Expand if needed
                if drafted_section.word_count < outline_section.target_length * 0.9:
                    drafted_section = await self._expand_section_efficiently(
                        drafted_section,
                        section_data,
                        outline_section.target_length,
                        case_context,
                    )

                drafted_sections.append(drafted_section)
                total_words += drafted_section.word_count

                # Update document context incrementally
                self._update_document_context(drafted_section)

                section_duration = datetime.utcnow() - section_start_time
                logger.info(
                    f"[MOTION_DRAFTER] Section {i + 1} completed in {section_duration}"
                )

            except Exception as e:
                logger.error(
                    f"[MOTION_DRAFTER] Error drafting section {i + 1}: {str(e)}"
                )
                # Create placeholder section
                drafted_section = DraftedSection(
                    outline_section=outline_section,
                    content=f"[ERROR: Section could not be drafted - {str(e)}]",
                    word_count=0,
                    citations_used=[],
                    citations_verified={},
                    expansion_cycles=0,
                    confidence_score=0.0,
                )
                drafted_sections.append(drafted_section)

        # Create motion draft
        motion_draft = MotionDraft(
            title=motion_title or self._extract_title_from_cache(outline),
            case_name=database_name.replace("_", " ").title(),
            sections=drafted_sections,
            total_word_count=total_words,
            total_page_estimate=total_words // self.words_per_page,
            creation_timestamp=datetime.utcnow(),
            outline_source={
                "outline_id": outline_id
            },  # Store reference instead of full outline
        )

        # Simplified review process
        logger.info("[MOTION_DRAFTER] Starting review process")
        motion_draft = await self._lightweight_review_process(motion_draft)

        # Build citation index
        motion_draft.citation_index = self._build_citation_index(motion_draft)

        timeout_monitor.finish(success=True)

        total_duration = datetime.utcnow() - start_time
        logger.info(f"[MOTION_DRAFTER] Completed in {total_duration}")

        return motion_draft

    async def _convert_structure_to_sections(
        self, outline_id: str, outline_structure: List[Dict[str, Any]]
    ) -> List[OutlineSection]:
        """Convert outline structure back to OutlineSection objects for compatibility"""
        sections = []

        for section_info in outline_structure:
            try:
                # Get the full section data from cache
                section_data = await outline_cache.get_section(
                    outline_id, section_info["index"]
                )
                if not section_data:
                    logger.warning(
                        f"Could not retrieve section {section_info['index']} from cache"
                    )
                    continue

                # Extract content points safely - handle both strings and dicts
                content_points = []
                content_items = section_data.get("content", [])[:3]  # Limit to first 3
                for item in content_items:
                    if isinstance(item, str):
                        content_points.append(item)
                    elif isinstance(item, dict):
                        # Extract text from dict structure
                        text = item.get(
                            "text", item.get("content", item.get("value", str(item)))
                        )
                        content_points.append(str(text))
                    else:
                        content_points.append(str(item))

                # Create minimal OutlineSection with safe defaults
                section = OutlineSection(
                    id=f"section_{section_info['index']}",
                    title=section_info.get(
                        "heading", f"Section {section_info['index'] + 1}"
                    ),
                    section_type=SectionType.ARGUMENT,  # Default type
                    content_points=content_points,
                    legal_authorities=section_data.get("legal_authorities", []),
                    target_length=500,  # Default target length
                    context_summary=section_data.get("summary", ""),
                    hook_options=section_data.get("hook_options", []),
                    themes=section_data.get("themes", []),
                    key_facts=section_data.get("key_facts", []),
                    counter_arguments=section_data.get("counter_arguments", []),
                )
                sections.append(section)

            except Exception as e:
                logger.warning(
                    f"Error converting section {section_info['index']}: {str(e)}"
                )
                continue

        return sections

    # Add this helper method
    def _create_minimal_outline_section(
        self,
        section_data: Dict[str, Any],
        section_info: Dict[str, Any],
        target_length: int,
    ) -> OutlineSection:
        """Create a minimal OutlineSection object for processing"""

        # Extract only essential content points (limit to prevent overflow)
        content_points = []
        content_items = section_data.get("content", [])[:3]  # Only first 3 items

        for item in content_items:
            if item.get("type") == "field":
                label = item.get("label", "")
                value = item.get("value", "")[:200]  # Truncate to 200 chars
                content_points.append(f"{label}: {value}")
            elif item.get("type") == "paragraph":
                content_points.append(item.get("text", "")[:200])

        # Extract limited authorities
        legal_authorities = []
        for item in content_items:
            if (
                item.get("type") == "list"
                and "authorities" in item.get("label", "").lower()
            ):
                legal_authorities.extend(item.get("items", [])[:5])  # Max 5 authorities

        return OutlineSection(
            id=f"section_{section_info['index']}",
            title=section_info["heading"],
            section_type=self._determine_section_type(section_info["heading"]),
            content_points=content_points[:5],  # Max 5 points
            legal_authorities=legal_authorities[:5],  # Max 5 authorities
            target_length=target_length,
            context_summary=f"Section {section_info['index'] + 1} of {section_info.get('type', 'standard')} type",
        )

    # Add this efficient drafting method
    async def _draft_section_efficiently(
        self,
        outline_section: OutlineSection,
        full_section_data: Dict[str, Any],
        case_context: Dict[str, Any],
        cumulative_context: Dict[str, Any],
    ) -> DraftedSection:
        """Draft a section efficiently without token overflow"""

        try:
            # Extract the most important content from full section data
            essential_content = self._extract_essential_content(full_section_data)

            # Create focused prompt
            drafting_prompt = f"""Draft this section of the legal motion:

    Section: {outline_section.title}
    Type: {outline_section.section_type.value}
    Target Length: {outline_section.target_length} words (MINIMUM)

    Key Points to Address:
    {chr(10).join(f"- {point}" for point in essential_content["key_points"][:5])}

    Required Authorities:
    {chr(10).join(f"- {auth}" for auth in essential_content["authorities"][:5])}

    Context from Previous Sections:
    {cumulative_context.get("summary", "This is the first section.")}

    REQUIREMENTS:
    1. Write {outline_section.target_length} words MINIMUM
    2. Use formal legal writing style
    3. Include all required authorities with parentheticals
    4. Develop each point with detailed analysis
    5. Include transitions and topic sentences
    6. Apply IRAC/CRAC structure where appropriate"""

            # Generate content
            result = await asyncio.wait_for(
                self.section_writer.run(drafting_prompt),
                timeout=45,  # 45 second timeout per section
            )

            content = str(result.data) if hasattr(result, "data") else str(result)

            # Create drafted section
            word_count = len(content.split())
            return DraftedSection(
                outline_section=outline_section,
                content=content,
                word_count=word_count,
                citations_used=self._extract_citations_from_text(content),
                citations_verified={},
                expansion_cycles=1,
                confidence_score=self._calculate_section_confidence(
                    content, outline_section, word_count
                ),
            )

        except asyncio.TimeoutError:
            logger.error(f"Timeout drafting section: {outline_section.title}")
            raise
        except Exception as e:
            logger.error(f"Error drafting section efficiently: {str(e)}")
            raise

    # Add this helper to extract essential content
    def _extract_essential_content(
        self, section_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Extract essential content from section data without overwhelming tokens"""

        essential = {"key_points": [], "authorities": [], "facts": [], "themes": []}

        # Process content items
        for item in section_data.get("content", []):
            item_type = item.get("type", "")

            if item_type == "field":
                label = item.get("label", "")
                value = item.get("value", "")

                # Handle multi-option values
                if "||" in value:
                    # Take first option only
                    value = value.split("||")[0].strip()
                    if value.startswith("Option"):
                        value = (
                            value.split(":", 1)[1].strip() if ":" in value else value
                        )

                # Truncate and clean
                value = value.strip('"').strip()[:300]  # 300 char limit

                if label == "Theme":
                    essential["themes"].append(value)
                elif label == "Summary":
                    essential["key_points"].insert(0, value)  # Put summary first
                else:
                    essential["key_points"].append(f"{label}: {value}")

            elif item_type == "list":
                label = item.get("label", "").lower()
                items = item.get("items", [])[:5]  # Max 5 items per list

                if "authorit" in label:
                    essential["authorities"].extend(items)
                elif "fact" in label:
                    essential["facts"].extend(items)
                else:
                    # Convert list items to key points
                    for list_item in items[:3]:  # Max 3 items
                        if isinstance(list_item, str):
                            essential["key_points"].append(list_item[:200])

        # Limit totals
        essential["key_points"] = essential["key_points"][:7]
        essential["authorities"] = essential["authorities"][:7]
        essential["facts"] = essential["facts"][:5]
        essential["themes"] = essential["themes"][:3]

        return essential

    # Add lightweight context builder
    def _build_lightweight_cumulative_context(
        self, drafted_sections: List[DraftedSection], case_context: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Build minimal cumulative context to prevent token overflow"""

        context = {"summary": "", "citations_used": set(), "key_holdings": []}

        if not drafted_sections:
            context["summary"] = "This is the first section of the motion."
            return context

        # Build brief summary from last 2 sections only
        recent_sections = drafted_sections[-2:]
        summaries = []

        for section in recent_sections:
            # Extract brief conclusion (first 100 chars of last paragraph)
            paragraphs = section.content.split("\n\n")
            if paragraphs:
                last_para = paragraphs[-1][:100]
                summaries.append(f"{section.outline_section.title}: {last_para}...")

            # Collect citations (limit to 10)
            for citation in section.citations_used[:10]:
                context["citations_used"].add(citation)

        context["summary"] = " | ".join(summaries)
        context["citations_used"] = list(context["citations_used"])[:15]

        return context

    # Add this helper for word distribution from structure
    def _calculate_word_distribution_from_structure(
        self, outline_structure: List[Dict[str, Any]], target_words: int
    ) -> Dict[str, int]:
        """Calculate word distribution based on section structure"""

        distribution = {}
        weights = {}
        total_weight = 0

        for i, section in enumerate(outline_structure):
            section_type = section.get("type", "standard")
            heading = section.get("heading", "").lower()

            # Assign weights based on type and heading
            if section_type == "introduction" or "introduction" in heading:
                weight = 0.5
            elif section_type == "facts" or "facts" in heading:
                weight = 1.5
            elif section_type == "conclusion" or "conclusion" in heading:
                weight = 0.5
            elif section_type == "argument":
                weight = 1.2
                # Add weight for content complexity
                weight += 0.1 * min(section.get("content_items", 0), 5)
            else:
                weight = 1.0

            weights[f"section_{i}"] = weight
            total_weight += weight

        # Distribute words based on weights
        for section_id, weight in weights.items():
            distribution[section_id] = int((weight / total_weight) * target_words)

        return distribution

    # Add simplified expansion method
    async def _expand_section_efficiently(
        self,
        drafted_section: DraftedSection,
        full_section_data: Dict[str, Any],
        target_length: int,
        case_context: Dict[str, Any],
    ) -> DraftedSection:
        """Expand section efficiently without overwhelming context"""

        try:
            current_length = drafted_section.word_count
            expansion_needed = target_length - current_length

            if expansion_needed <= 0:
                return drafted_section

            # Extract unused content from full section data
            unused_content = self._extract_unused_content(
                drafted_section.content, full_section_data
            )

            expansion_prompt = f"""Expand this legal section by adding {expansion_needed} words.

    Current Section ({current_length} words):
    {drafted_section.content[:2000]}... [truncated]

    Unused Content Points:
    {chr(10).join(f"- {point}" for point in unused_content[:5])}

    REQUIREMENTS:
    1. Add {expansion_needed} words of substantive content
    2. DO NOT remove any existing content
    3. Add deeper analysis and more detailed application
    4. Maintain formal legal style
    5. Include smooth transitions

    Provide the COMPLETE expanded section."""

            result = await asyncio.wait_for(
                self.section_expander.run(expansion_prompt), timeout=30
            )

            expanded_content = (
                str(result.data) if hasattr(result, "data") else str(result)
            )

            # Update section
            drafted_section.content = expanded_content
            drafted_section.word_count = len(expanded_content.split())
            drafted_section.expansion_cycles += 1

            return drafted_section

        except Exception as e:
            logger.error(f"Error in efficient expansion: {str(e)}")
            return drafted_section

    # Add helper to extract unused content
    def _extract_unused_content(
        self, current_content: str, full_section_data: Dict[str, Any]
    ) -> List[str]:
        """Extract content points not yet used in the section"""

        unused = []
        content_lower = current_content.lower()

        # Check all content items
        for item in full_section_data.get("content", []):
            if item.get("type") == "field":
                value = item.get("value", "")
                # Check if this content is not yet used
                if value and value[:50].lower() not in content_lower:
                    unused.append(value[:200])

            elif item.get("type") == "list":
                for list_item in item.get("items", [])[:3]:
                    if (
                        isinstance(list_item, str)
                        and list_item[:30].lower() not in content_lower
                    ):
                        unused.append(list_item[:150])

        return unused[:7]  # Return max 7 unused points

    # Add lightweight review process
    async def _lightweight_review_process(
        self, motion_draft: MotionDraft
    ) -> MotionDraft:
        """Simplified review process that doesn't overwhelm the system"""

        try:
            # Basic quality score calculation
            total_score = 0

            for section in motion_draft.sections:
                # Check basic quality metrics
                if section.word_count >= section.outline_section.target_length * 0.9:
                    total_score += 0.2
                if len(section.citations_used) >= 2:
                    total_score += 0.2
                if section.confidence_score > 0.7:
                    total_score += 0.1

            motion_draft.coherence_score = min(
                total_score / len(motion_draft.sections), 0.9
            )

            # Add basic review notes
            if motion_draft.coherence_score < 0.7:
                motion_draft.review_notes.append(
                    "Some sections may need expansion or additional citations"
                )

            motion_draft.quality_metrics = {
                "sections_complete": len(
                    [s for s in motion_draft.sections if s.word_count > 0]
                ),
                "average_confidence": sum(
                    s.confidence_score for s in motion_draft.sections
                )
                / len(motion_draft.sections),
                "total_citations": len(motion_draft.citation_index),
            }

            return motion_draft

        except Exception as e:
            logger.error(f"Error in lightweight review: {str(e)}")
            motion_draft.coherence_score = 0.7
            return motion_draft

    # Helper to extract title from cached outline
    def _extract_title_from_cache(self, outline: Dict[str, Any]) -> str:
        """Extract title from outline with minimal processing"""
        return outline.get("title") or outline.get("motion_title") or "Legal Motion"

    async def draft_proper_legal_motion(
        self,
        outline: Dict[str, Any],
        database_name: str,
        target_length: DocumentLength = DocumentLength.MEDIUM,
        motion_title: Optional[str] = None,
        case_caption: Optional[str] = None,
        opposing_motion_text: Optional[str] = None,
    ) -> MotionDraft:
        """
        Draft a properly structured legal motion that follows standard format
        """
        start_time = datetime.utcnow()
        logger.info(
            f"[MOTION_DRAFTER] Starting proper legal motion draft for database: {database_name}"
        )

        # Initialize timeout monitor
        timeout_monitor = TimeoutMonitor(
            operation_name=f"Legal Motion Drafting ({database_name})",
            warning_threshold=60,
            critical_threshold=180,
        )

        # Extract case information
        if not case_caption:
            case_caption = self._extract_case_caption(outline, database_name)

        if not motion_title:
            motion_title = self._extract_motion_title(outline)

        # Convert to proper structure
        logger.info("[MOTION_DRAFTER] Converting outline to proper legal structure")
        proper_structure = self._convert_to_proper_legal_structure(
            outline, target_length
        )

        # Retrieve case context
        logger.info("[MOTION_DRAFTER] Retrieving case context")
        case_context = await self._retrieve_enhanced_case_context(
            database_name, self._structure_to_outline_sections(proper_structure)
        )

        # Initialize document context
        self.document_context = {
            "themes": self._extract_themes_from_outline(outline),
            "key_arguments": [],
            "terminology": {},
            "citations_used": set(),
            "fact_chronology": [],
            "database_name": database_name,
            "opposing_motion_text": opposing_motion_text,
            "case_caption": case_caption,
            "motion_title": motion_title,
        }

        # Draft sections in proper order
        drafted_sections = []
        total_words = 0

        # 1. Introduction
        logger.info("[MOTION_DRAFTER] Drafting INTRODUCTION")
        intro_section = await self._draft_introduction_section(
            proper_structure["introduction"], case_context, opposing_motion_text
        )
        drafted_sections.append(intro_section)
        total_words += intro_section.word_count

        # 2. Statement of Facts (skip if ERROR as mentioned)
        logger.info("[MOTION_DRAFTER] Drafting STATEMENT OF FACTS")
        facts_section = await self._draft_facts_section(
            proper_structure["facts"], case_context
        )
        drafted_sections.append(facts_section)
        total_words += facts_section.word_count

        # 3. Legal Standard
        logger.info("[MOTION_DRAFTER] Drafting LEGAL STANDARD")
        legal_section = await self._draft_legal_standard_section(
            proper_structure["legal_standard"], case_context
        )
        drafted_sections.append(legal_section)
        total_words += legal_section.word_count

        # 4. Arguments
        logger.info("[MOTION_DRAFTER] Drafting ARGUMENT sections")
        for i, arg_section in enumerate(proper_structure["arguments"]):
            logger.info(
                f"[MOTION_DRAFTER] Drafting Argument {i + 1}: {arg_section.title}"
            )

            argument = await self._draft_argument_section_proper(
                arg_section, case_context, i + 1, len(proper_structure["arguments"])
            )
            drafted_sections.append(argument)
            total_words += argument.word_count

        # 5. Conclusion
        logger.info("[MOTION_DRAFTER] Drafting CONCLUSION")
        conclusion_section = await self._draft_conclusion_section(
            proper_structure["conclusion"], case_context, drafted_sections
        )
        drafted_sections.append(conclusion_section)
        total_words += conclusion_section.word_count

        # Create motion draft
        motion_draft = MotionDraft(
            title=motion_title,
            case_name=case_caption,
            sections=drafted_sections,
            total_word_count=total_words,
            total_page_estimate=total_words // self.words_per_page,
            creation_timestamp=datetime.utcnow(),
            outline_source=outline,
        )

        # Quick review
        motion_draft = await self._quick_legal_review(motion_draft)

        # Build citation index
        motion_draft.citation_index = self._build_citation_index(motion_draft)

        timeout_monitor.finish(success=True)

        logger.info(
            f"[MOTION_DRAFTER] Proper legal motion completed: {motion_draft.total_page_estimate} pages"
        )

        return motion_draft

    def _convert_to_proper_legal_structure(
        self, outline: Dict[str, Any], target_length: DocumentLength
    ) -> Dict[str, Any]:
        """Convert outline to proper legal motion structure"""

        min_pages, max_pages = target_length.value
        target_words = ((min_pages + max_pages) // 2) * self.words_per_page

        # Define standard structure with word allocation
        structure = {
            "introduction": OutlineSection(
                id="intro",
                title="Introduction",
                section_type=SectionType.INTRODUCTION,
                content_points=[],
                legal_authorities=[],
                target_length=int(target_words * 0.05),  # 5%
                hook_options=[],
                themes=[],
            ),
            "facts": OutlineSection(
                id="facts",
                title="Statement of Facts",
                section_type=SectionType.STATEMENT_OF_FACTS,
                content_points=[],
                legal_authorities=[],
                target_length=int(target_words * 0.20),  # 20%
                key_facts=[],
            ),
            "legal_standard": OutlineSection(
                id="legal",
                title="Legal Standard",
                section_type=SectionType.LEGAL_STANDARD,
                content_points=[],
                legal_authorities=[],
                target_length=int(target_words * 0.10),  # 10%
            ),
            "arguments": [],  # 60% total
            "conclusion": OutlineSection(
                id="conclusion",
                title="Conclusion",
                section_type=SectionType.CONCLUSION,
                content_points=[],
                legal_authorities=[],
                target_length=int(target_words * 0.05),  # 5%
            ),
        }

        # Extract content from outline sections
        outline_sections = outline.get("sections", [])
        hooks = []
        themes = []
        facts = []
        authorities = []
        arguments = []

        for section in outline_sections:
            # Parse content
            for item in section.get("content", []):
                if item.get("type") == "field":
                    label = item.get("label", "").lower()
                    value = item.get("value", "")

                    if "hook" in label:
                        hooks.append(self._extract_clean_value(value))
                    elif "theme" in label:
                        themes.extend(self._extract_themes_from_value(value))
                    elif "fact" in label:
                        facts.append(value)
                    elif "authorit" in label:
                        authorities.append(value)

            # Check if this is an argument section
            heading = section.get("heading", "").lower()
            if not any(
                skip in heading
                for skip in ["introduction", "facts", "conclusion", "standard"]
            ):
                # This is an argument
                arguments.append(
                    {
                        "title": section.get(
                            "heading", f"Argument {len(arguments) + 1}"
                        ),
                        "content": section,
                        "authorities": self._extract_section_authorities(section),
                    }
                )

        # Populate structure
        structure["introduction"].hook_options = hooks[:3]
        structure["introduction"].themes = themes[:5]
        structure["facts"].key_facts = facts
        structure["legal_standard"].legal_authorities = authorities[:10]

        # Create argument sections
        argument_words = int(target_words * 0.60)
        if arguments:
            words_per_argument = argument_words // len(arguments)
            for i, arg_data in enumerate(arguments):
                arg_section = OutlineSection(
                    id=f"arg_{i}",
                    title=arg_data["title"],
                    section_type=SectionType.ARGUMENT,
                    content_points=[],
                    legal_authorities=arg_data["authorities"],
                    target_length=words_per_argument,
                )
                # Store original content for reference
                arg_section.context_summary = str(arg_data["content"])
                structure["arguments"].append(arg_section)

        return structure

    async def _draft_introduction_section(
        self,
        intro_structure: OutlineSection,
        case_context: Dict[str, Any],
        opposing_motion_text: Optional[str],
    ) -> DraftedSection:
        """Draft proper introduction section"""

        prompt = f"""Draft the INTRODUCTION section for this legal motion.

    CRITICAL: This is the formal INTRODUCTION section, not a summary or outline.

    Target length: {intro_structure.target_length} words MINIMUM

    Available hooks:
    {chr(10).join(f"- {hook}" for hook in intro_structure.hook_options[:3])}

    Key themes:
    {chr(10).join(f"- {theme}" for theme in intro_structure.themes[:5])}

    Motion title: {self.document_context.get("motion_title", "MOTION")}
    Case caption: {self.document_context.get("case_caption", "")}

    Requirements:
    1. Start with "INTRODUCTION" as the section header
    2. First paragraph: Use one of the hooks to grab attention
    3. Second paragraph: State what relief is sought
    4. Third paragraph: Preview main arguments (briefly)
    5. Optional fourth paragraph: Frame response to opposing motion if applicable
    6. Use formal legal language throughout
    7. Do NOT include bullet points or lists
    8. Write in flowing paragraphs

    Example structure:
    "INTRODUCTION

    [Hook paragraph that frames the legal issue compellingly]

    [Relief paragraph stating what plaintiff seeks from the court]

    [Preview paragraph briefly mentioning the key arguments to follow]"

    Remember: This is the actual INTRODUCTION section of the motion, not an outline."""

        if opposing_motion_text:
            prompt += f"\n\nResponding to opposition:\n{opposing_motion_text[:500]}..."

        result = await self.section_writer.run(prompt)
        content = str(result.data) if hasattr(result, "data") else str(result)

        # Ensure it starts with proper header
        if not content.strip().startswith("INTRODUCTION"):
            content = "INTRODUCTION\n\n" + content

        return DraftedSection(
            outline_section=intro_structure,
            content=content,
            word_count=len(content.split()),
            citations_used=self._extract_citations_from_text(content),
            citations_verified={},
            expansion_cycles=1,
            confidence_score=0.9,
        )

    async def _draft_argument_section_proper(
        self,
        arg_structure: OutlineSection,
        case_context: Dict[str, Any],
        arg_number: int,
        total_args: int,
    ) -> DraftedSection:
        """Draft a proper argument section"""

        # Format section header
        roman = self._to_roman_numeral(arg_number)
        section_header = f"{roman}. {arg_structure.title.upper()}"

        prompt = f"""Draft legal argument section: {section_header}

    Target length: {arg_structure.target_length} words MINIMUM

    This is argument {arg_number} of {total_args} in the ARGUMENT portion of the motion.

    Legal authorities for this argument:
    {chr(10).join(f"- {auth}" for auth in arg_structure.legal_authorities[:7])}

    Requirements:
    1. Start with the section header: {section_header}
    2. First paragraph: Clear thesis statement of what this argument proves
    3. Follow IRAC structure:
    - Issue: State the specific legal question
    - Rule: Cite the applicable law with parentheticals
    - Application: Apply law to facts in detail (multiple paragraphs)
    - Conclusion: Explain why this supports the relief sought
    4. Use "Moreover," "Furthermore," "Additionally" to build argument
    5. Include "Even if" or "Even assuming" for counterarguments
    6. End with conclusion tying back to requested relief
    7. Minimum 5-7 substantial paragraphs
    8. All citations must include parentheticals

    Example structure:
    "{section_header}

    [Thesis paragraph stating the legal conclusion]

    [Rule paragraph(s) with legal standard and citations]

    [Multiple application paragraphs applying law to facts]

    [Additional analysis with Moreover/Furthermore]

    [Counter-argument paragraph with Even if]

    [Conclusion paragraph tying to relief]"

    Remember: This is a complete legal argument that must be thorough and persuasive."""

        result = await self.section_writer.run(prompt)
        content = str(result.data) if hasattr(result, "data") else str(result)

        # Ensure proper header
        if not content.strip().startswith(roman):
            content = f"{section_header}\n\n{content}"

        # Expand if needed
        word_count = len(content.split())
        if word_count < arg_structure.target_length * 0.9:
            content = await self._expand_argument_section(
                content, arg_structure, case_context
            )

        return DraftedSection(
            outline_section=arg_structure,
            content=content,
            word_count=len(content.split()),
            citations_used=self._extract_citations_from_text(content),
            citations_verified={},
            expansion_cycles=1,
            confidence_score=0.85,
        )

    def _to_roman_numeral(self, num: int) -> str:
        """Convert number to Roman numeral"""
        values = [(10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")]
        result = ""
        for value, letter in values:
            count, num = divmod(num, value)
            result += letter * count
        return result

    # Add this to export_to_docx method to format properly
    def export_proper_motion_to_docx(self, motion_draft: MotionDraft, output_path: str):
        """Export motion with proper legal formatting"""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE

        doc = Document()

        # Set up styles
        styles = doc.styles

        # Create heading style for section headers
        heading_style = styles.add_style("SectionHeading", WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = "Times New Roman"
        heading_style.font.size = Pt(12)
        heading_style.font.bold = True
        heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_style.paragraph_format.space_after = Pt(12)

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1)

        # Add caption
        caption_lines = motion_draft.case_name.split("\n")
        for line in caption_lines:
            p = doc.add_paragraph(line)
            if "CIRCUIT COURT" in line.upper():
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Space after caption

        # Add title
        title_p = doc.add_paragraph(motion_draft.title)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.style.font.bold = True
        title_p.style.font.size = Pt(14)

        doc.add_paragraph()  # Space after title

        # Add each section
        for section in motion_draft.sections:
            # Parse content to find section headers
            lines = section.content.split("\n")

            for i, line in enumerate(lines):
                line = line.strip()

                if not line:
                    continue

                # Check if this is a section header
                is_header = False
                if line.upper() == line and len(line) < 50:  # All caps, short
                    is_header = True
                elif line.startswith(
                    ("I.", "II.", "III.", "IV.", "V.")
                ):  # Roman numerals
                    is_header = True

                if is_header:
                    p = doc.add_paragraph(line, style="SectionHeading")
                else:
                    p = doc.add_paragraph(line)
                    p.style.font.name = "Times New Roman"
                    p.style.font.size = Pt(12)
                    p.paragraph_format.line_spacing = 2  # Double spaced
                    p.paragraph_format.first_line_indent = Inches(0.5)

        # Add signature block
        doc.add_page_break()
        doc.add_paragraph("Respectfully submitted,")
        doc.add_paragraph()
        doc.add_paragraph("_______________________________")
        doc.add_paragraph("[ATTORNEY NAME]")
        doc.add_paragraph("Attorney for Plaintiff")
        doc.add_paragraph("[FIRM NAME]")
        doc.add_paragraph("[ADDRESS]")
        doc.add_paragraph("[PHONE]")
        doc.add_paragraph("[EMAIL]")
        doc.add_paragraph("Florida Bar No. [NUMBER]")

        # Certificate of Service
        doc.add_paragraph()
        cert = doc.add_paragraph("CERTIFICATE OF SERVICE")
        cert.style.font.bold = True
        cert.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(
            "I HEREBY CERTIFY that a true and correct copy of the foregoing was "
            "served via [METHOD] on [DATE] to all counsel of record."
        )

        doc.save(output_path)
        logger.info(f"Proper motion exported to {output_path}")


# Create enhanced singleton instance
motion_drafter = EnhancedMotionDraftingAgent()
