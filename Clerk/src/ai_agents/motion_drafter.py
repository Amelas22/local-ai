"""
AI Motion Drafting Agent - Enhanced Version with Database Fix and Firm Knowledge Integration
Implements section-by-section legal motion generation with improved database access
"""

import asyncio
import logging
from typing import List, Dict, Any, Optional
from datetime import datetime
from dataclasses import dataclass, field
from enum import Enum
import re

from pydantic_ai import Agent
from pydantic_ai.models.openai import OpenAIModel
from openai import AsyncOpenAI
import tiktoken

from src.vector_storage.qdrant_store import QdrantVectorStore
from src.vector_storage.embeddings import EmbeddingGenerator
from src.ai_agents.outline_cache_manager import outline_cache
from src.ai_agents.rag_research_agent import rag_research_agent, ResearchRequest
from src.ai_agents.legal_formatter import legal_formatter
from src.utils.timeout_monitor import TimeoutMonitor
from config.settings import settings

# Use the same logger as the main API to ensure messages are visible
logger = logging.getLogger("clerk_api")


class SectionType(Enum):
    """Types of sections in a legal motion"""

    PREINTRODUCTION = "preintroduction"  # Case caption and title
    INTRODUCTION = "introduction"  # Brief 1-2 paragraph introduction
    PROCEDURAL_BACKGROUND = "procedural_background"  # Procedural history
    STATEMENT_OF_FACTS = "statement_of_facts"  # Factual background (skip for now)
    MEMORANDUM_OF_LAW = "memorandum_of_law"  # Main legal analysis section
    LEGAL_STANDARD = "legal_standard"  # Legal standards subsection
    ARGUMENT = "argument"  # Main arguments (part of memorandum)
    SUB_ARGUMENT = "sub_argument"  # Sub-arguments within main arguments
    CONCLUSION = "conclusion"  # Brief 3-4 sentence conclusion
    PRAYER_FOR_RELIEF = "prayer_for_relief"  # WHEREFORE clause and signature block


class DocumentLength(Enum):
    """Target document lengths with expanded ranges"""

    SHORT = (15, 20)  # 15-20 pages
    MEDIUM = (20, 30)  # 20-30 pages
    LONG = (30, 40)  # 30-40 pages
    COMPREHENSIVE = (35, 50)  # 35-50 pages for complex motions


@dataclass
class OutlineSection:
    """Enhanced outline section with better structure"""

    id: str
    title: str
    section_type: SectionType
    content_points: List[str]
    legal_authorities: List[str]
    target_length: int  # Target word count
    parent_id: Optional[str] = None
    children: List["OutlineSection"] = field(default_factory=list)
    context_summary: Optional[str] = None
    hook_options: List[str] = field(default_factory=list)
    themes: List[str] = field(default_factory=list)
    key_facts: List[Dict[str, Any]] = field(default_factory=list)
    counter_arguments: List[Dict[str, str]] = field(default_factory=list)


@dataclass
class DraftedSection:
    """Enhanced drafted section with quality metrics"""

    outline_section: OutlineSection
    content: str
    word_count: int
    citations_used: List[str]
    citations_verified: Dict[str, bool]  # Citation -> verification status
    expansion_cycles: int
    confidence_score: float
    needs_revision: bool = False
    revision_notes: List[str] = field(default_factory=list)
    consistency_score: float = 0.0
    factual_accuracy_score: float = 0.0
    argument_strength_score: float = 0.0
    transitions: Dict[str, str] = field(default_factory=dict)  # To next/from previous


@dataclass
class MotionDraft:
    """Enhanced motion draft with comprehensive metadata"""

    title: str
    case_name: str  # Can be derived from database_name or set as a display value
    sections: List[DraftedSection]
    total_word_count: int
    total_page_estimate: int
    creation_timestamp: datetime
    outline_source: Dict[str, Any]
    coherence_score: float = 0.0
    review_notes: List[str] = field(default_factory=list)
    quality_metrics: Dict[str, float] = field(default_factory=dict)
    citation_index: Dict[str, List[str]] = field(
        default_factory=dict
    )  # Citation -> sections using it
    consistency_issues: List[Dict[str, Any]] = field(default_factory=list)


class EnhancedMotionDraftingAgent:
    """Enhanced AI agent for drafting legal motions with fixed database access and firm knowledge"""

    def __init__(self):
        """Initialize the enhanced motion drafting agent"""
        self.vector_store = QdrantVectorStore()
        self.embedding_generator = EmbeddingGenerator()
        self.openai_client = AsyncOpenAI(api_key=settings.openai.api_key)

        # Initialize OpenAI models with proper configuration
        self.gpt4_model = OpenAIModel("gpt-4.1-mini-2025-04-14")

        # For now, use GPT-4 for all agents to ensure compatibility
        self.primary_model = self.gpt4_model

        # Initialize AI agents with enhanced prompts
        self.section_writer = self._create_enhanced_section_writer()

        # Initialize tokenizer
        self.tokenizer = tiktoken.encoding_for_model("gpt-4")

        # Enhanced configuration
        self.words_per_page = 250
        self.max_expansion_cycles = 5
        self.min_confidence_threshold = 0.75
        self.citation_patterns = self._compile_citation_patterns()

        # Track document-wide context
        self.document_context = {
            "themes": [],
            "key_arguments": [],
            "terminology": {},
            "citations_used": set(),
            "fact_chronology": [],
        }

        logger.info("EnhancedMotionDraftingAgent initialized successfully")

    def _create_enhanced_section_writer(self) -> Agent:
        """Create enhanced section writer with evidence-focused prompting"""
        try:
            # Use primary model for section writing
            agent = Agent(
                self.primary_model,
                system_prompt="""You are an expert legal writer specializing in professional legal motion drafting for Florida courts.

## CRITICAL EVIDENCE INTEGRATION REQUIREMENTS

### EVERY PARAGRAPH MUST INCLUDE SPECIFIC EVIDENCE
- Each factual assertion MUST be supported by a specific citation
- Use exact page:line references for depositions (e.g., "Smith Dep. 45:12-23")
- Include exhibit numbers (e.g., "Ex. 15", "Maintenance Log Nos. 52-63, Ex. 22")
- Reference specific documents with dates (e.g., "Email from Jones dated 3/15/2023")
- Quote or paraphrase actual evidence content

### CITATION FORMAT EXAMPLES
- Deposition: "MR. DACOSTA testified that he was instructed to disable the telematics system. (DACOSTA Dep. 34:12-25)"
- Exhibit: "The maintenance logs show brake failures were reported but not repaired. (Maintenance Log Nos. 52-63, Ex. 22)"
- Expert Report: "Dr. Reynolds concluded the crash was preventable with proper maintenance. (Expert Report of Dr. Reynolds at 15, Ex. 19)"
- Internal Document: "PFG's internal emails confirm management knew of the safety risks. (Email from Fleet Manager dated 8/15/2021, Ex. 17)"
- Video Evidence: "The dashcam video shows DACOSTA stopped in the travel lane. (Dashcam Recording at 3:52:15, Ex. 5)"

## SECTION-SPECIFIC REQUIREMENTS

### INTRODUCTION SECTIONS
- 1-2 paragraphs maximum
- State the motion type and relief sought
- Brief preview of key arguments
- End with "for the reasons set forth below"

### PROCEDURAL BACKGROUND
- Numbered paragraphs for each filing
- Include docket entry numbers (DE #)
- Chronological order

### MEMORANDUM OF LAW / ARGUMENTS (80-90% of motion)
STRUCTURE FOR EACH ARGUMENT:
1. **Heading**: Clear statement of legal position
2. **Rule Paragraph**: State the applicable law with case citations
3. **Application Paragraphs**: 
   - MUST include 3-5 specific pieces of evidence per argument
   - Connect each piece of evidence to the legal standard
   - Use transitions like "Moreover," "Additionally," "Furthermore"
4. **Conclusion**: Brief summary tying evidence to relief sought

EVIDENCE REQUIREMENTS PER ARGUMENT:
- Minimum 3 citations to case-specific evidence
- At least 1 deposition excerpt
- At least 1 exhibit reference
- At least 1 internal document or expert opinion

### CONCLUSION SECTIONS
- 3-4 sentences maximum
- Summarize why evidence compels relief
- Reference strongest evidence briefly
- Lead to WHEREFORE paragraph

### WHEREFORE/PRAYER FOR RELIEF
- Standard format with specific relief
- Professional signature block

## EVIDENCE INTEGRATION TECHNIQUES
1. **Lead with Evidence**: "The video evidence establishes that..."
2. **Parenthetical Support**: "PFG knew of the danger (Email from Safety Director dated 1/5/2023)"
3. **Block Quotes**: Use for critical deposition testimony
4. **Sequential Evidence**: "First... Second... Third..." with different evidence types
5. **Corroboration**: "This is confirmed by..." linking multiple pieces

## QUALITY CHECKS
Before completing any section, verify:
- Every factual claim has a specific citation
- Citations include precise references (page:line, exhibit #, dates)
- Evidence directly supports the legal argument
- Multiple evidence types are used
- No generic or unsupported assertions

Remember: Judges want SPECIFIC EVIDENCE, not general allegations.""",
                result_type=str,
            )
            logger.info("Section writer agent created successfully")
            return agent
        except Exception as e:
            logger.error(f"Failed to create section writer agent: {str(e)}")
            raise

    def _compile_citation_patterns(self) -> List[re.Pattern]:
        """Compile regex patterns for citation extraction"""
        patterns = [
            # Federal cases
            re.compile(r"\d+\s+U\.S\.\s+\d+(?:\s+\(\d{4}\))?"),
            re.compile(r"\d+\s+F\.\d+d\s+\d+(?:\s+\([^)]+\d{4}\))?"),
            re.compile(r"\d+\s+F\.\s*Supp\.\s*\d+d?\s+\d+(?:\s+\([^)]+\d{4}\))?"),
            re.compile(r"\d+\s+S\.\s*Ct\.\s+\d+(?:\s+\(\d{4}\))?"),
            # State cases
            re.compile(r"\d+\s+[A-Z][a-z]+\.(?:\s+\d+d)?\s+\d+(?:\s+\([^)]+\d{4}\))?"),
            # Case names
            re.compile(r"[A-Z][a-zA-Z]+\s+v\.\s+[A-Z][a-zA-Z]+(?:,\s+\d+)?"),
            # Statutes
            re.compile(r"\d+\s+U\.S\.C\.\s+§+\s*\d+[\w\-\.]*"),
            re.compile(r"[A-Z][a-z]+\.\s+Stat\.\s+(?:Ann\.\s+)?§+\s*\d+[\w\-\.]*"),
            # Regulations
            re.compile(r"\d+\s+C\.F\.R\.\s+§*\s*\d+\.\d+"),
            # Rules
            re.compile(r"Fed\.\s*R\.\s*(?:Civ|Crim|Evid|App)\.\s*P\.\s*\d+[a-z]*"),
        ]
        return patterns

    def _determine_section_type(self, title: str) -> SectionType:
        """Determine section type from title with proper legal motion structure"""
        title_lower = title.lower()

        # Check for pre-introduction/caption
        if any(
            term in title_lower for term in ["caption", "title", "court", "case no"]
        ):
            return SectionType.PREINTRODUCTION

        # Check for procedural background
        elif any(
            term in title_lower for term in ["procedural", "history", "proceedings"]
        ):
            return SectionType.PROCEDURAL_BACKGROUND

        # Check for memorandum of law (main section)
        elif any(term in title_lower for term in ["memorandum", "memo", "law"]):
            return SectionType.MEMORANDUM_OF_LAW

        # Check for introduction (keep brief)
        elif "introduction" in title_lower and "argument" not in title_lower:
            return SectionType.INTRODUCTION

        # Check for facts/statement of facts
        elif (
            any(term in title_lower for term in ["facts", "factual", "background"])
            and "procedural" not in title_lower
        ):
            return SectionType.STATEMENT_OF_FACTS

        # Check for legal standard
        elif any(
            term in title_lower
            for term in ["standard", "legal framework", "applicable law"]
        ):
            return SectionType.LEGAL_STANDARD

        # Check for conclusion (keep very brief)
        elif (
            any(term in title_lower for term in ["conclusion", "summary"])
            and "argument" not in title_lower
        ):
            return SectionType.CONCLUSION

        # Check for prayer/relief (WHEREFORE clause)
        elif any(
            term in title_lower for term in ["prayer", "relief", "wherefore", "request"]
        ):
            return SectionType.PRAYER_FOR_RELIEF

        # Default to argument (part of memorandum of law)
        else:
            return SectionType.ARGUMENT

    async def _extract_argument_sections(
        self, outline_id: str, outline_structure: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """Extract argument sections from outline content"""

        new_structure = []

        for section_info in outline_structure:
            # Get the full section data
            section_data = await outline_cache.get_section(
                outline_id, section_info["index"]
            )
            if not section_data:
                new_structure.append(section_info)
                continue

            # Check if this section contains argument paragraphs
            content_items = section_data.get("content", [])
            added_arguments = False

            for item in content_items:
                if item.get("type") == "paragraph":
                    text = item.get("text", "")
                    # Check if this is an argument heading (e.g., "ARGUMENT I – ...")
                    if text.startswith("ARGUMENT") and " – " in text:
                        # Create a new section for this argument
                        arg_section = {
                            "index": len(new_structure),
                            "heading": text,
                            "type": "argument",
                            "content_items": 0,  # Will be populated from following items
                            "has_subsections": False,
                            "_original_index": section_info[
                                "index"
                            ],  # Keep reference to original
                        }
                        new_structure.append(arg_section)
                        added_arguments = True
                        logger.info(f"[EXTRACT_ARGS] Found argument section: {text}")

            # If we didn't extract any arguments, keep the original section
            # But only if it's not a facts section that contained arguments
            if not added_arguments or section_info["type"] not in [
                "facts",
                "statement_of_facts",
            ]:
                new_structure.append(section_info)

        logger.info(
            f"[EXTRACT_ARGS] Extracted structure: {len(outline_structure)} -> {len(new_structure)} sections"
        )
        return new_structure

    async def _extract_argument_content(
        self, original_section_data: Dict[str, Any], argument_heading: str
    ) -> Dict[str, Any]:
        """Extract content for a specific argument from the original section data"""

        content_items = original_section_data.get("content", [])
        argument_content = []
        found_argument = False

        for i, item in enumerate(content_items):
            # Check if this is the argument we're looking for
            if item.get("type") == "paragraph" and item.get("text") == argument_heading:
                found_argument = True
                # Look for the following content items that belong to this argument
                for j in range(i + 1, len(content_items)):
                    next_item = content_items[j]
                    # Stop if we hit another argument
                    if (
                        next_item.get("type") == "paragraph"
                        and next_item.get("text", "").startswith("ARGUMENT")
                        and " – " in next_item.get("text", "")
                    ):
                        break
                    # Collect content for this argument
                    if next_item.get("label") in [
                        "Summary",
                        "Structure",
                        "Key Authorities",
                        "Fact Integration",
                        "Counter-Argument Response",
                    ]:
                        argument_content.append(next_item)
                break

        if not found_argument:
            logger.warning(f"Could not find argument content for: {argument_heading}")
            return {"content": []}

        return {"content": argument_content}

    def _deduplicate_sections(
        self, outline_structure: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """Deduplicate sections by type to prevent multiple conclusions"""
        seen_types = {}
        deduplicated = []

        for section in outline_structure:
            section_type = section.get("type", "standard")
            heading = section.get("heading", "").lower()

            # Determine actual section type
            if "conclusion" in heading:
                section_type = "conclusion"
            elif "introduction" in heading:
                section_type = "introduction"
            elif "facts" in heading:
                section_type = "facts"
            elif "prayer" in heading or "relief" in heading:
                section_type = "prayer"

            # For conclusion and prayer sections, only keep the first one
            if section_type in ["conclusion", "prayer", "introduction"]:
                if section_type not in seen_types:
                    seen_types[section_type] = True
                    deduplicated.append(section)
                else:
                    logger.warning(
                        f"Skipping duplicate {section_type} section: {section.get('heading')}"
                    )
            else:
                # Keep all argument and fact sections
                deduplicated.append(section)

        logger.info(
            f"Deduplicated sections: {len(outline_structure)} -> {len(deduplicated)}"
        )
        return deduplicated

    def _clean_outline_content(self, content: str, max_length: int = 500) -> str:
        """Clean and truncate outline content to prevent token overflow"""
        if not content:
            return ""

        # Handle multi-option format (e.g., "Option 1: ... || Option 2: ...")
        if "||" in content:
            options = content.split("||")
            # Take only the first option and clean it
            content = options[0].strip()
            # Remove "Option 1:" prefix if present
            if content.startswith("Option"):
                content = (
                    content.split(":", 1)[1].strip() if ":" in content else content
                )

        # Truncate if too long
        if len(content) > max_length:
            content = content[:max_length] + "..."

        return content

    # NOTE: Old query building methods removed - now using RAG research agent

    async def _retrieve_enhanced_case_context(
        self, database_name: str, outline_sections: List[OutlineSection]
    ) -> Dict[str, Any]:
        """Retrieve comprehensive case context using RAG research agent"""

        logger.info(
            f"[CASE_CONTEXT] Starting RAG agent-based context retrieval for {database_name}"
        )
        logger.info(
            f"[CASE_CONTEXT] Number of outline sections: {len(outline_sections)}"
        )

        try:
            # Step 1: Generate questions from outline sections
            questions = self._generate_research_questions_from_outline(outline_sections)
            logger.info(
                f"[CASE_CONTEXT] Generated {len(questions)} research questions from outline"
            )

            # Step 2: Create research request
            research_request = ResearchRequest(
                questions=questions,
                database_name=database_name,
                research_context=f"Motion drafting research for case database '{database_name}' with {len(outline_sections)} sections",
                section_type="mixed",  # Since we have multiple section types
            )

            # Step 3: Execute RAG research with timeout
            logger.info(
                f"[CASE_CONTEXT] Executing RAG research with {len(questions)} questions"
            )
            research_response = await asyncio.wait_for(
                rag_research_agent.research_questions(research_request),
                timeout=90.0,  # 90 second timeout for RAG research
            )

            # Step 4: Convert research response to legacy context format
            context = self._convert_research_response_to_context(research_response)

            # Step 5: Add additional processing
            context["opposing_motion"] = self.document_context.get(
                "opposing_motion_text", ""
            )
            context["themes"] = self._extract_themes_from_research(research_response)

            logger.info("[CASE_CONTEXT] RAG research completed successfully")
            logger.info(
                f"[CASE_CONTEXT] Results: {len(context['legal_authorities'])} legal authorities, "
                f"{len(context['case_facts'])} case facts, {len(context['expert_reports'])} expert reports, "
                f"{len(context['firm_knowledge'])} firm examples"
            )

            return context

        except asyncio.TimeoutError:
            logger.error(
                "[CASE_CONTEXT] RAG research timed out - falling back to minimal context"
            )
            return self._get_minimal_case_context()
        except Exception as e:
            logger.error(
                f"[CASE_CONTEXT] Error in RAG research: {str(e)}", exc_info=True
            )
            return self._get_minimal_case_context()

    def _generate_research_questions_from_outline(
        self, outline_sections: List[OutlineSection]
    ) -> List[str]:
        """Generate research questions from outline sections for RAG agent"""
        questions = []

        try:
            for section in outline_sections:
                # Process content points to generate questions
                for content_point in section.content_points[
                    :5
                ]:  # Limit to 5 per section
                    if not content_point or len(content_point) < 10:
                        continue

                    content_lower = content_point.lower()

                    # Generate fact-based questions
                    if any(
                        keyword in content_lower
                        for keyword in ["evidence", "fact", "document", "record"]
                    ):
                        questions.append(
                            f"What evidence or documentation exists about {content_point[:80]}?"
                        )

                    # Generate legal precedent questions
                    if any(
                        keyword in content_lower
                        for keyword in ["court", "case", "precedent", "law", "legal"]
                    ):
                        questions.append(
                            f"What legal precedents or authorities support arguments about {content_point[:80]}?"
                        )

                    # Generate strategy questions for firm knowledge
                    if any(
                        keyword in content_lower
                        for keyword in ["argument", "claim", "negligent", "liability"]
                    ):
                        questions.append(
                            f"How do successful motions argue {content_point[:80]}?"
                        )

                # Process legal authorities
                for authority in section.legal_authorities[
                    :3
                ]:  # Limit to 3 per section
                    if authority and len(authority) > 5:
                        questions.append(
                            f"How has {authority} been successfully used in similar legal arguments?"
                        )
                        questions.append(
                            f"What evidence supports the application of {authority} in this case?"
                        )

                # Process key facts
                for fact_item in section.key_facts[:3]:  # Limit to 3 per section
                    if isinstance(fact_item, dict):
                        fact_text = fact_item.get("description", "") or fact_item.get(
                            "text", ""
                        )
                    elif isinstance(fact_item, str):
                        fact_text = fact_item
                    else:
                        continue

                    if fact_text and len(fact_text) > 10:
                        questions.append(
                            f"What documentation supports the fact that {fact_text[:80]}?"
                        )

            # Add some general strategic questions
            questions.extend(
                [
                    "What are the strongest pieces of evidence for negligence claims?",
                    "How do effective motions structure arguments about corporate liability?",
                    "What regulatory violations or compliance issues exist in this case?",
                    "What expert testimony or opinions support the legal claims?",
                ]
            )

            # Remove duplicates and limit total
            unique_questions = list(dict.fromkeys(questions))  # Remove duplicates
            limited_questions = unique_questions[:25]  # Limit to 25 questions max

            logger.info(
                f"[RESEARCH_QUESTIONS] Generated {len(limited_questions)} unique questions from {len(outline_sections)} sections"
            )

            return limited_questions

        except Exception as e:
            logger.error(f"Error generating research questions: {str(e)}")
            # Return fallback questions
            return [
                "What are the main facts and evidence in this case?",
                "What legal precedents support the claims being made?",
                "How do successful motions argue similar legal theories?",
                "What expert testimony or analysis exists?",
                "What regulatory violations or compliance issues are present?",
            ]

    def _convert_research_response_to_context(
        self, research_response
    ) -> Dict[str, Any]:
        """Convert RAG research response to legacy context format"""
        context = {
            "case_facts": [],
            "legal_authorities": [],
            "expert_reports": [],
            "regulatory_evidence": [],
            "procedural_documents": [],
            "fact_chronology": [],
            "themes": [],
            "firm_knowledge": [],
        }

        try:
            # Convert case facts
            for fact in research_response.case_facts:
                context["case_facts"].append(
                    type(
                        "obj",
                        (object,),
                        {
                            "content": fact["content"],
                            "metadata": fact["metadata"],
                            "score": fact["score"],
                        },
                    )()
                )

            # Convert legal authorities (keep as dicts for compatibility)
            context["legal_authorities"] = research_response.legal_precedents

            # Convert expert reports
            for expert in research_response.expert_evidence:
                context["expert_reports"].append(
                    type(
                        "obj",
                        (object,),
                        {
                            "content": expert["content"],
                            "metadata": expert["metadata"],
                            "score": expert["score"],
                        },
                    )()
                )

            # Convert regulatory evidence
            for reg in research_response.regulatory_compliance:
                context["regulatory_evidence"].append(
                    type(
                        "obj",
                        (object,),
                        {
                            "content": reg["content"],
                            "metadata": reg["metadata"],
                            "score": reg["score"],
                        },
                    )()
                )

            # Convert procedural documents
            for proc in research_response.procedural_history:
                context["procedural_documents"].append(
                    type(
                        "obj",
                        (object,),
                        {
                            "content": proc["content"],
                            "metadata": proc["metadata"],
                            "score": proc["score"],
                        },
                    )()
                )

            # Convert firm knowledge (keep as dicts)
            context["firm_knowledge"] = research_response.argument_strategies

            # Build fact chronology from case facts
            context["fact_chronology"] = self._build_fact_chronology_from_research(
                research_response.case_facts
            )

            logger.info(
                "[CONVERT_RESPONSE] Converted research response to legacy context format"
            )

        except Exception as e:
            logger.error(f"Error converting research response: {str(e)}")

        return context

    def _extract_themes_from_research(self, research_response) -> List[str]:
        """Extract themes from RAG research response"""
        themes = []

        try:
            # Analyze all content for common themes
            all_content = []

            for category in [
                research_response.case_facts,
                research_response.legal_precedents,
                research_response.expert_evidence,
                research_response.argument_strategies,
            ]:
                for item in category[:5]:  # Limit to prevent overflow
                    content = item.get("content", "") if isinstance(item, dict) else ""
                    if content:
                        all_content.append(content.lower())

            combined_text = " ".join(all_content)

            # Theme detection patterns
            theme_patterns = {
                "negligent_hiring": [
                    "hiring",
                    "background",
                    "qualification",
                    "employment",
                ],
                "negligent_supervision": [
                    "supervision",
                    "monitoring",
                    "oversight",
                    "management",
                ],
                "safety_violations": [
                    "safety",
                    "violation",
                    "regulation",
                    "compliance",
                ],
                "maintenance_issues": [
                    "maintenance",
                    "inspection",
                    "repair",
                    "equipment",
                ],
                "corporate_responsibility": [
                    "corporate",
                    "company",
                    "policy",
                    "procedure",
                ],
            }

            for theme, keywords in theme_patterns.items():
                if sum(1 for kw in keywords if kw in combined_text) >= 2:
                    themes.append(theme)

        except Exception as e:
            logger.error(f"Error extracting themes from research: {str(e)}")

        return themes[:5]  # Top 5 themes

    def _build_fact_chronology_from_research(
        self, case_facts: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """Build chronology from research case facts"""
        chronology = []

        try:
            for fact in case_facts[:20]:  # Limit processing
                content = fact.get("content", "")
                if not content:
                    continue

                # Extract dates
                date_patterns = [
                    r"(\d{1,2}/\d{1,2}/\d{2,4})",
                    r"(\w+\s+\d{1,2},\s+\d{4})",
                    r"(\d{4}-\d{2}-\d{2})",
                ]

                for pattern in date_patterns:
                    dates = re.findall(pattern, content)
                    for date in dates[:2]:  # Limit dates per fact
                        chronology.append(
                            {
                                "date": date,
                                "description": content[:100],
                                "source": fact.get("source", "Unknown"),
                            }
                        )

        except Exception as e:
            logger.debug(f"Error building chronology from research: {e}")

        return chronology[:15]  # Top 15 events

    # NOTE: Old firm knowledge search methods removed - now handled by RAG research agent

    def _build_fact_chronology(self, facts: List[Any]) -> List[Dict[str, Any]]:
        """Build chronological ordering of facts"""

        chronology = []

        for fact in facts[:50]:  # Limit processing
            try:
                content = fact.content if hasattr(fact, "content") else str(fact)

                # Extract dates
                date_patterns = [
                    r"(\d{1,2}/\d{1,2}/\d{2,4})",
                    r"(\w+\s+\d{1,2},\s+\d{4})",
                    r"(\d{4}-\d{2}-\d{2})",
                ]

                for pattern in date_patterns:
                    dates = re.findall(pattern, content)
                    for date in dates[:2]:  # Limit dates per fact
                        chronology.append(
                            {
                                "date": date,
                                "description": content[:100],
                                "source": fact.metadata.get("document_name", "Unknown")
                                if hasattr(fact, "metadata")
                                else "Unknown",
                            }
                        )
            except Exception as e:
                logger.debug(f"Error processing fact for chronology: {e}")
                continue

        # Sort by date (simple approach - would need proper date parsing in production)
        return chronology[:20]  # Top 20 events

    def _extract_case_themes(self, context: Dict[str, Any]) -> List[str]:
        """Extract overarching themes from case context"""

        themes = []

        try:
            # Analyze content for common themes
            all_content = []
            for category in ["case_facts", "expert_reports", "legal_authorities"]:
                for item in context.get(category, [])[:10]:
                    if hasattr(item, "content"):
                        all_content.append(item.content)
                    elif isinstance(item, dict):
                        all_content.append(item.get("content", ""))

            combined_text = " ".join(all_content).lower()

            # Theme detection patterns
            theme_patterns = {
                "safety": ["safety", "dangerous", "hazard", "risk", "harm"],
                "corporate_responsibility": [
                    "corporate",
                    "company",
                    "business",
                    "employer",
                ],
                "negligence": ["negligent", "breach", "duty", "reasonable care"],
                "causation": ["caused", "resulted", "led to", "because of"],
                "damages": ["damages", "injury", "harm", "loss", "suffering"],
            }

            for theme, keywords in theme_patterns.items():
                if sum(1 for kw in keywords if kw in combined_text) >= 2:
                    themes.append(theme)

        except Exception as e:
            logger.error(f"Error extracting themes: {e}")

        return themes[:5]  # Top 5 themes

    def _extract_themes_from_value(self, value: str) -> List[str]:
        """Extract themes from a field value string"""
        if not value:
            return []

        themes = []

        # Handle multi-option format
        if "||" in value:
            options = value.split("||")
            for option in options:
                option = option.strip()
                if option.startswith("Option"):
                    option = (
                        option.split(":", 1)[1].strip() if ":" in option else option
                    )
                themes.append(option.strip('"').strip())
        else:
            themes.append(value.strip('"').strip())

        # Clean and filter themes
        cleaned_themes = []
        for theme in themes:
            if theme and len(theme) > 3:  # Skip very short themes
                cleaned_themes.append(theme[:100])  # Limit length

        return cleaned_themes[:3]  # Max 3 themes

    async def draft_motion_with_cache(
        self,
        outline: Dict[str, Any],
        database_name: str,
        target_length: DocumentLength = DocumentLength.MEDIUM,
        motion_title: Optional[str] = None,
        opposing_motion_text: Optional[str] = None,
        outline_id: Optional[str] = None,  # If already cached
    ) -> MotionDraft:
        """
        Enhanced motion drafting with fixed database access and firm knowledge
        """
        start_time = datetime.utcnow()
        logger.info(
            f"[MOTION_DRAFTER] Starting cached motion draft for database: {database_name}"
        )

        # Cache the outline if not already cached
        if not outline_id:
            outline_id = await outline_cache.cache_outline(outline, database_name)
            logger.info(f"[MOTION_DRAFTER] Cached outline with ID: {outline_id}")
        else:
            # Verify outline exists in cache
            cached_outline = await outline_cache.get_outline(outline_id)
            if not cached_outline:
                # Re-cache if missing
                outline_id = await outline_cache.cache_outline(outline, database_name)
                logger.info(
                    f"[MOTION_DRAFTER] Re-cached missing outline with ID: {outline_id}"
                )

        # Get outline structure for processing
        outline_structure = await outline_cache.get_outline_structure(outline_id)
        if not outline_structure:
            raise ValueError(f"Failed to get outline structure for ID: {outline_id}")

        # Extract and restructure argument sections from content
        outline_structure = await self._extract_argument_sections(
            outline_id, outline_structure
        )

        # Deduplicate sections to prevent multiple conclusions
        outline_structure = self._deduplicate_sections(outline_structure)

        logger.info(
            f"[MOTION_DRAFTER] Processing {len(outline_structure)} sections after deduplication"
        )

        # Initialize timeout monitor
        timeout_monitor = TimeoutMonitor(
            operation_name=f"Motion Drafting ({database_name})",
            warning_threshold=60,
            critical_threshold=300,  # Increased to 5 minutes
        )
        timeout_monitor.log_progress("Cached motion drafting started")

        # Initialize document context with minimal data
        self.document_context = {
            "themes": outline.get("themes", outline.get("style_notes", [])),
            "key_arguments": [],
            "terminology": {},
            "citations_used": set(),
            "fact_chronology": [],
            "database_name": database_name,
            "opposing_motion_text": opposing_motion_text,
            "outline_id": outline_id,  # Store for reference
        }

        # Calculate word distribution based on structure
        min_pages, max_pages = target_length.value
        target_words = ((min_pages + max_pages) // 2) * self.words_per_page
        word_distribution = self._calculate_word_distribution_from_structure(
            outline_structure, target_words
        )

        # Retrieve case context with timeout protection
        case_context = {}
        try:
            logger.info(
                f"[MOTION_DRAFTER] Starting case context retrieval for database: {database_name}"
            )
            full_outline_sections = await self._convert_structure_to_sections(
                outline_id, outline_structure
            )
            logger.info(
                f"[MOTION_DRAFTER] Converted {len(full_outline_sections)} outline sections"
            )

            # Use longer timeout for context retrieval
            logger.info(
                "[MOTION_DRAFTER] Calling _retrieve_enhanced_case_context with timeout=60s"
            )
            case_context = await asyncio.wait_for(
                self._retrieve_enhanced_case_context(
                    database_name, full_outline_sections
                ),
                timeout=60.0,  # 60 second timeout for all context retrieval
            )
            logger.info("[MOTION_DRAFTER] Case context retrieved successfully")
        except asyncio.TimeoutError:
            logger.error(
                "[MOTION_DRAFTER] Timeout retrieving case context - attempting basic search"
            )
            case_context = self._get_minimal_case_context()
            # Try to do at least one basic search
            try:
                case_context = await self._perform_basic_case_search(database_name)
            except Exception as search_error:
                logger.error(
                    f"[MOTION_DRAFTER] Basic search also failed: {search_error}"
                )
        except Exception as e:
            logger.error(
                f"[MOTION_DRAFTER] Error retrieving case context: {str(e)}",
                exc_info=True,
            )
            case_context = self._get_minimal_case_context()
            # Try to do at least one basic search
            try:
                case_context = await self._perform_basic_case_search(database_name)
            except Exception as search_error:
                logger.error(
                    f"[MOTION_DRAFTER] Basic search also failed: {search_error}"
                )

        # Process sections individually
        drafted_sections = []
        total_words = 0

        # Log case context summary
        logger.info(
            f"[MOTION_DRAFTER] Case context summary - "
            f"Facts: {len(case_context.get('case_facts', []))}, "
            f"Legal Authorities: {len(case_context.get('legal_authorities', []))}, "
            f"Expert Reports: {len(case_context.get('expert_reports', []))}, "
            f"Regulatory Evidence: {len(case_context.get('regulatory_evidence', []))}, "
            f"Firm Knowledge: {len(case_context.get('firm_knowledge', []))}, "
            f"Search Status: {case_context.get('search_status', 'completed')}"
        )

        for i, section_info in enumerate(outline_structure):
            section_start_time = datetime.utcnow()
            logger.info(
                f"[MOTION_DRAFTER] Processing section {i + 1}/{len(outline_structure)}: {section_info['heading']}"
            )

            # Update timeout monitor
            timeout_monitor.log_progress(
                f"Drafting section {i + 1}: {section_info['heading']}"
            )

            try:
                # Retrieve section from cache - handle extracted argument sections
                if "_original_index" in section_info:
                    # This is an extracted argument - get from original section
                    original_data = await outline_cache.get_section(
                        outline_id, section_info["_original_index"]
                    )
                    if not original_data:
                        logger.error(
                            f"Failed to retrieve original section for argument {i}"
                        )
                        continue
                    # Extract the relevant argument content
                    section_data = await self._extract_argument_content(
                        original_data, section_info["heading"]
                    )
                else:
                    # Regular section
                    section_data = await outline_cache.get_section(
                        outline_id, section_info["index"]
                    )
                    if not section_data:
                        logger.error(f"Failed to retrieve section {i} from cache")
                        continue

                # Create minimal OutlineSection for this section only
                outline_section = self._create_minimal_outline_section(
                    section_data,
                    section_info,
                    word_distribution.get(f"section_{i}", 500),
                )

                # Build cumulative context (lightweight)
                cumulative_context = self._build_lightweight_cumulative_context(
                    drafted_sections, case_context
                )

                # Draft the section with longer timeout
                drafted_section = await asyncio.wait_for(
                    self._draft_section_efficiently(
                        outline_section,
                        section_data,  # Pass full section data separately
                        case_context,
                        cumulative_context,
                    ),
                    timeout=60.0,  # 60 seconds per section
                )

                # Expand if needed
                if drafted_section.word_count < outline_section.target_length * 0.9:
                    drafted_section = await self._expand_section_efficiently(
                        drafted_section,
                        section_data,
                        outline_section.target_length,
                        case_context,
                    )

                drafted_sections.append(drafted_section)
                total_words += drafted_section.word_count

                # Update document context incrementally
                self._update_document_context(drafted_section)

                section_duration = datetime.utcnow() - section_start_time
                logger.info(
                    f"[MOTION_DRAFTER] Section {i + 1} completed in {section_duration}"
                )

            except asyncio.TimeoutError:
                logger.error(f"[MOTION_DRAFTER] Timeout drafting section {i + 1}")
                # Create placeholder section
                drafted_section = self._create_placeholder_section(
                    outline_section, "Section timed out"
                )
                drafted_sections.append(drafted_section)
            except Exception as e:
                logger.error(
                    f"[MOTION_DRAFTER] Error drafting section {i + 1}: {str(e)}",
                    exc_info=True,
                )
                # Create placeholder section
                outline_section = self._create_minimal_outline_section(
                    {"content": []}, section_info, 500
                )
                drafted_section = self._create_placeholder_section(
                    outline_section, str(e)
                )
                drafted_sections.append(drafted_section)

        # Post-process sections to ensure proper structure
        logger.info("[MOTION_DRAFTER] Post-processing sections for proper structure")
        drafted_sections = self._restructure_motion_sections(drafted_sections)

        # Recalculate total words after restructuring
        total_words = sum(section.word_count for section in drafted_sections)

        # Create motion draft
        motion_draft = MotionDraft(
            title=motion_title or self._extract_title_from_cache(outline),
            case_name=database_name.replace("_", " ").title(),
            sections=drafted_sections,
            total_word_count=total_words,
            total_page_estimate=total_words // self.words_per_page,
            creation_timestamp=datetime.utcnow(),
            outline_source={
                "outline_id": outline_id
            },  # Store reference instead of full outline
        )

        # Simplified review process
        logger.info("[MOTION_DRAFTER] Starting review process")
        motion_draft = await self._lightweight_review_process(motion_draft)

        # Build citation index
        motion_draft.citation_index = self._build_citation_index(motion_draft)

        timeout_monitor.finish(success=True)

        total_duration = datetime.utcnow() - start_time
        logger.info(f"[MOTION_DRAFTER] Completed in {total_duration}")

        return motion_draft

    def _get_minimal_case_context(self) -> Dict[str, Any]:
        """Return minimal case context when retrieval fails"""
        logger.warning(
            "[MOTION_DRAFTER] Using minimal case context - database searches were skipped"
        )
        return {
            "case_facts": [],
            "legal_authorities": [],
            "expert_reports": [],
            "key_documents": [],
            "medical_records": [],
            "motion_practice": [],
            "fact_chronology": [],
            "themes": ["negligence", "liability"],
            "opposing_motion": "",
            "firm_knowledge": [],
            "search_status": "skipped",  # Flag to indicate searches were skipped
        }

    async def _perform_basic_case_search(self, database_name: str) -> Dict[str, Any]:
        """Perform basic RAG search as fallback"""
        logger.info(f"[BASIC_SEARCH] Attempting basic RAG search on {database_name}")

        context = self._get_minimal_case_context()
        basic_rag_queries = [
            "What are the main facts and circumstances of this case?",
            "What evidence exists about the cause of the incident?",
            "What safety violations or negligence occurred?",
        ]

        try:
            for query in basic_rag_queries:
                logger.info(f"[BASIC_SEARCH] RAG query: {query}")
                (
                    query_embedding,
                    _,
                ) = await self.embedding_generator.generate_embedding_async(query)

                results = await self.vector_store.hybrid_search(
                    collection_name=database_name,
                    query=query,
                    query_embedding=query_embedding,
                    limit=3,
                    enable_reranking=False,
                )

                logger.info(f"[BASIC_SEARCH] Found {len(results)} results")
                for result in results:
                    logger.info(f"[BASIC_SEARCH] Result score: {result.score:.3f}")
                    if result.score > 0.2:  # Low threshold for basic search
                        context["case_facts"].append(result)
                        logger.info(
                            f"[BASIC_SEARCH] Added fact: {result.content[:50]}..."
                        )

            context["search_status"] = "basic_completed"
            logger.info(
                f"[BASIC_SEARCH] Collected {len(context['case_facts'])} facts from basic search"
            )
        except Exception as e:
            logger.error(f"[BASIC_SEARCH] Error: {str(e)}")
            context["search_status"] = "failed"

        return context

    def _create_placeholder_section(
        self, outline_section: OutlineSection, error_msg: str
    ) -> DraftedSection:
        """Create a placeholder section when drafting fails"""
        return DraftedSection(
            outline_section=outline_section,
            content=f"[Section could not be drafted due to error: {error_msg}]",
            word_count=0,
            citations_used=[],
            citations_verified={},
            expansion_cycles=0,
            confidence_score=0.0,
            needs_revision=True,
            revision_notes=[f"Section failed to draft: {error_msg}"],
        )

    async def _draft_section_efficiently(
        self,
        outline_section: OutlineSection,
        full_section_data: Dict[str, Any],
        case_context: Dict[str, Any],
        cumulative_context: Dict[str, Any],
    ) -> DraftedSection:
        """Draft a section efficiently with case facts and firm knowledge"""

        try:
            # Extract the most important content from full section data
            essential_content = self._extract_essential_content(full_section_data)

            # Build relevant case facts for this section
            relevant_facts = self._extract_relevant_facts(
                outline_section, case_context.get("case_facts", []), limit=5
            )

            # Find relevant firm knowledge examples
            relevant_examples = self._extract_relevant_firm_examples(
                outline_section, case_context.get("firm_knowledge", []), limit=2
            )

            # Create section-specific prompts based on type
            if outline_section.section_type == SectionType.INTRODUCTION:
                drafting_prompt = self._create_introduction_prompt(
                    outline_section, essential_content, cumulative_context
                )
            elif outline_section.section_type == SectionType.CONCLUSION:
                drafting_prompt = self._create_conclusion_prompt(
                    outline_section, essential_content, cumulative_context
                )
            elif outline_section.section_type == SectionType.PRAYER_FOR_RELIEF:
                drafting_prompt = self._create_prayer_prompt(
                    outline_section, essential_content, cumulative_context
                )
            elif outline_section.section_type in [
                SectionType.ARGUMENT,
                SectionType.SUB_ARGUMENT,
                SectionType.MEMORANDUM_OF_LAW,
                SectionType.LEGAL_STANDARD,
            ]:
                # Import citation formatter
                from src.ai_agents.citation_formatter import citation_formatter

                # Format evidence with proper citations
                formatted_evidence = []
                for fact in case_context.get("case_facts", [])[:10]:
                    try:
                        # Handle both object and dict formats
                        if hasattr(fact, "content"):
                            content = fact.content
                            metadata = (
                                fact.metadata if hasattr(fact, "metadata") else {}
                            )
                        else:
                            content = fact.get("content", "")
                            metadata = fact.get("metadata", {})

                        citation = citation_formatter.format_search_result_as_citation(
                            {"content": content, "metadata": metadata}
                        )
                        if citation:
                            formatted_evidence.append(
                                f"{content[:200]}... ({citation.formatted_citation})"
                            )
                        else:
                            formatted_evidence.append(f"{content[:200]}...")
                    except Exception as e:
                        logger.debug(f"Error formatting evidence: {e}")
                        continue

                # Full argument prompt with formatted evidence
                drafting_prompt = f"""Draft this ARGUMENT section of the legal motion:

Section: {outline_section.title}
Type: MEMORANDUM OF LAW - {outline_section.section_type.value}
Target Length: {outline_section.target_length} words (MINIMUM - this is a substantive section)

Key Arguments to Develop:
{chr(10).join(f"- {point}" for point in essential_content["key_points"][:7])}

Required Legal Authorities:
{chr(10).join(f"- {auth}" for auth in essential_content["authorities"][:7])}

SPECIFIC EVIDENCE TO INCORPORATE (USE THESE EXACT CITATIONS):
{chr(10).join(f"{i + 1}. {evidence}" for i, evidence in enumerate(formatted_evidence[:8]))}

Expert Testimony Available:
{self._format_expert_evidence(case_context.get("expert_reports", [])[:3])}

Regulatory Violations:
{self._format_regulatory_evidence(case_context.get("regulatory_evidence", [])[:3])}

Previous Arguments Made:
{cumulative_context.get("summary", "This is the first argument section.")}

CRITICAL REQUIREMENTS:
1. EVERY factual assertion MUST have a specific citation
2. Use the EXACT citations provided above
3. Include page:line references for depositions
4. Reference exhibit numbers
5. Quote key testimony or documents
6. Connect each piece of evidence to your legal argument
7. Use transitions between evidence presentations
8. Write AT LEAST {outline_section.target_length} words

EVIDENCE INTEGRATION EXAMPLE:
"The evidence establishes that PFG knew of the safety risks but chose profits over safety. MR. DACOSTA testified that his supervisor instructed him to disable the telematics monitoring system to avoid speed alerts. (DACOSTA Dep. 45:12-23). This is corroborated by an email from the Fleet Manager stating 'disable those annoying alerts - they're slowing down deliveries.' (Email from Jones dated 8/15/2021, Ex. 17). Moreover, the maintenance logs reveal that brake issues were reported three times but never repaired. (Maintenance Log Nos. 52-63, Ex. 22)."

Remember: Use SPECIFIC EVIDENCE with EXACT CITATIONS throughout."""
            else:
                # Default prompt for other sections
                drafting_prompt = f"""Draft this section of the legal motion:

Section: {outline_section.title}
Type: {outline_section.section_type.value}
Target Length: {outline_section.target_length} words

Key Points:
{chr(10).join(f"- {point}" for point in essential_content["key_points"][:5])}

Requirements:
1. Use appropriate legal formatting
2. Be concise and direct
3. Include relevant facts and authorities
4. Maintain professional tone"""

            # Generate content
            result = await asyncio.wait_for(
                self.section_writer.run(drafting_prompt),
                timeout=45,  # 45 second timeout per section
            )

            content = str(result.data) if hasattr(result, "data") else str(result)

            # Create drafted section
            word_count = len(content.split())
            return DraftedSection(
                outline_section=outline_section,
                content=content,
                word_count=word_count,
                citations_used=self._extract_citations_from_text(content),
                citations_verified={},
                expansion_cycles=1,
                confidence_score=self._calculate_section_confidence(
                    content, outline_section, word_count
                ),
            )

        except asyncio.TimeoutError:
            logger.error(f"Timeout drafting section: {outline_section.title}")
            raise
        except Exception as e:
            logger.error(f"Error drafting section efficiently: {str(e)}")
            raise

    def _extract_relevant_facts(
        self, outline_section: OutlineSection, case_facts: List[Any], limit: int = 5
    ) -> List[str]:
        """Extract facts relevant to the current section"""
        relevant_facts = []

        try:
            # Keywords to match based on section type
            if outline_section.section_type == SectionType.STATEMENT_OF_FACTS:
                # For facts section, get chronological facts
                for fact in case_facts[: limit * 2]:
                    if hasattr(fact, "content"):
                        relevant_facts.append(fact.content[:200])
            else:
                # For other sections, match based on content
                section_keywords = []
                for point in outline_section.content_points:
                    # Extract keywords from content points
                    words = point.lower().split()
                    section_keywords.extend([w for w in words if len(w) > 4][:3])

                # Find facts matching keywords
                for fact in case_facts:
                    if hasattr(fact, "content"):
                        fact_lower = fact.content.lower()
                        if any(kw in fact_lower for kw in section_keywords):
                            relevant_facts.append(fact.content[:200])
                            if len(relevant_facts) >= limit:
                                break

        except Exception as e:
            logger.debug(f"Error extracting relevant facts: {e}")

        return relevant_facts[:limit]

    def _extract_relevant_firm_examples(
        self,
        outline_section: OutlineSection,
        firm_knowledge: List[Dict[str, Any]],
        limit: int = 2,
    ) -> List[Dict[str, Any]]:
        """Extract firm knowledge examples relevant to current section"""
        relevant_examples = []

        try:
            # Match based on section type and authorities
            for example in firm_knowledge:
                # Check if motion type matches
                if outline_section.section_type == SectionType.ARGUMENT:
                    if "argument" in example.get("motion_type", "").lower():
                        relevant_examples.append(example)

                # Check if any authorities match
                for auth in outline_section.legal_authorities:
                    if auth.lower() in example.get("content", "").lower():
                        if example not in relevant_examples:
                            relevant_examples.append(example)

                if len(relevant_examples) >= limit:
                    break

        except Exception as e:
            logger.debug(f"Error extracting firm examples: {e}")

        return relevant_examples[:limit]

    def _format_expert_evidence(self, expert_reports: List[Any]) -> str:
        """Format expert evidence for inclusion in prompts"""
        formatted = []

        for expert in expert_reports:
            try:
                # Handle both object and dict formats
                if hasattr(expert, "content"):
                    content = expert.content
                    metadata = expert.metadata if hasattr(expert, "metadata") else {}
                else:
                    content = expert.get("content", "")
                    metadata = expert.get("metadata", {})

                # Try to extract expert name and format citation
                doc_name = metadata.get("document_name", "Expert")
                formatted_cite = f"Expert Report of {doc_name}"

                formatted.append(f"- {content[:200]}... ({formatted_cite})")
            except Exception as e:
                logger.debug(f"Error formatting expert evidence: {e}")
                continue

        return "\n".join(formatted) if formatted else "- No expert testimony available"

    def _format_regulatory_evidence(self, regulatory_evidence: List[Any]) -> str:
        """Format regulatory evidence for inclusion in prompts"""
        formatted = []

        for reg in regulatory_evidence:
            try:
                # Handle both object and dict formats
                if hasattr(reg, "content"):
                    content = reg.content
                    metadata = reg.metadata if hasattr(reg, "metadata") else {}
                else:
                    content = reg.get("content", "")
                    metadata = reg.get("metadata", {})

                # Try to extract regulation section
                section = metadata.get("section", "")
                if not section:
                    # Try to extract from content
                    import re

                    section_match = re.search(r"(\d+\.\d+)", content)
                    section = section_match.group(1) if section_match else "XXX"

                formatted.append(f"- {content[:150]}... (49 C.F.R. § {section})")
            except Exception as e:
                logger.debug(f"Error formatting regulatory evidence: {e}")
                continue

        return (
            "\n".join(formatted)
            if formatted
            else "- No regulatory violations identified"
        )

    def _create_introduction_prompt(
        self,
        outline_section: OutlineSection,
        essential_content: Dict[str, Any],
        cumulative_context: Dict[str, Any],
    ) -> str:
        """Create specific prompt for introduction sections"""
        return f"""Draft a BRIEF INTRODUCTION for this legal motion:

Motion Title: {outline_section.title}
Maximum Length: 150-200 words (1-2 paragraphs ONLY)

Key Arguments to Preview (mention briefly):
{chr(10).join(f"- {point}" for point in essential_content["key_points"][:3])}

STRICT REQUIREMENTS:
1. First sentence: State party name, representation, and the specific motion/objection being filed
2. Second sentence (optional): Very brief statement of why relief should be granted
3. End with: "for the reasons set forth below" or similar transition
4. DO NOT exceed 200 words total
5. DO NOT include detailed arguments - save those for the memorandum section

Example format:
"Plaintiff, [NAME], as Personal Representative of the Estate of [NAME], by and through undersigned counsel, files this [Motion Type] and respectfully shows this Court that [very brief reason]. For the reasons set forth below, [relief sought]."

Write the introduction now:"""

    def _create_conclusion_prompt(
        self,
        outline_section: OutlineSection,
        essential_content: Dict[str, Any],
        cumulative_context: Dict[str, Any],
    ) -> str:
        """Create specific prompt for conclusion sections"""
        prev_args = cumulative_context.get("summary", "the arguments presented above")
        return f"""Draft a BRIEF CONCLUSION for this legal motion:

Maximum Length: 100-150 words (3-4 sentences ONLY)

Main Arguments Made in Motion:
{prev_args}

Key Relief Sought:
{chr(10).join(f"- {point}" for point in essential_content["key_points"][:2])}

STRICT REQUIREMENTS:
1. Summarize in 2-3 sentences why the law and facts require granting relief
2. Final sentence: State what the defendant's motion should be (granted/denied)
3. DO NOT exceed 150 words total
4. DO NOT rehash arguments - just state the conclusion
5. Lead naturally into the WHEREFORE paragraph

Example format:
"The law and equity both require that [core principle/finding]. Based on [key reason], [secondary point if needed]. Accordingly, Defendant's Motion should be [denied/granted]."

Write the conclusion now:"""

    def _create_prayer_prompt(
        self,
        outline_section: OutlineSection,
        essential_content: Dict[str, Any],
        cumulative_context: Dict[str, Any],
    ) -> str:
        """Create specific prompt for prayer/wherefore sections"""
        return f"""Draft the WHEREFORE/PRAYER FOR RELIEF section:

Relief Requested:
{chr(10).join(f"- {point}" for point in essential_content["key_points"][:3])}

STRICT FORMAT:
1. Start with: "WHEREFORE, Plaintiff respectfully requests that this Honorable Court"
2. List specific relief sought
3. End with: "and grant such other relief as the Court deems just and equitable."
4. Add signature block

Example:
WHEREFORE, Plaintiff respectfully requests that this Honorable Court deny Defendant's Motion for [specific relief] and grant such other relief as the Court deems just and equitable.

[LAW FIRM NAME]

/s/ [Attorney Name]
[Attorney Name], Esq.
Florida Bar No.: [Number]
[Address]
[Phone]
[Email]

Write the WHEREFORE section now:"""

    # Keep all other methods unchanged from the original...
    async def _convert_structure_to_sections(
        self, outline_id: str, outline_structure: List[Dict[str, Any]]
    ) -> List[OutlineSection]:
        """Convert outline structure back to OutlineSection objects for compatibility"""
        sections = []
        logger.info(
            f"[CONVERT] Converting {len(outline_structure)} structure items to sections"
        )

        for i, section_info in enumerate(outline_structure):
            try:
                logger.debug(f"[CONVERT] Processing section {i}: {section_info}")
                # Get the full section data from cache
                section_data = await outline_cache.get_section(
                    outline_id, section_info["index"]
                )
                if not section_data:
                    logger.warning(
                        f"Could not retrieve section {section_info['index']} from cache"
                    )
                    continue

                # Extract ALL content for search purposes (not just first 3)
                content_points = []
                legal_authorities = []
                key_facts = []
                themes = []

                # Process ALL content items for database search
                for item in section_data.get("content", []):
                    item_type = item.get("type", "")

                    if item_type == "field":
                        label = item.get("label", "")
                        value = item.get("value", "")

                        # Clean multi-option values but keep full content
                        if "||" in value:
                            options = value.split("||")
                            # Add first option as primary, but keep others for search
                            for i, option in enumerate(options[:3]):
                                clean_option = option.strip()
                                if clean_option.startswith("Option"):
                                    clean_option = (
                                        clean_option.split(":", 1)[1].strip()
                                        if ":" in clean_option
                                        else clean_option
                                    )
                                content_points.append(clean_option.strip('"').strip())
                        else:
                            content_points.append(value.strip('"').strip())

                        # Special handling for themes
                        if "theme" in label.lower():
                            themes.extend(self._extract_themes_from_value(value))

                    elif item_type == "list":
                        label = item.get("label", "").lower()
                        items = item.get("items", [])

                        if "authorit" in label:
                            legal_authorities.extend(items)
                        elif "fact" in label:
                            key_facts.extend(items)
                        else:
                            # Add list items as content points
                            content_points.extend([str(item) for item in items])

                    elif item_type == "paragraph":
                        text = item.get("text", "")
                        if text:
                            content_points.append(text)

                # Create OutlineSection with FULL content for search
                section = OutlineSection(
                    id=f"section_{section_info['index']}",
                    title=section_info.get(
                        "heading", f"Section {section_info['index'] + 1}"
                    ),
                    section_type=self._determine_section_type(
                        section_info.get("heading", "")
                    ),
                    content_points=content_points,  # Keep ALL content points
                    legal_authorities=legal_authorities,  # Keep ALL authorities
                    target_length=500,  # Will be updated later
                    context_summary=section_data.get("summary", ""),
                    hook_options=section_data.get("hook_options", []),
                    themes=themes,
                    key_facts=key_facts,
                    counter_arguments=section_data.get("counter_arguments", []),
                )
                sections.append(section)

            except Exception as e:
                logger.warning(
                    f"Error converting section {section_info['index']}: {str(e)}"
                )
                continue

        logger.info(
            f"[CONVERT] Successfully converted {len(sections)} sections with content"
        )
        return sections

    def _create_minimal_outline_section(
        self,
        section_data: Dict[str, Any],
        section_info: Dict[str, Any],
        target_length: int,
    ) -> OutlineSection:
        """Create a minimal OutlineSection object for processing"""

        # Extract only essential content points (limit to prevent overflow)
        content_points = []
        content_items = section_data.get("content", [])[:3]  # Only first 3 items

        for item in content_items:
            if item.get("type") == "field":
                label = item.get("label", "")
                value = item.get("value", "")[:200]  # Truncate to 200 chars
                content_points.append(f"{label}: {value}")
            elif item.get("type") == "paragraph":
                content_points.append(item.get("text", "")[:200])

        # Extract limited authorities
        legal_authorities = []
        for item in content_items:
            if (
                item.get("type") == "list"
                and "authorities" in item.get("label", "").lower()
            ):
                legal_authorities.extend(item.get("items", [])[:5])  # Max 5 authorities

        return OutlineSection(
            id=f"section_{section_info['index']}",
            title=section_info["heading"],
            section_type=self._determine_section_type(section_info["heading"]),
            content_points=content_points[:5],  # Max 5 points
            legal_authorities=legal_authorities[:5],  # Max 5 authorities
            target_length=target_length,
            context_summary=f"Section {section_info['index'] + 1} of {section_info.get('type', 'standard')} type",
        )

    def _extract_essential_content(
        self, section_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Extract essential content from section data without overwhelming tokens"""

        essential = {"key_points": [], "authorities": [], "facts": [], "themes": []}

        # Process content items
        for item in section_data.get("content", []):
            item_type = item.get("type", "")

            if item_type == "field":
                label = item.get("label", "")
                value = item.get("value", "")

                # Handle multi-option values
                if "||" in value:
                    # Take first option only
                    value = value.split("||")[0].strip()
                    if value.startswith("Option"):
                        value = (
                            value.split(":", 1)[1].strip() if ":" in value else value
                        )

                # Truncate and clean
                value = value.strip('"').strip()[:300]  # 300 char limit

                if label == "Theme":
                    essential["themes"].append(value)
                elif label == "Summary":
                    essential["key_points"].insert(0, value)  # Put summary first
                else:
                    essential["key_points"].append(f"{label}: {value}")

            elif item_type == "list":
                label = item.get("label", "").lower()
                items = item.get("items", [])[:5]  # Max 5 items per list

                if "authorit" in label:
                    essential["authorities"].extend(items)
                elif "fact" in label:
                    essential["facts"].extend(items)
                else:
                    # Convert list items to key points
                    for list_item in items[:3]:  # Max 3 items
                        if isinstance(list_item, str):
                            essential["key_points"].append(list_item[:200])

        # Limit totals
        essential["key_points"] = essential["key_points"][:7]
        essential["authorities"] = essential["authorities"][:7]
        essential["facts"] = essential["facts"][:5]
        essential["themes"] = essential["themes"][:3]

        return essential

    def _build_lightweight_cumulative_context(
        self, drafted_sections: List[DraftedSection], case_context: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Build minimal cumulative context to prevent token overflow"""

        context = {"summary": "", "citations_used": set(), "key_holdings": []}

        if not drafted_sections:
            context["summary"] = "This is the first section of the motion."
            return context

        # Build brief summary from last 2 sections only
        recent_sections = drafted_sections[-2:]
        summaries = []

        for section in recent_sections:
            # Extract brief conclusion (first 100 chars of last paragraph)
            paragraphs = section.content.split("\n\n")
            if paragraphs:
                last_para = paragraphs[-1][:100]
                summaries.append(f"{section.outline_section.title}: {last_para}...")

            # Collect citations (limit to 10)
            for citation in section.citations_used[:10]:
                context["citations_used"].add(citation)

        context["summary"] = " | ".join(summaries)
        context["citations_used"] = list(context["citations_used"])[:15]

        return context

    def _calculate_word_distribution_from_structure(
        self, outline_structure: List[Dict[str, Any]], target_words: int
    ) -> Dict[str, int]:
        """Calculate word distribution based on flexible legal motion structure guidelines"""

        distribution = {}

        # First pass: categorize sections
        intro_sections = []
        argument_sections = []
        conclusion_sections = []
        procedural_sections = []
        other_sections = []

        for i, section in enumerate(outline_structure):
            section_type = self._determine_section_type(section.get("heading", ""))
            section_id = f"section_{i}"

            if section_type in [SectionType.PREINTRODUCTION, SectionType.INTRODUCTION]:
                intro_sections.append(section_id)
            elif section_type in [
                SectionType.ARGUMENT,
                SectionType.SUB_ARGUMENT,
                SectionType.MEMORANDUM_OF_LAW,
                SectionType.LEGAL_STANDARD,
            ]:
                argument_sections.append(section_id)
            elif section_type in [
                SectionType.CONCLUSION,
                SectionType.PRAYER_FOR_RELIEF,
            ]:
                conclusion_sections.append(section_id)
            elif section_type == SectionType.PROCEDURAL_BACKGROUND:
                procedural_sections.append(section_id)
            else:
                other_sections.append(section_id)

        # FLEXIBLE PROPORTIONAL GUIDELINES (not strict limits)
        # Target ranges:
        # - Introduction: 5-10% (but allow flexibility for complex cases)
        # - Arguments: 75-85% (main body should dominate)
        # - Conclusion: 5-10% (keep concise but not rigid)
        # - Other sections: As needed

        # Calculate base allocations
        intro_target = int(target_words * 0.07)  # 7% baseline
        conclusion_target = int(target_words * 0.07)  # 7% baseline
        procedural_target = (
            int(target_words * 0.06) if procedural_sections else 0
        )  # 6% if exists

        # Remaining goes to arguments (usually 75-80%)
        argument_target = (
            target_words - intro_target - conclusion_target - procedural_target
        )

        # Distribute intro words with flexibility
        if intro_sections:
            words_per_intro = intro_target // len(intro_sections)
            for section_id in intro_sections:
                # Suggest but don't hard cap - allow 150-300 words typically
                distribution[section_id] = max(150, min(words_per_intro, 300))

        # Distribute argument words based on complexity
        if argument_sections:
            # Weight arguments by content richness
            arg_weights = {}
            total_arg_weight = 0

            for idx, section_id in enumerate(argument_sections):
                section = outline_structure[int(section_id.split("_")[1])]
                content_items = len(section.get("content", []))

                # More complex arguments get more weight
                # Base weight 1.0, up to 2.0 for content-rich sections
                weight = 1.0 + min(content_items * 0.1, 1.0)
                arg_weights[section_id] = weight
                total_arg_weight += weight

            # Distribute proportionally with minimum thresholds
            for section_id, weight in arg_weights.items():
                allocated = int((weight / total_arg_weight) * argument_target)
                # Ensure substantial arguments get adequate space
                distribution[section_id] = max(
                    allocated, 400
                )  # Min 400 words per argument

        # Distribute conclusion words with flexibility
        if conclusion_sections:
            words_per_conclusion = conclusion_target // len(conclusion_sections)
            for section_id in conclusion_sections:
                # Suggest 100-250 words for conclusions
                distribution[section_id] = max(100, min(words_per_conclusion, 250))

        # Handle procedural background
        if procedural_sections:
            words_per_procedural = procedural_target // len(procedural_sections)
            for section_id in procedural_sections:
                distribution[section_id] = words_per_procedural

        # Handle other sections flexibly
        remaining_words = target_words - sum(distribution.values())
        if other_sections and remaining_words > 0:
            words_per_other = remaining_words // len(other_sections)
            for section_id in other_sections:
                distribution[section_id] = max(words_per_other, 200)

        # Log proportions for guidance (not enforcement)
        total_allocated = sum(distribution.values())
        intro_proportion = (
            sum(distribution.get(s, 0) for s in intro_sections) / total_allocated * 100
            if total_allocated > 0
            else 0
        )
        arg_proportion = (
            sum(distribution.get(s, 0) for s in argument_sections)
            / total_allocated
            * 100
            if total_allocated > 0
            else 0
        )
        conclusion_proportion = (
            sum(distribution.get(s, 0) for s in conclusion_sections)
            / total_allocated
            * 100
            if total_allocated > 0
            else 0
        )

        logger.info(
            f"Word distribution guidance - "
            f"Intro: {intro_proportion:.1f}% (target: 5-10%), "
            f"Arguments: {arg_proportion:.1f}% (target: 75-85%), "
            f"Conclusion: {conclusion_proportion:.1f}% (target: 5-10%)"
        )

        return distribution

    async def _expand_section_efficiently(
        self,
        drafted_section: DraftedSection,
        full_section_data: Dict[str, Any],
        target_length: int,
        case_context: Dict[str, Any],
    ) -> DraftedSection:
        """Expand section efficiently without overwhelming context"""

        try:
            current_length = drafted_section.word_count
            expansion_needed = target_length - current_length

            if expansion_needed <= 0:
                return drafted_section

            # Flexible expansion guidance based on section type
            section_type = drafted_section.outline_section.section_type
            if section_type == SectionType.INTRODUCTION:
                # Allow some flexibility but discourage excessive expansion
                if current_length >= 300:
                    logger.info(
                        "Introduction already substantial - consider keeping concise"
                    )
                    return drafted_section
                elif current_length >= 250:
                    # Limit expansion for already decent-sized intros
                    expansion_needed = min(expansion_needed, 50)
            elif section_type == SectionType.CONCLUSION:
                # Similar flexibility for conclusions
                if current_length >= 250:
                    logger.info(
                        "Conclusion already substantial - consider keeping concise"
                    )
                    return drafted_section
                elif current_length >= 200:
                    # Limit expansion for already decent-sized conclusions
                    expansion_needed = min(expansion_needed, 50)

            # Extract unused content from full section data
            unused_content = self._extract_unused_content(
                drafted_section.content, full_section_data
            )

            expansion_prompt = f"""Expand this legal section by adding {expansion_needed} words.

Current Section ({current_length} words):
{drafted_section.content[:2000]}... [truncated]

Unused Content Points:
{chr(10).join(f"- {point}" for point in unused_content[:5])}

REQUIREMENTS:
1. Add {expansion_needed} words of substantive content
2. DO NOT remove any existing content
3. Add deeper analysis and more detailed application
4. Maintain formal legal style
5. Include smooth transitions

Provide the COMPLETE expanded section."""

            result = await asyncio.wait_for(
                self.section_writer.run(expansion_prompt), timeout=30
            )

            expanded_content = (
                str(result.data) if hasattr(result, "data") else str(result)
            )

            # Update section
            drafted_section.content = expanded_content
            drafted_section.word_count = len(expanded_content.split())
            drafted_section.expansion_cycles += 1

            return drafted_section

        except asyncio.TimeoutError:
            logger.warning(
                f"Timeout during section expansion for: {drafted_section.outline_section.title}"
            )
            return drafted_section
        except Exception as e:
            logger.error(
                f"Error in efficient expansion for section '{drafted_section.outline_section.title}': {str(e)}",
                exc_info=True,
            )
            # Add a note about the expansion failure
            drafted_section.revision_notes.append(f"Expansion failed: {str(e)[:100]}")
            return drafted_section

    def _extract_unused_content(
        self, current_content: str, full_section_data: Dict[str, Any]
    ) -> List[str]:
        """Extract content points not yet used in the section"""

        unused = []
        content_lower = current_content.lower()

        # Check all content items
        for item in full_section_data.get("content", []):
            if item.get("type") == "field":
                value = item.get("value", "")
                # Check if this content is not yet used
                if value and value[:50].lower() not in content_lower:
                    unused.append(value[:200])

            elif item.get("type") == "list":
                for list_item in item.get("items", [])[:3]:
                    if (
                        isinstance(list_item, str)
                        and list_item[:30].lower() not in content_lower
                    ):
                        unused.append(list_item[:150])

        return unused[:7]  # Return max 7 unused points

    async def _lightweight_review_process(
        self, motion_draft: MotionDraft
    ) -> MotionDraft:
        """Enhanced review process with comprehensive quality metrics tracking"""

        try:
            # Import citation formatter for analysis
            from src.ai_agents.citation_formatter import citation_formatter

            # Validate section structure and lengths
            structure_issues = self._validate_motion_structure(motion_draft)
            motion_draft.review_notes.extend(structure_issues)

            # Initialize comprehensive metrics
            evidence_quality_scores = []
            citation_density_scores = []
            argument_strength_scores = []

            for section in motion_draft.sections:
                section_type = section.outline_section.section_type

                # Calculate evidence quality metrics
                if section_type in [
                    SectionType.ARGUMENT,
                    SectionType.MEMORANDUM_OF_LAW,
                ]:
                    # Extract citations from content
                    citations = citation_formatter.extract_citations_from_text(
                        section.content
                    )

                    # Citation density (citations per 100 words)
                    citation_density = (
                        (len(citations) / section.word_count * 100)
                        if section.word_count > 0
                        else 0
                    )
                    citation_density_scores.append(citation_density)

                    # Evidence specificity score
                    specific_citations = 0
                    for citation in citations:
                        # Check for specific references (page:line, exhibit numbers, dates)
                        if (
                            citation.page_reference
                            or citation.line_reference
                            or citation.exhibit_number
                            or citation.date
                        ):
                            specific_citations += 1

                    evidence_specificity = (
                        (specific_citations / len(citations)) if citations else 0
                    )
                    evidence_quality_scores.append(evidence_specificity)

                    # Argument strength based on evidence integration
                    has_deposition = any(
                        c.citation_type.value == "deposition" for c in citations
                    )
                    has_exhibit = any(
                        c.citation_type.value == "exhibit" for c in citations
                    )
                    has_expert = any(
                        c.citation_type.value == "expert_report" for c in citations
                    )

                    argument_strength = (
                        sum([has_deposition, has_exhibit, has_expert]) / 3.0
                    )
                    argument_strength_scores.append(argument_strength)

                # Section-specific quality notes
                if section_type == SectionType.INTRODUCTION:
                    if section.word_count > 300:
                        motion_draft.review_notes.append(
                            f"Introduction is {section.word_count} words - consider making more concise"
                        )

                elif section_type == SectionType.CONCLUSION:
                    if section.word_count > 250:
                        motion_draft.review_notes.append(
                            f"Conclusion is {section.word_count} words - consider making more concise"
                        )

            # Calculate overall coherence score
            coherence_components = []

            # Structure balance component
            total_words = motion_draft.total_word_count
            if total_words > 0:
                intro_words = sum(
                    s.word_count
                    for s in motion_draft.sections
                    if s.outline_section.section_type
                    in [SectionType.INTRODUCTION, SectionType.PREINTRODUCTION]
                )
                argument_words = sum(
                    s.word_count
                    for s in motion_draft.sections
                    if s.outline_section.section_type
                    in [SectionType.ARGUMENT, SectionType.MEMORANDUM_OF_LAW]
                )

                intro_proportion = intro_words / total_words
                argument_proportion = argument_words / total_words

                # Score based on how close to ideal proportions
                structure_score = 0.0
                if 0.05 <= intro_proportion <= 0.10:
                    structure_score += 0.5
                elif intro_proportion <= 0.15:
                    structure_score += 0.3

                if 0.75 <= argument_proportion <= 0.85:
                    structure_score += 0.5
                elif argument_proportion >= 0.70:
                    structure_score += 0.3

                coherence_components.append(structure_score)

            # Evidence quality component
            if evidence_quality_scores:
                avg_evidence_quality = sum(evidence_quality_scores) / len(
                    evidence_quality_scores
                )
                coherence_components.append(avg_evidence_quality)

            # Citation density component (target: 2-4 citations per 100 words)
            if citation_density_scores:
                avg_citation_density = sum(citation_density_scores) / len(
                    citation_density_scores
                )
                density_score = min(avg_citation_density / 3.0, 1.0)  # Cap at 1.0
                coherence_components.append(density_score)

            # Argument strength component
            if argument_strength_scores:
                avg_argument_strength = sum(argument_strength_scores) / len(
                    argument_strength_scores
                )
                coherence_components.append(avg_argument_strength)

            # Calculate final coherence score
            if coherence_components:
                motion_draft.coherence_score = sum(coherence_components) / len(
                    coherence_components
                )
            else:
                motion_draft.coherence_score = 0.7  # Default

            # Compile comprehensive quality metrics
            motion_draft.quality_metrics = {
                # Basic metrics
                "sections_complete": float(
                    len([s for s in motion_draft.sections if s.word_count > 0])
                ),
                "total_word_count": float(total_words),
                "total_pages": float(motion_draft.total_page_estimate),
                # Structure metrics (as proportions, not hard limits)
                "intro_proportion": round((intro_words / total_words * 100), 1)
                if total_words > 0
                else 0.0,
                "argument_proportion": round((argument_words / total_words * 100), 1)
                if total_words > 0
                else 0.0,
                "conclusion_proportion": round(
                    (
                        sum(
                            s.word_count
                            for s in motion_draft.sections
                            if s.outline_section.section_type
                            in [SectionType.CONCLUSION, SectionType.PRAYER_FOR_RELIEF]
                        )
                        / total_words
                        * 100
                    ),
                    1,
                )
                if total_words > 0
                else 0.0,
                # Evidence metrics
                "average_citation_density": round(
                    sum(citation_density_scores) / len(citation_density_scores), 2
                )
                if citation_density_scores
                else 0.0,
                "evidence_specificity_score": round(
                    sum(evidence_quality_scores) / len(evidence_quality_scores), 2
                )
                if evidence_quality_scores
                else 0.0,
                "argument_strength_average": round(
                    sum(argument_strength_scores) / len(argument_strength_scores), 2
                )
                if argument_strength_scores
                else 0.0,
                # Citation breakdown
                "total_unique_citations": float(len(motion_draft.citation_index)),
                "depositions_cited": float(
                    sum(1 for c in motion_draft.citation_index if "Dep." in c)
                ),
                "exhibits_cited": float(
                    sum(1 for c in motion_draft.citation_index if "Ex." in c)
                ),
                # Quality indicators
                "coherence_score": round(motion_draft.coherence_score, 2),
                "average_section_confidence": round(
                    sum(s.confidence_score for s in motion_draft.sections)
                    / len(motion_draft.sections),
                    2,
                )
                if motion_draft.sections
                else 0.0,
            }

            # Add actionable guidance notes (not hard requirements)
            if motion_draft.quality_metrics["average_citation_density"] < 1.0:
                motion_draft.review_notes.append(
                    "Consider adding more specific evidence citations to support arguments"
                )
            elif motion_draft.quality_metrics["average_citation_density"] > 5.0:
                motion_draft.review_notes.append(
                    "High citation density - ensure citations flow naturally with the text"
                )

            if motion_draft.quality_metrics["evidence_specificity_score"] < 0.5:
                motion_draft.review_notes.append(
                    "Many citations lack specific page/line references - consider adding for stronger support"
                )

            if motion_draft.quality_metrics["argument_strength_average"] < 0.5:
                motion_draft.review_notes.append(
                    "Arguments could benefit from more diverse evidence types (depositions, exhibits, expert reports)"
                )

            return motion_draft

        except Exception as e:
            logger.error(f"Error in review process: {str(e)}")
            motion_draft.coherence_score = 0.7
            motion_draft.quality_metrics = {"error": str(e)}
            return motion_draft

    def _validate_motion_structure(self, motion_draft: MotionDraft) -> List[str]:
        """Validate the structure of the motion"""
        issues = []

        # Check for required sections
        section_types = [s.outline_section.section_type for s in motion_draft.sections]

        if SectionType.INTRODUCTION not in section_types:
            issues.append("Missing introduction section")

        if (
            SectionType.MEMORANDUM_OF_LAW not in section_types
            and SectionType.ARGUMENT not in section_types
        ):
            issues.append("Missing memorandum of law/argument sections")

        if SectionType.CONCLUSION not in section_types:
            issues.append("Missing conclusion section")

        # Check section order
        expected_order = [
            SectionType.PREINTRODUCTION,
            SectionType.INTRODUCTION,
            SectionType.PROCEDURAL_BACKGROUND,
            SectionType.STATEMENT_OF_FACTS,
            SectionType.MEMORANDUM_OF_LAW,
            SectionType.ARGUMENT,
            SectionType.CONCLUSION,
            SectionType.PRAYER_FOR_RELIEF,
        ]

        # Validate order (simplified check)
        last_index = -1
        for section in motion_draft.sections:
            if section.outline_section.section_type in expected_order:
                current_index = expected_order.index(
                    section.outline_section.section_type
                )
                if current_index < last_index:
                    issues.append(
                        f"Section order issue: {section.outline_section.section_type.value} appears out of order"
                    )
                last_index = max(last_index, current_index)

        return issues

    def _extract_title_from_cache(self, outline: Dict[str, Any]) -> str:
        """Extract title from outline with minimal processing"""
        return outline.get("title") or outline.get("motion_title") or "Legal Motion"

    def _calculate_section_confidence(
        self, content: str, section: OutlineSection, word_count: int
    ) -> float:
        """Calculate confidence score for section"""
        score = 0.0
        max_score = 100.0

        # Length score (30 points)
        if word_count >= section.target_length:
            score += 30
        else:
            score += (word_count / section.target_length) * 30

        # Content points addressed (25 points)
        points_addressed = sum(
            1
            for point in section.content_points
            if any(keyword in content.lower() for keyword in point.lower().split()[:5])
        )
        if section.content_points:
            score += (points_addressed / len(section.content_points)) * 25
        else:
            score += 25

        # Citations included (25 points)
        citations_found = sum(
            1
            for auth in section.legal_authorities
            if auth in content or auth.replace(" ", "") in content
        )
        if section.legal_authorities:
            score += (citations_found / len(section.legal_authorities)) * 25
        else:
            score += 25

        # Structure quality (20 points)
        structure_indicators = [
            "first",
            "second",
            "third",
            "moreover",
            "furthermore",
            "additionally",
            "however",
            "nevertheless",
            "therefore",
            "accordingly",
            "issue",
            "rule",
            "application",
            "conclusion",
        ]
        structure_score = sum(
            2 for indicator in structure_indicators if indicator in content.lower()
        )
        score += min(structure_score, 20)

        return score / max_score

    def _update_document_context(self, section: DraftedSection):
        """Update document-wide context with section information"""

        # Update citations used
        self.document_context["citations_used"].update(section.citations_used)

        # Update terminology
        terms = self._extract_key_terms(section.content)
        for term in terms:
            if term not in self.document_context["terminology"]:
                self.document_context["terminology"][term] = section.outline_section.id

    def _extract_key_terms(self, content: str) -> List[str]:
        """Extract key legal terms from content"""
        # Simple implementation - could be enhanced
        terms = []
        legal_terms = [
            "negligence",
            "liability",
            "duty",
            "breach",
            "damages",
            "causation",
            "standard of care",
            "reasonable person",
            "proximate cause",
            "foreseeability",
        ]

        content_lower = content.lower()
        for term in legal_terms:
            if term in content_lower:
                terms.append(term)

        return terms[:5]

    def _restructure_motion_sections(
        self, sections: List[DraftedSection]
    ) -> List[DraftedSection]:
        """Restructure sections to ensure proper legal motion format with professional formatting"""

        # Group sections by type
        caption_sections = []
        intro_sections = []
        procedural_sections = []
        argument_sections = []
        conclusion_sections = []
        prayer_sections = []
        other_sections = []

        for section in sections:
            section_type = section.outline_section.section_type

            if section_type == SectionType.PREINTRODUCTION:
                caption_sections.append(section)
            elif section_type == SectionType.INTRODUCTION:
                intro_sections.append(section)
            elif section_type == SectionType.PROCEDURAL_BACKGROUND:
                procedural_sections.append(section)
            elif section_type in [
                SectionType.ARGUMENT,
                SectionType.SUB_ARGUMENT,
                SectionType.MEMORANDUM_OF_LAW,
                SectionType.LEGAL_STANDARD,
            ]:
                argument_sections.append(section)
            elif section_type == SectionType.CONCLUSION:
                conclusion_sections.append(section)
            elif section_type == SectionType.PRAYER_FOR_RELIEF:
                prayer_sections.append(section)
            else:
                other_sections.append(section)

        # Create restructured list in proper order
        restructured = []

        # 1. Add professional case caption if not present
        if not caption_sections:
            # Create a professional caption
            caption_content = legal_formatter.format_case_caption(
                case_name=self.document_context.get(
                    "case_name", "Case Name v. Defendant"
                ),
                case_number="2024-XXXXX-CA-01",
                plaintiff_name="Plaintiff Name",
                defendant_names=["Defendant Name"],
            )

            # Add motion title
            motion_title = self.document_context.get(
                "motion_title", "MOTION TO DISMISS"
            )
            caption_content += legal_formatter.format_motion_title(motion_title)

            caption_section = OutlineSection(
                id="caption",
                title="Caption",
                section_type=SectionType.PREINTRODUCTION,
                content_points=[],
                legal_authorities=[],
                target_length=100,
            )

            caption_drafted = DraftedSection(
                outline_section=caption_section,
                content=caption_content,
                word_count=len(caption_content.split()),
                citations_used=[],
                citations_verified={},
                expansion_cycles=0,
                confidence_score=1.0,
            )

            restructured.append(caption_drafted)
        else:
            restructured.extend(caption_sections)

        # 2. Introduction sections
        restructured.extend(intro_sections)

        # 3. Procedural background (if any)
        if procedural_sections:
            # Format procedural background with numbered paragraphs
            for proc_section in procedural_sections:
                # Ensure numbered paragraphs
                paragraphs = proc_section.content.split("\n\n")
                formatted_paragraphs = []
                para_num = 1

                for para in paragraphs:
                    if para.strip() and not para.strip()[0].isdigit():
                        formatted_paragraphs.append(f"{para_num}. {para.strip()}")
                        para_num += 1
                    else:
                        formatted_paragraphs.append(para)

                proc_section.content = "\n\n".join(formatted_paragraphs)

            restructured.extend(procedural_sections)

        # 4. Merge all arguments under MEMORANDUM OF LAW
        if argument_sections:
            # Create a single MEMORANDUM OF LAW section that contains all arguments
            memo_content = "\nMEMORANDUM OF LAW\n\n"
            total_word_count = 0
            all_citations = []

            argument_count = 0
            for i, arg_section in enumerate(argument_sections):
                # Add Roman numeral for each main argument
                if arg_section.outline_section.section_type == SectionType.ARGUMENT:
                    argument_count += 1
                    roman_num = self._to_roman(argument_count)

                    # Format the argument with legal formatter
                    arg_title = arg_section.outline_section.title
                    if not arg_title.startswith("ARGUMENT"):
                        arg_title = f"{arg_title}"

                    # Use legal formatter for proper argument structure
                    formatted_arg = legal_formatter.format_legal_argument(
                        argument_number=roman_num,
                        argument_heading=arg_title.upper(),
                        argument_text=arg_section.content,
                    )
                    memo_content += formatted_arg
                else:
                    # Sub-arguments or supporting sections
                    memo_content += f"\n{arg_section.content}\n\n"

                total_word_count += arg_section.word_count
                all_citations.extend(arg_section.citations_used)

            # Create merged memorandum section
            memo_section = OutlineSection(
                id="memorandum_of_law",
                title="MEMORANDUM OF LAW",
                section_type=SectionType.MEMORANDUM_OF_LAW,
                content_points=[],
                legal_authorities=[],
                target_length=total_word_count,
            )

            merged_memo = DraftedSection(
                outline_section=memo_section,
                content=memo_content.strip(),
                word_count=total_word_count,
                citations_used=list(set(all_citations)),
                citations_verified={},
                expansion_cycles=1,
                confidence_score=0.85,
            )

            restructured.append(merged_memo)

        # 5. Other sections (facts, etc. - usually skipped)
        restructured.extend(other_sections)

        # 6. Conclusion (ensure only one)
        if conclusion_sections:
            restructured.append(conclusion_sections[0])
            if len(conclusion_sections) > 1:
                logger.warning(
                    "Multiple conclusion sections found, using only the first"
                )

        # 7. Prayer for relief with proper WHEREFORE formatting
        if prayer_sections:
            prayer_section = prayer_sections[0]

            # Extract relief requests from content
            relief_requests = []
            content_lines = prayer_section.content.split("\n")
            for line in content_lines:
                if line.strip() and not line.strip().startswith("WHEREFORE"):
                    relief_requests.append(line.strip())

            # Format with legal formatter
            formatted_wherefore = legal_formatter.format_wherefore_clause(
                party_name="Plaintiff",
                relief_requested=relief_requests[:3],  # Main relief items
                include_costs=True,
                include_further_relief=True,
            )

            # Add signature block
            formatted_wherefore += legal_formatter.format_signature_block(
                attorney_name="Attorney Name",
                bar_number="12345",
                firm_name="Law Firm Name",
                address=["123 Main Street", "Suite 100", "Miami, FL 33131"],
                phone="(305) 555-1234",
                email="attorney@lawfirm.com",
            )

            prayer_section.content = formatted_wherefore
            restructured.append(prayer_section)

        logger.info(
            f"Restructured {len(sections)} sections into {len(restructured)} sections with professional formatting"
        )
        return restructured

    def _to_roman(self, num: int) -> str:
        """Convert number to Roman numeral"""
        values = [
            (1000, "M"),
            (900, "CM"),
            (500, "D"),
            (400, "CD"),
            (100, "C"),
            (90, "XC"),
            (50, "L"),
            (40, "XL"),
            (10, "X"),
            (9, "IX"),
            (5, "V"),
            (4, "IV"),
            (1, "I"),
        ]
        result = ""
        for value, letter in values:
            count = num // value
            if count:
                result += letter * count
                num -= value * count
        return result

    def _build_citation_index(self, motion_draft: MotionDraft) -> Dict[str, List[str]]:
        """Build index of citations to sections using them"""
        citation_index = {}

        for section in motion_draft.sections:
            for citation in section.citations_used:
                if citation not in citation_index:
                    citation_index[citation] = []
                citation_index[citation].append(section.outline_section.id)

        return citation_index

    def _extract_citations_from_text(self, text: str) -> List[str]:
        """Extract all citations from text using compiled patterns"""
        citations = []

        for pattern in self.citation_patterns:
            matches = pattern.findall(text)
            citations.extend(matches)

        # Clean and deduplicate
        citations = list(set(citation.strip() for citation in citations))

        return citations

    def export_to_docx(self, motion_draft: MotionDraft, output_path: str):
        """Export motion draft to DOCX format"""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Create document
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add title
        title = doc.add_heading(motion_draft.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add case caption
        doc.add_paragraph(motion_draft.case_name)
        doc.add_paragraph()

        # Add each section with transitions
        for i, section in enumerate(motion_draft.sections):
            # Section heading
            doc.add_heading(section.outline_section.title, level=1)

            # Add transition from previous if available
            if "from_previous" in section.transitions and i > 0:
                transition_para = doc.add_paragraph(
                    section.transitions["from_previous"]
                )
                transition_para.style.font.italic = True
                doc.add_paragraph()

            # Section content
            for paragraph in section.content.split("\n\n"):
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph)
                    p.style.font.size = Pt(12)
                    p.style.font.name = "Times New Roman"

            # Add transition to next if available
            if "to_next" in section.transitions and i < len(motion_draft.sections) - 1:
                doc.add_paragraph()
                transition_para = doc.add_paragraph(section.transitions["to_next"])
                transition_para.style.font.italic = True

        # Add metadata page
        doc.add_page_break()
        doc.add_heading("Document Metadata", level=1)

        # Basic metadata
        doc.add_paragraph(f"Total Words: {motion_draft.total_word_count:,}")
        doc.add_paragraph(f"Estimated Pages: {motion_draft.total_page_estimate}")
        doc.add_paragraph(
            f"Created: {motion_draft.creation_timestamp.strftime('%B %d, %Y at %I:%M %p')}"
        )
        doc.add_paragraph(f"Document Quality Score: {motion_draft.coherence_score:.2%}")

        # Quality metrics
        if motion_draft.quality_metrics:
            doc.add_heading("Quality Metrics", level=2)
            for metric, score in motion_draft.quality_metrics.items():
                doc.add_paragraph(f"• {metric.replace('_', ' ').title()}: {score:.2%}")

        # Citation index
        if motion_draft.citation_index:
            doc.add_heading("Citation Index", level=2)
            doc.add_paragraph(
                f"Total Unique Citations: {len(motion_draft.citation_index)}"
            )

            # List top citations
            sorted_citations = sorted(
                motion_draft.citation_index.items(),
                key=lambda x: len(x[1]),
                reverse=True,
            )[:10]

            for citation, sections in sorted_citations:
                doc.add_paragraph(f"• {citation} (used in {len(sections)} sections)")

        # Review notes
        if motion_draft.review_notes:
            doc.add_heading("Review Notes", level=2)
            for note in motion_draft.review_notes:
                doc.add_paragraph(f"• {note}")

        # Save document
        doc.save(output_path)
        logger.info(f"Motion exported to {output_path}")


# Create global instance
motion_drafter = EnhancedMotionDraftingAgent()
