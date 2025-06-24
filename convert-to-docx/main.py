from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import FastAPI, Request, File, UploadFile, HTTPException
from fastapi.responses import FileResponse, JSONResponse
from tempfile import NamedTemporaryFile
import re
from typing import List, Dict, Any
import io

app = FastAPI()

def set_doc_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

def add_bullet_points(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(6)

def add_numbered_points(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.space_after = Pt(6)

def parse_docx_to_sections(doc):
    """Parse a DOCX document into sections for sequential processing"""
    result = {
        "title": "",
        "sections": [],
        "style_notes": []
    }
    
    current_section = None
    current_subsection = None
    collecting_bullets = False
    bullet_collector = []
    in_style_notes = False
    
    # Helper function to save current section
    def save_current_section():
        nonlocal current_section
        if current_section:
            result["sections"].append(current_section)
            current_section = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Check if it's a heading
        style_name = para.style.name
        is_heading = 'Heading' in style_name
        is_bullet = para.style.name == 'List Bullet' or para.style.name == 'List Number'
        
        # Title (centered, bold, no heading style)
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER and para.runs and para.runs[0].bold and not result["title"]:
            result["title"] = text
            continue
            
        # Main section headings (Heading 1 or Heading 2 with Roman numerals)
        if is_heading and ('Heading 1' in style_name or ('Heading 2' in style_name and re.match(r'^[IVX]+\.', text))):
            # Save any pending bullets
            if collecting_bullets and bullet_collector:
                if in_style_notes:
                    result["style_notes"] = bullet_collector.copy()
                elif current_subsection:
                    current_subsection["items"] = bullet_collector.copy()
                bullet_collector = []
                collecting_bullets = False
                
            # Check if this is Style Notes
            if "Style Notes" in text:
                save_current_section()
                in_style_notes = True
                continue
                
            # Save current section
            save_current_section()
            in_style_notes = False
            
            # Determine section type
            section_type = "standard"
            if "Introduction" in text:
                section_type = "introduction"
            elif "Statement of Facts" in text:
                section_type = "facts"
            elif "Conclusion" in text:
                section_type = "conclusion"
            elif re.match(r'^[IVX]+\.', text):  # Roman numeral arguments
                section_type = "argument"
                
            current_section = {
                "type": section_type,
                "heading": text,
                "content": []
            }
            current_subsection = None
            continue
            
        # Handle bullet points
        if is_bullet:
            bullet_collector.append(text)
            collecting_bullets = True
            continue
            
        # If we hit a non-bullet after collecting bullets, save them
        if collecting_bullets and bullet_collector:
            if in_style_notes:
                result["style_notes"] = bullet_collector.copy()
            elif current_subsection:
                current_subsection["items"] = bullet_collector.copy()
            bullet_collector = []
            collecting_bullets = False
            
        # Process content within sections
        if in_style_notes:
            # Style notes are being collected as bullets
            continue
        elif current_section:
            # Check for field labels
            field_match = re.match(r'^(Hook|Theme|Preview|Organization|Summary|Final Theme):\s*(.+)', text)
            if field_match:
                current_section["content"].append({
                    "type": "field",
                    "label": field_match.group(1),
                    "value": field_match.group(2).strip()
                })
                current_subsection = None
            # Check for subsection headers
            elif text in ["Key Facts to Emphasize:", "Bad Facts to Address:", "Fact Themes:", 
                         "Structure:", "Key Authorities:", "Fact Integration:", 
                         "Counter-Argument Response:", "Specific Relief:"]:
                current_subsection = {
                    "type": "list",
                    "label": text.rstrip(':'),
                    "items": []
                }
                current_section["content"].append(current_subsection)
            # Regular paragraph
            else:
                current_section["content"].append({
                    "type": "paragraph",
                    "text": text
                })
                current_subsection = None
    
    # Save any remaining section or style notes
    if collecting_bullets and bullet_collector:
        if in_style_notes:
            result["style_notes"] = bullet_collector.copy()
        elif current_subsection:
            current_subsection["items"] = bullet_collector.copy()
    save_current_section()
    
    return result

def parse_docx_to_json(doc):
    """Parse a DOCX document back to the original JSON structure (legacy format)"""
    sections_data = parse_docx_to_sections(doc)
    
    # Convert sections format to original JSON structure
    result = {
        "title": sections_data["title"],
        "introduction": {
            "hook": "",
            "theme": "",
            "preview": ""
        },
        "fact_section": {
            "organization": "",
            "key_facts_to_emphasize": [],
            "bad_facts_to_address": [],
            "fact_themes": []
        },
        "arguments": [],
        "conclusion": {
            "specific_relief": [],
            "final_theme": ""
        },
        "style_notes": []
    }
    
    # Process each section
    for section in sections_data["sections"]:
        if section["type"] == "introduction":
            for item in section["content"]:
                if item["type"] == "field":
                    if item["label"] == "Hook":
                        result["introduction"]["hook"] = item["value"]
                    elif item["label"] == "Theme":
                        result["introduction"]["theme"] = item["value"]
                    elif item["label"] == "Preview":
                        result["introduction"]["preview"] = item["value"]
                        
        elif section["type"] == "facts":
            for item in section["content"]:
                if item["type"] == "field" and item["label"] == "Organization":
                    result["fact_section"]["organization"] = item["value"]
                elif item["type"] == "list":
                    if item["label"] == "Key Facts to Emphasize":
                        result["fact_section"]["key_facts_to_emphasize"] = item["items"]
                    elif item["label"] == "Bad Facts to Address":
                        result["fact_section"]["bad_facts_to_address"] = item["items"]
                    elif item["label"] == "Fact Themes":
                        result["fact_section"]["fact_themes"] = item["items"]
                        
        elif section["type"] == "argument":
            arg = {
                "heading": section["heading"],
                "summary": "",
                "structure": [],
                "key_authorities": [],
                "fact_integration": [],
                "counter_argument_response": []
            }
            
            for item in section["content"]:
                if item["type"] == "field" and item["label"] == "Summary":
                    arg["summary"] = item["value"]
                elif item["type"] == "list":
                    if item["label"] == "Structure":
                        arg["structure"] = item["items"]
                    elif item["label"] == "Key Authorities":
                        for auth_text in item["items"]:
                            arg["key_authorities"].append(_parse_authority(auth_text))
                    elif item["label"] == "Fact Integration":
                        for fact_text in item["items"]:
                            arg["fact_integration"].append(_parse_fact_integration(fact_text))
                    elif item["label"] == "Counter-Argument Response":
                        for counter_text in item["items"]:
                            arg["counter_argument_response"].append(_parse_counter_argument(counter_text))
                            
            result["arguments"].append(arg)
            
        elif section["type"] == "conclusion":
            for item in section["content"]:
                if item["type"] == "field" and item["label"] == "Final Theme":
                    result["conclusion"]["final_theme"] = item["value"]
                elif item["type"] == "list" and item["label"] == "Specific Relief":
                    result["conclusion"]["specific_relief"] = item["items"]
    
    # Add style notes from the sections data
    result["style_notes"] = sections_data.get("style_notes", [])
    
    return result

def _parse_authority(text):
    """Parse key authority text back to structured format"""
    # Try to match pattern: "Case Name – Citation: Principle (Why it matters)"
    pattern = r'^(.+?)\s*–\s*(.+?):\s*(.+?)\s*\((.+?)\)$'
    match = re.match(pattern, text)
    
    if match:
        return {
            "case_name": match.group(1).strip(),
            "citation": match.group(2).strip(),
            "principle": match.group(3).strip(),
            "why_it_matters": match.group(4).strip()
        }
    else:
        # Return as simple string if pattern doesn't match
        return text

def _parse_fact_integration(text):
    """Parse fact integration text back to structured format"""
    # Try to match pattern: "Fact – Relevance"
    pattern = r'^(.+?)\s*–\s*(.+)$'
    match = re.match(pattern, text)
    
    if match:
        return {
            "fact": match.group(1).strip(),
            "relevance": match.group(2).strip()
        }
    else:
        return text

def _parse_counter_argument(text):
    """Parse counter-argument text back to structured format"""
    # Try to match pattern: "Opposing argument → Response (Strategic value)"
    pattern = r'^(.+?)\s*→\s*(.+?)\s*\((.+?)\)$'
    match = re.match(pattern, text)
    
    if match:
        return {
            "opposing_argument": match.group(1).strip(),
            "response": match.group(2).strip(),
            "strategic_value": match.group(3).strip()
        }
    else:
        return text

@app.get("/health")
def health_check():
    return {"status": "ok"}

@app.post("/generate-outline/")
async def generate_outline(request: Request):
    data = await request.json()
    doc = Document()
    set_doc_style(doc)

    title = doc.add_paragraph(data['title'])
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(f"Hook: {data['introduction']['hook']}")
    doc.add_paragraph(f"Theme: {data['introduction']['theme']}")
    doc.add_paragraph(f"Preview: {data['introduction']['preview']}")

    doc.add_heading("Statement of Facts", level=1)
    doc.add_paragraph(f"Organization: {data['fact_section']['organization']}")

    doc.add_paragraph("Key Facts to Emphasize:")
    add_bullet_points(doc, data['fact_section']['key_facts_to_emphasize'])

    doc.add_paragraph("Bad Facts to Address:")
    add_bullet_points(doc, data['fact_section']['bad_facts_to_address'])

    doc.add_paragraph("Fact Themes:")
    add_bullet_points(doc, data['fact_section']['fact_themes'])

    for arg in data['arguments']:
        doc.add_heading(arg['heading'], level=2)
        doc.add_paragraph(f"Summary: {arg['summary']}")

        doc.add_paragraph("Structure:")
        add_bullet_points(doc, arg['structure'])

        doc.add_paragraph("Key Authorities:")
        for auth in arg['key_authorities']:
            if isinstance(auth, str):
                doc.add_paragraph(auth, style='List Bullet')
            else:
                summary = f"{auth['case_name']} – {auth['citation']}: {auth['principle']} ({auth['why_it_matters']})"
                doc.add_paragraph(summary, style='List Bullet')

        doc.add_paragraph("Fact Integration:")
        for fact in arg['fact_integration']:
            if isinstance(fact, str):
                doc.add_paragraph(fact, style='List Bullet')
            else:
                summary = f"{fact['fact']} – {fact['relevance']}"
                doc.add_paragraph(summary, style='List Bullet')

        doc.add_paragraph("Counter-Argument Response:")
        for reply in arg['counter_argument_response']:
            if isinstance(reply, str):
                doc.add_paragraph(reply, style='List Bullet')
            else:
                summary = f"{reply['opposing_argument']} → {reply['response']} ({reply['strategic_value']})"
                doc.add_paragraph(summary, style='List Bullet')

    doc.add_heading("Conclusion", level=1)
    if data['conclusion']['specific_relief']:
        doc.add_paragraph("Specific Relief:")
        add_bullet_points(doc, data['conclusion']['specific_relief'])
    doc.add_paragraph(f"Final Theme: {data['conclusion']['final_theme']}")

    doc.add_heading("Style Notes", level=1)
    add_bullet_points(doc, data['style_notes'])

    tmp = NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return FileResponse(tmp.name, filename="Legal_Outline.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.post("/parse-outline/")
async def parse_outline(file: UploadFile = File(...)):
    """Parse an uploaded DOCX file back to JSON structure"""
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="File must be a .docx document")
    
    try:
        # Read the uploaded file
        contents = await file.read()
        
        # Load the document
        doc = Document(io.BytesIO(contents))
        
        # Parse the document - legacy format by default
        parsed_data = parse_docx_to_json(doc)
        
        return JSONResponse(content=parsed_data)
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error parsing document: {str(e)}")

@app.post("/parse-outline-sections/")
async def parse_outline_sections(file: UploadFile = File(...)):
    """Parse an uploaded DOCX file into sections for sequential processing"""
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="File must be a .docx document")
    
    try:
        # Read the uploaded file
        contents = await file.read()
        
        # Load the document
        doc = Document(io.BytesIO(contents))
        
        # Parse the document into sections
        parsed_data = parse_docx_to_sections(doc)
        
        return JSONResponse(content=parsed_data)
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error parsing document: {str(e)}")